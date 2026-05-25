from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "绘图尝试_高质量版"
MANIFEST = OUT / "figure_manifest.csv"
DOCX = OUT / "地铁环控系统论文高质量图集.docx"


PALETTE = {
    "ink": "1F2937",
    "muted": "64748B",
    "blue": "1D4ED8",
    "light_blue": "EAF2FF",
    "line": "CBD5E1",
}


def set_run_font(run, east_asia: str = "宋体", latin: str = "Times New Roman") -> None:
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def rgb(hex_color: str) -> RGBColor:
    h = hex_color.replace("#", "")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.replace("#", ""))
    tc_pr.append(shd)


def set_cell_text(
    cell,
    text: str,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    font_size: float = 9,
) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text))
    set_run_font(r)
    r.font.size = Pt(font_size)
    r.bold = bold


def set_cell_width(cell, width_cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def chapter_short(text: str) -> str:
    text = str(text)
    return text.split("、", 1)[0] if "、" in text else text[:2]


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.9)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(10.5)
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Times New Roman"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        styles[name].font.color.rgb = rgb(PALETTE["ink"])


def add_cover(doc: Document, manifest: pd.DataFrame) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(80)
    r = p.add_run("地铁车站环控系统容量优化论文高质量图集")
    set_run_font(r, east_asia="黑体")
    r.font.size = Pt(22)
    r.bold = True
    r.font.color.rgb = rgb(PALETTE["ink"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("MATLAB 论文级制图版")
    set_run_font(r)
    r.font.size = Pt(13)
    r.font.color.rgb = rgb(PALETTE["muted"])

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    rows = [
        ("图件数量", f"{len(manifest)} 张"),
        ("图件格式", "600 dpi PNG；同步导出 PDF"),
        ("数据来源", "data/fuzhou_metro_dongjiekou_2025.csv；output/tables/*.csv；output/models/*.mat"),
        ("质量策略", "统一字体、字号、线宽、配色、单位和图注；DOCX 渲染检查后交付"),
    ]
    for i, (k, v) in enumerate(rows):
        shade_cell(table.cell(i, 0), PALETTE["light_blue"])
        set_cell_text(table.cell(i, 0), k, bold=True)
        set_cell_text(table.cell(i, 1), v, align=WD_ALIGN_PARAGRAPH.LEFT)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("使用建议：")
    set_run_font(r, east_asia="黑体")
    r.bold = True
    r.font.size = Pt(11)
    r2 = p.add_run("优先将本图集中每章的核心图放入论文正文，其余图可用于附录、答辩或结果解释。")
    set_run_font(r2)
    r2.font.size = Pt(10.5)
    doc.add_page_break()


def add_toc(doc: Document, manifest: pd.DataFrame) -> None:
    doc.add_heading("图件目录", level=1)
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(5)
    r = note.add_run("目录采用双栏索引排版；每幅图的详细说明见对应图件下方。")
    set_run_font(r)
    r.font.size = Pt(9)
    r.font.color.rgb = rgb(PALETTE["muted"])

    left = manifest.iloc[:18].reset_index(drop=True)
    right = manifest.iloc[18:].reset_index(drop=True)
    row_count = max(len(left), len(right)) + 1

    table = doc.add_table(rows=row_count, cols=6)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    headers = ["序号", "章", "图名", "序号", "章", "图名"]
    for i, h in enumerate(headers):
        shade_cell(table.cell(0, i), PALETTE["light_blue"])
        set_cell_text(table.cell(0, i), h, bold=True, font_size=8)
        set_cell_width(table.cell(0, i), [0.8, 0.8, 6.7, 0.8, 0.8, 6.7][i])
    set_repeat_table_header(table.rows[0])
    for i in range(row_count - 1):
        cells = table.rows[i + 1].cells
        for col, width in enumerate([0.8, 0.8, 6.7, 0.8, 0.8, 6.7]):
            set_cell_width(cells[col], width)
        if i < len(left):
            row = left.iloc[i]
            set_cell_text(cells[0], int(row["序号"]), font_size=8)
            set_cell_text(cells[1], chapter_short(row["章节"]), font_size=8)
            set_cell_text(cells[2], row["图名"], align=WD_ALIGN_PARAGRAPH.LEFT, font_size=8)
        if i < len(right):
            row = right.iloc[i]
            set_cell_text(cells[3], int(row["序号"]), font_size=8)
            set_cell_text(cells[4], chapter_short(row["章节"]), font_size=8)
            set_cell_text(cells[5], row["图名"], align=WD_ALIGN_PARAGRAPH.LEFT, font_size=8)
    doc.add_page_break()


def add_figure_block(doc: Document, row: pd.Series) -> None:
    title = str(row["图名"])
    note = str(row["说明"])
    png = Path(row["PNG"])

    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    r = p.add_run(title)
    set_run_font(r, east_asia="黑体")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = rgb(PALETTE["ink"])

    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.paragraph_format.keep_with_next = True
    img_p.add_run().add_picture(str(png), width=Inches(6.35))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(3)
    r = cap.add_run(title)
    set_run_font(r)
    r.bold = True
    r.font.size = Pt(10)

    desc = doc.add_paragraph()
    desc.paragraph_format.line_spacing = 1.15
    desc.paragraph_format.space_after = Pt(10)
    r = desc.add_run("图示说明：")
    set_run_font(r, east_asia="黑体")
    r.bold = True
    r.font.size = Pt(10.5)
    r2 = desc.add_run(note)
    set_run_font(r2)
    r2.font.size = Pt(10.5)


def add_figures(doc: Document, manifest: pd.DataFrame) -> None:
    current_section = None
    for idx, row in manifest.iterrows():
        section = str(row["章节"])
        if section != current_section:
            if current_section is not None:
                doc.add_page_break()
            current_section = section
            doc.add_heading(section, level=1)
        add_figure_block(doc, row)
        if (idx + 1) % 2 == 0:
            # Keep pages visually calm; most figures are wide and benefit from breathing room.
            doc.add_paragraph()


def main() -> None:
    manifest = pd.read_csv(MANIFEST, encoding="utf-8")
    doc = Document()
    configure_doc(doc)
    add_cover(doc, manifest)
    add_toc(doc, manifest)
    add_figures(doc, manifest)
    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
