from __future__ import annotations

import csv
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]


def text_of(paragraph: Paragraph) -> str:
    return "".join(run.text for run in paragraph.runs).strip()


def set_run_font(run, size_pt: float = 12, bold: bool = False) -> None:
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = "Times New Roman"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "宋体")
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


def clear_paragraph(paragraph: Paragraph) -> None:
    p_pr = paragraph._p.pPr
    for child in list(paragraph._p):
        if child is not p_pr:
            paragraph._p.remove(child)


def set_paragraph_text(paragraph: Paragraph, text: str, size: float = 12, bold: bool = False) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size, bold)


def format_body(paragraph: Paragraph) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_font(run, 12)


def find_paragraph(doc: Document, target: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if text_of(paragraph) == target:
            return paragraph
    raise ValueError(f"paragraph not found: {target}")


def find_paragraph_starts(doc: Document, prefix: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if text_of(paragraph).startswith(prefix):
            return paragraph
    raise ValueError(f"paragraph not found: {prefix}")


def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        run = new_para.add_run(text)
        set_run_font(run, 12)
    return new_para


def insert_caption_after(paragraph: Paragraph, text: str) -> Paragraph:
    caption = insert_paragraph_after(paragraph, text)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(3)
    for run in caption.runs:
        set_run_font(run, 10.5)
    return caption


def append_paragraph_after_table(table, text: str) -> Paragraph:
    p = OxmlElement("w:p")
    table._tbl.addnext(p)
    para = Paragraph(p, table._parent)
    run = para.add_run(text)
    set_run_font(run, 12)
    format_body(para)
    return para


def create_table_after(paragraph: Paragraph, headers: list[str], rows: list[list[str]], widths_cm: list[float], font_size: float = 10.0):
    table = paragraph._parent.add_table(rows=1, cols=len(headers), width=Cm(sum(widths_cm)))
    paragraph._p.addnext(table._tbl)
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
    for row_values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_values):
            cells[i].text = value
    set_table_widths(table, widths_cm)
    format_table(table, font_size=font_size)
    return table


def set_table_widths(table, widths_cm: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)


def format_table(table, font_size: float = 10.0) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, font_size, False)
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, font_size, True)


def read_csv_dicts(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def replace_text(doc: Document, replacements: dict[str, str]) -> None:
    for paragraph in doc.paragraphs:
        t = paragraph.text
        new_t = t
        for old, new in replacements.items():
            if old in new_t:
                new_t = new_t.replace(old, new)
        if new_t != t:
            set_paragraph_text(paragraph, new_t)
            if paragraph.style.name == "Normal":
                format_body(paragraph)


def round_str(value: str, digits: int = 2) -> str:
    return f"{float(value):.{digits}f}"


def add_dataset_table(doc: Document) -> None:
    anchor = find_paragraph_starts(doc, "原始运行数据来自客流、气象、站内环境和设备运行等多个系统")
    intro = insert_paragraph_after(
        anchor,
        "为增强数据来源和复现口径的可追溯性，本文进一步对原始数据集和建模数据集的基本情况进行汇总。原始数据覆盖 2025 年全年 15 min 时间尺度，经过时间对齐、缺失修复和特征构造后形成建模数据，具体如表3-1所示。",
    )
    format_body(intro)
    caption = insert_caption_after(intro, "表3-1 数据集基本信息与预处理结果")
    rows = [
        ["研究对象", "福州地铁东街口站站台区域通风空调系统", "与容量配置优化对象保持一致"],
        ["时间范围", "2025-01-01 00:00 至 2025-12-31 23:45", "覆盖全年运行周期"],
        ["采样间隔", "15 min", "理论样本数为 35040 条"],
        ["原始字段", "25 个", "包括客流、气象、站内环境、负荷分项和总冷负荷"],
        ["原始缺失记录", "1050 个单元格", "通过插值和邻近有效值补齐"],
        ["建模数据", "35040 条、21 个字段", "清洗后缺失值为 0"],
        ["数据划分", "训练集:验证集:测试集 = 7:2:1", "按时间顺序划分，避免时间泄漏"],
    ]
    table = create_table_after(caption, ["项目", "取值或范围", "说明"], rows, [3.0, 6.2, 6.0], 9.5)
    append_paragraph_after_table(
        table,
        "表3-1 说明本文建模数据在时间范围、采样间隔、字段构成和缺失处理上的基本口径。由于负荷预测属于时间序列问题，本文不采用随机打乱划分，而按时间顺序划分训练集、验证集和测试集，以保持预测评价与实际运行场景一致。",
    )


def add_cluster_ablation_table(doc: Document) -> None:
    anchor = find_paragraph_starts(doc, "同时需要看到，负荷预测误差仍可能来自三类因素")
    intro = insert_paragraph_after(
        anchor,
        "为检验第3章典型日聚类结果是否真正服务于预测模型，本文进一步比较不加入聚类标签和加入聚类标签两种 LSTM 输入设置。两组实验保持序列长度、数据划分和评价指标一致，仅改变是否引入 cluster_label 特征，结果如表4-3所示。",
    )
    format_body(intro)
    caption = insert_caption_after(intro, "表4-3 聚类标签消融实验结果")
    csv_rows = read_csv_dicts(ROOT / "output" / "tables" / "step3_cluster_ablation_metrics.csv")
    rows = []
    for item in csv_rows:
        rows.append(
            [
                item["模型"].replace("LSTM", " LSTM").strip(),
                "是" if item["是否使用聚类标签"] == "1" else "否",
                item["特征数量"],
                round_str(item["均方根误差_kW"], 2),
                round_str(item["平均绝对误差_kW"], 2),
                round_str(item["平均绝对百分比误差_百分比"], 2),
                round_str(item["决定系数R2"], 4),
            ]
        )
    table = create_table_after(
        caption,
        ["模型", "聚类标签", "特征数", "RMSE/kW", "MAE/kW", "MAPE/%", "R²"],
        rows,
        [3.6, 1.8, 1.5, 2.0, 2.0, 2.0, 1.8],
        9.2,
    )
    append_paragraph_after_table(
        table,
        "表4-3 表明，加入聚类标签后 LSTM 的 RMSE 由 10.26 kW 降至 9.40 kW，MAPE 由 4.09% 降至 3.74%，R² 由 0.9666 提高至 0.9720。该结果说明典型日聚类不仅具有运行场景解释意义，也能够为预测模型提供一定的负荷模式信息。",
    )


def add_scenario_demand_table(doc: Document) -> None:
    anchor = find_paragraph_starts(doc, "在本文计算中，设计负荷用于约束冷水机组总容量")
    caption = insert_caption_after(anchor, "表5-2 不同负荷情景下的容量需求")
    csv_rows = read_csv_dicts(ROOT / "output" / "tables" / "step3_scenario_demand.csv")
    rows = []
    for item in csv_rows:
        rows.append(
            [
                item["情景"],
                f"P{int(float(item['分位数']) * 100)}",
                round_str(item["总冷负荷_kW"], 2),
                round_str(item["冷机需求_kW"], 2),
                round_str(item["风机需求_kW"], 2),
                round_str(item["水泵需求_kW"], 2),
                round_str(item["AHU风量需求_m3_h"], 0),
            ]
        )
    table = create_table_after(
        caption,
        ["情景", "分位数", "总冷负荷/kW", "冷机/kW", "风机/kW", "水泵/kW", "AHU风量/(m³/h)"],
        rows,
        [2.0, 1.5, 2.4, 2.0, 2.0, 2.0, 3.1],
        8.8,
    )
    append_paragraph_after_table(
        table,
        "表5-2 将预测负荷转换为典型、峰值和极端三类容量需求。本文选取 P99 极端情景作为容量优化的主要校核边界，是为了在降低过度冗余的同时保留对高负荷工况的覆盖能力；P50 和 P95 情景则用于解释设备在常规与高峰运行区间的匹配情况。",
    )


def add_sensitivity_table(doc: Document) -> None:
    anchor = find_paragraph_starts(doc, "不过，既有车站改造还需要进一步结合现场条件进行校核")
    intro = insert_paragraph_after(
        anchor,
        "此外，容量优化结果并非只由单一目标函数决定，还受到设计分位数、安全系数、TOPSIS 权重、电价和预测误差等参数影响。为说明推荐方案的适用边界，本文对主要敏感因素及工程校核要求进行归纳，如表6-3所示。",
    )
    format_body(intro)
    caption = insert_caption_after(intro, "表6-3 容量优化敏感性与工程校核建议")
    csv_rows = read_csv_dicts(ROOT / "output" / "tables" / "step4_sensitivity_analysis.csv")
    concise = {
        "设计分位数": ["extreme / P99", "分位数提高会增大装机容量并降低欠配风险"],
        "安全系数": ["冷机1.09；风机/水泵1.07；AHU1.05", "系数提高可增强可靠性，但会削弱成本和冗余率降幅"],
        "TOPSIS权重": ["成本0.55；冗余0.45", "成本权重高偏低成本，冗余权重高偏紧凑容量"],
        "电价": ["0.85 元/kWh", "电价越高，部分负荷高效运行价值越大"],
        "预测误差": ["LSTM RMSE 与 6.55% 最小极端裕量比较", "负向峰值误差会侵蚀安全裕量，需用实测峰值校准"],
    }
    rows = []
    for item in csv_rows:
        current, influence = concise.get(
            item["敏感性类型"], [item["当前设置"], item["预期影响"]]
        )
        rows.append([item["敏感性类型"], current, influence, item["风险等级"]])
    table = create_table_after(
        caption,
        ["敏感因素", "当前设置", "影响方向", "风险等级"],
        rows,
        [2.4, 4.8, 6.2, 1.6],
        8.0,
    )
    append_paragraph_after_table(
        table,
        "表6-3 表明，预测误差对极端工况裕量影响最大，应在工程落地前结合实测峰值负荷进行校准；设计分位数、安全系数和 TOPSIS 权重属于中等风险因素，需要在施工图设计或设备选型阶段开展复算。电价主要影响生命周期成本估算，对推荐容量边界的影响相对较低。",
    )


def add_conclusion_limitations(doc: Document) -> None:
    anchor = find_paragraph_starts(doc, "综上，本文提出的基于负荷特性的地铁车站通风空调系统容量配置优化方法")
    p1 = insert_paragraph_after(
        anchor,
        "同时，本文研究仍存在一定局限。第一，数据集主要用于验证方法链条，工程应用前仍需结合真实 BMS 设备运行数据、AFC 客流数据、气象数据和站内环境监测数据进行校准。第二，本文容量优化侧重设备容量和台数组合，能耗计算采用简化 COP/PLR 和变频曲线，尚未展开完整风管阻力、水泵扬程、管网平衡、厂家设备性能曲线和 N+1 备用逻辑校核。第三，TOPSIS 权重、设计分位数和安全系数会影响推荐方案，后续应结合项目需求开展多情景敏感性分析。",
    )
    format_body(p1)
    p2 = insert_paragraph_after(
        p1,
        "后续研究可从三个方向继续深化：一是扩展到多车站、多线路和不同气候区样本，验证模型泛化能力；二是将容量配置优化与运行控制策略联合考虑，进一步评估全年节能潜力；三是引入舒适性约束、室内空气品质约束和设备厂家选型数据，使容量配置结果更接近施工图阶段的工程设计要求。",
    )
    format_body(p2)


def renumber_tables(doc: Document) -> None:
    replacements = {
        "Pearson 相关分析结果如表3-1所示": "Pearson 相关分析结果如表3-2所示",
        "表3-1 关键影响因素相关性排序": "表3-2 关键影响因素相关性排序",
        "结果如表3-2所示": "结果如表3-3所示",
        "表3-2 负荷分项平均占比": "表3-3 负荷分项平均占比",
        "表3-2 用平均负荷和占比量化": "表3-3 用平均负荷和占比量化",
        "如表3-3所示": "如表3-4所示",
        "表3-3 不同聚类数的轮廓系数": "表3-4 不同聚类数的轮廓系数",
        "表3-3 用平均轮廓系数比较": "表3-4 用平均轮廓系数比较",
        "表5-2 经济性计算参数": "表5-3 经济性计算参数",
        "表5-2 列出了": "表5-3 列出了",
        "表5-3 TOPSIS排序结果": "表5-4 TOPSIS排序结果",
        "表5-3 展示了": "表5-4 展示了",
        "如表5-4所示": "如表5-5所示",
        "表5-4 基准方案与优化方案主要指标": "表5-5 基准方案与优化方案主要指标",
        "表5-4 直接比较": "表5-5 直接比较",
        "由表5-4可知": "由表5-5可知",
    }
    replace_text(doc, replacements)


def create_toc_paragraph(text: str, page: str, level: int, doc: Document) -> OxmlElement:
    p = OxmlElement("w:p")
    para = Paragraph(p, doc)
    para.style = "Normal"
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.line_spacing = 1.0
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    if level == 1:
        para.paragraph_format.left_indent = Inches(0.0)
        font_size = 12
        bold = True
    elif level == 2:
        para.paragraph_format.left_indent = Inches(0.32)
        font_size = 11.5
        bold = False
    else:
        para.paragraph_format.left_indent = Inches(0.64)
        font_size = 11
        bold = False
    para.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.35), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
    )
    run = para.add_run(text)
    set_run_font(run, font_size, bold)
    run = para.add_run("\t" + page)
    set_run_font(run, font_size, bold)
    return p


def create_toc_title(doc: Document) -> OxmlElement:
    p = OxmlElement("w:p")
    para = Paragraph(p, doc)
    para.style = "Normal"
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(12)
    run = para.add_run("目    次")
    set_run_font(run, 18, True)
    return p


def create_page_break_paragraph(doc: Document) -> OxmlElement:
    p = OxmlElement("w:p")
    para = Paragraph(p, doc)
    para.add_run().add_break(WD_BREAK.PAGE)
    return p


def replace_static_toc(doc: Document, pages: dict[str, str]) -> None:
    toc_title = find_paragraph(doc, "目    次")
    actual_heading = None
    for paragraph in doc.paragraphs:
        if paragraph._p is toc_title._p:
            continue
        if text_of(paragraph) == "1  绪论" and paragraph.style.name.startswith("Heading"):
            actual_heading = paragraph
            break
    if actual_heading is None:
        raise ValueError("actual chapter 1 heading not found")

    node = toc_title._p
    section_break = actual_heading._p.getprevious()
    while node is not None and node is not section_break:
        next_node = node.getnext()
        node.getparent().remove(node)
        node = next_node

    entries_page_1 = [
        ("摘    要", pages["摘要"], 1),
        ("Abstract:", pages["Abstract"], 1),
        ("1  绪论", pages["1"], 1),
        ("1.1 研究背景", pages["1.1"], 2),
        ("1.2 研究目的与意义", pages["1.2"], 2),
        ("1.3 国内外研究现状", pages["1.3"], 2),
        ("1.4 研究内容与技术路线", pages["1.4"], 2),
        ("1.5 本章小结", pages["1.5"], 2),
        ("2  地铁车站通风空调系统负荷特性与容量配置理论基础", pages["2"], 1),
        ("2.1 地铁车站通风空调系统组成", pages["2.1"], 2),
        ("2.2 站台区域空调负荷构成", pages["2.2"], 2),
        ("2.3 负荷影响因素分析", pages["2.3"], 2),
        ("2.4 典型负荷模式识别理论", pages["2.4"], 2),
        ("2.5 容量配置优化问题描述", pages["2.5"], 2),
        ("2.6 多目标优化与综合评价方法", pages["2.6"], 2),
        ("2.7 本章小结", pages["2.7"], 2),
        ("3  数据预处理与负荷特性分析", pages["3"], 1),
        ("3.1 数据清洗与预处理", pages["3.1"], 2),
        ("3.2 基于 Pearson 相关系数的关键影响因素筛选", pages["3.2"], 2),
        ("3.3 负荷分项分解", pages["3.3"], 2),
        ("3.4 基于 K-Means 的典型负荷曲线聚类", pages["3.4"], 2),
        ("3.5 本章小结", pages["3.5"], 2),
        ("4  基于 LSTM 的地铁车站空调负荷预测模型", pages["4"], 1),
        ("4.1 负荷预测问题描述", pages["4.1"], 2),
        ("4.2 预测输入特征构建", pages["4.2"], 2),
        ("4.3 LSTM 神经网络模型构建", pages["4.3"], 2),
    ]
    entries_page_2 = [
        ("4.4 BP 神经网络对比模型", pages["4.4"], 2),
        ("4.5 模型评价指标", pages["4.5"], 2),
        ("4.6 预测结果与对比分析", pages["4.6"], 2),
        ("4.7 本章小结", pages["4.7"], 2),
        ("5  地铁车站通风空调系统容量配置多目标优化", pages["5"], 1),
        ("5.1 容量配置优化对象与决策变量", pages["5.1"], 2),
        ("5.2 设计负荷确定", pages["5.2"], 2),
        ("5.3 优化目标函数", pages["5.3"], 2),
        ("5.3.1 全生命周期成本目标", pages["5.3.1"], 3),
        ("5.3.2 容量冗余率目标", pages["5.3.2"], 3),
        ("5.4 工程约束条件", pages["5.4"], 2),
        ("5.5 NSGA-II 多目标优化模型", pages["5.5"], 2),
        ("5.6 基于 TOPSIS 的综合最优方案选择", pages["5.6"], 2),
        ("5.7 优化结果概述", pages["5.7"], 2),
        ("5.8 本章小结", pages["5.8"], 2),
        ("6  优化结果分析", pages["6"], 1),
        ("6.1 基准方案与优化方案对比", pages["6.1"], 2),
        ("6.2 容量冗余率分析", pages["6.2"], 2),
        ("6.3 全生命周期成本分析", pages["6.3"], 2),
        ("6.4 年运行能耗与节能效果分析", pages["6.4"], 2),
        ("6.5 设备匹配性分析", pages["6.5"], 2),
        ("6.6 调节灵活性与运行适应性分析", pages["6.6"], 2),
        ("6.7 既有车站改造适配性分析", pages["6.7"], 2),
        ("6.8 本章小结", pages["6.8"], 2),
        ("7  结论", pages["7"], 1),
        ("致  谢", pages["致谢"], 1),
        ("参考文献", pages["参考文献"], 1),
        ("附录1  主要程序文件与数据字段说明", pages["附录1"], 1),
    ]

    insert_before = section_break
    elements = [create_toc_title(doc)]
    elements.extend(create_toc_paragraph(text, page, level, doc) for text, page, level in entries_page_1)
    elements.append(create_page_break_paragraph(doc))
    elements.extend(create_toc_paragraph(text, page, level, doc) for text, page, level in entries_page_2)
    for element in elements:
        insert_before.addprevious(element)


def default_pages() -> dict[str, str]:
    return {
        "摘要": "I",
        "Abstract": "II",
        "1": "1",
        "1.1": "1",
        "1.2": "1",
        "1.3": "2",
        "1.4": "3",
        "1.5": "5",
        "2": "6",
        "2.1": "6",
        "2.2": "7",
        "2.3": "9",
        "2.4": "9",
        "2.5": "10",
        "2.6": "12",
        "2.7": "13",
        "3": "14",
        "3.1": "14",
        "3.2": "18",
        "3.3": "20",
        "3.4": "22",
        "3.5": "24",
        "4": "25",
        "4.1": "25",
        "4.2": "26",
        "4.3": "27",
        "4.4": "29",
        "4.5": "30",
        "4.6": "31",
        "4.7": "35",
        "5": "36",
        "5.1": "36",
        "5.2": "38",
        "5.3": "39",
        "5.3.1": "39",
        "5.3.2": "40",
        "5.4": "41",
        "5.5": "42",
        "5.6": "44",
        "5.7": "45",
        "5.8": "47",
        "6": "48",
        "6.1": "48",
        "6.2": "50",
        "6.3": "50",
        "6.4": "51",
        "6.5": "51",
        "6.6": "52",
        "6.7": "52",
        "6.8": "54",
        "7": "55",
        "致谢": "58",
        "参考文献": "59",
        "附录1": "62",
    }


def main() -> None:
    if len(sys.argv) != 3:
        print("usage: revise_draft08.py <input.docx> <output.docx>")
        raise SystemExit(2)

    src = Path(sys.argv[1])
    dst = Path(sys.argv[2])
    shutil.copyfile(src, dst)

    doc = Document(str(dst))
    renumber_tables(doc)
    add_dataset_table(doc)
    add_cluster_ablation_table(doc)
    add_scenario_demand_table(doc)
    add_sensitivity_table(doc)
    add_conclusion_limitations(doc)
    replace_static_toc(doc, default_pages())
    doc.save(str(dst))
    print(dst)


if __name__ == "__main__":
    main()
