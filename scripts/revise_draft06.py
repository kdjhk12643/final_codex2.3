from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


def text_of(paragraph: Paragraph) -> str:
    return "".join(run.text for run in paragraph.runs).strip()


def set_run_font(run, size_pt: float | None = None, bold: bool = False) -> None:
    run.bold = bold
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    run.font.name = "宋体"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "宋体")
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


def clear_paragraph(paragraph: Paragraph) -> None:
    p_pr = paragraph._p.pPr
    for child in list(paragraph._p):
        if child is not p_pr:
            paragraph._p.remove(child)


def set_paragraph_text(paragraph: Paragraph, text: str, style: str | None = None) -> None:
    clear_paragraph(paragraph)
    if style:
        paragraph.style = style
    run = paragraph.add_run(text)
    set_run_font(run, 12)


def format_body(paragraph: Paragraph) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_font(run, 12)


def format_caption(paragraph: Paragraph) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)
    for run in paragraph.runs:
        set_run_font(run, 10.5)


def format_picture(paragraph: Paragraph) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = Paragraph(new_p, paragraph._parent)
    if style:
        new_paragraph.style = style
    if text:
        run = new_paragraph.add_run(text)
        set_run_font(run, 12)
    return new_paragraph


def insert_paragraph_after_element(element, parent, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    element.addnext(new_p)
    new_paragraph = Paragraph(new_p, parent)
    if text:
        run = new_paragraph.add_run(text)
        set_run_font(run, 12)
    return new_paragraph


def find_paragraph(doc: Document, prefix: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if text_of(paragraph).startswith(prefix):
            return paragraph
    raise ValueError(f"paragraph not found: {prefix}")


def find_paragraph_exact(doc: Document, target: str) -> Paragraph | None:
    for paragraph in doc.paragraphs:
        if text_of(paragraph) == target:
            return paragraph
    return None


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def find_next_table_element(paragraph: Paragraph):
    node = paragraph._p.getnext()
    while node is not None:
        if node.tag == qn("w:tbl"):
            return node
        if node.tag == qn("w:p") and "".join(node.itertext()).strip():
            return None
        node = node.getnext()
    return None


def replace_previous_drawing(caption: Paragraph, image_path: Path, width_in: float = 5.7) -> None:
    node = caption._p.getprevious()
    while node is not None:
        if node.tag == qn("w:p"):
            candidate = Paragraph(node, caption._parent)
            has_drawing = bool(
                node.findall(".//" + qn("w:drawing")) or node.findall(".//" + qn("w:pict"))
            )
            if has_drawing:
                clear_paragraph(candidate)
                format_picture(candidate)
                candidate.add_run().add_picture(str(image_path), width=Inches(width_in))
                return
            if text_of(candidate):
                break
        node = node.getprevious()
    raise ValueError(f"no drawing paragraph before caption: {text_of(caption)}")


def replace_paragraph_with_picture(paragraph: Paragraph, image_path: Path, width_in: float = 5.7) -> None:
    clear_paragraph(paragraph)
    format_picture(paragraph)
    paragraph.add_run().add_picture(str(image_path), width=Inches(width_in))


def insert_figure_after(
    paragraph: Paragraph,
    image_path: Path,
    caption_text: str,
    explanation: str,
    width_in: float = 5.7,
) -> Paragraph:
    pic = insert_paragraph_after(paragraph)
    format_picture(pic)
    pic.add_run().add_picture(str(image_path), width=Inches(width_in))
    cap = insert_paragraph_after(pic, caption_text)
    format_caption(cap)
    exp = insert_paragraph_after(cap, explanation)
    format_body(exp)
    return exp


def insert_body_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_para = insert_paragraph_after(paragraph, text)
    format_body(new_para)
    return new_para


def next_nonempty_paragraph(paragraph: Paragraph) -> Paragraph | None:
    node = paragraph._p.getnext()
    while node is not None:
        if node.tag == qn("w:p"):
            candidate = Paragraph(node, paragraph._parent)
            if text_of(candidate):
                return candidate
        node = node.getnext()
    return None


def add_transition_after_heading(doc: Document, heading_text: str, transition: str) -> None:
    heading = find_paragraph(doc, heading_text)
    nxt = next_nonempty_paragraph(heading)
    if nxt is not None and text_of(nxt).startswith(transition[:12]):
        return
    para = insert_paragraph_after(heading, transition)
    format_body(para)


def make_box_diagram(path: Path, boxes: list[dict], width: int = 2200, height: int = 620) -> None:
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    font_path = Path(r"C:\Windows\Fonts\msyh.ttc")
    if not font_path.exists():
        font_path = Path(r"C:\Windows\Fonts\simhei.ttf")
    title_font = ImageFont.truetype(str(font_path), 58)
    body_font = ImageFont.truetype(str(font_path), 42)
    small_font = ImageFont.truetype(str(font_path), 34)

    border = (63, 105, 190)
    arrow = (70, 85, 105)
    for idx, box in enumerate(boxes):
        x1, y1, x2, y2 = box["rect"]
        fill = box.get("fill", (238, 245, 255))
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, fill=fill, outline=border, width=7)
        lines = box["lines"]
        fonts = [title_font] + [body_font] * (len(lines) - 1)
        total_h = 0
        measured = []
        for line, font in zip(lines, fonts):
            bbox = draw.textbbox((0, 0), line, font=font)
            measured.append((bbox[2] - bbox[0], bbox[3] - bbox[1], font))
            total_h += bbox[3] - bbox[1]
        total_h += 18 * (len(lines) - 1)
        y = y1 + ((y2 - y1) - total_h) / 2
        for line, (tw, th, font) in zip(lines, measured):
            draw.text((x1 + (x2 - x1 - tw) / 2, y), line, fill="black", font=font)
            y += th + 18

        note = box.get("note")
        if note:
            nb = draw.textbbox((0, 0), note, font=small_font)
            draw.text((x1 + (x2 - x1 - (nb[2] - nb[0])) / 2, y2 + 26), note, fill=(55, 65, 80), font=small_font)

        if idx < len(boxes) - 1:
            nx1 = boxes[idx + 1]["rect"][0]
            y_mid = (y1 + y2) // 2
            start = (x2 + 18, y_mid)
            end = (nx1 - 24, y_mid)
            draw.line((start, end), fill=arrow, width=10)
            draw.polygon(
                [(end[0], end[1]), (end[0] - 34, end[1] - 20), (end[0] - 34, end[1] + 20)],
                fill=arrow,
            )
    img.save(path, quality=95)


def generate_assets(asset_dir: Path) -> dict[str, Path]:
    asset_dir.mkdir(parents=True, exist_ok=True)
    lstm_path = asset_dir / "fig4_2_lstm_architecture.png"
    bp_path = asset_dir / "fig4_3_bp_architecture.png"

    make_box_diagram(
        lstm_path,
        [
            {
                "rect": (40, 135, 360, 410),
                "fill": (232, 242, 255),
                "lines": ["序列输入", "16步 × 多特征"],
            },
            {
                "rect": (510, 105, 850, 440),
                "fill": (241, 236, 255),
                "lines": ["LSTM层1", "96个隐含单元"],
                "note": "输出完整序列",
            },
            {
                "rect": (1010, 105, 1350, 440),
                "fill": (241, 236, 255),
                "lines": ["LSTM层2", "48个隐含单元"],
                "note": "输出最后状态",
            },
            {
                "rect": (1510, 135, 1790, 410),
                "fill": (231, 248, 239),
                "lines": ["全连接层", "1维输出"],
            },
            {
                "rect": (1940, 135, 2170, 410),
                "fill": (255, 246, 230),
                "lines": ["回归输出", "总冷负荷/kW"],
            },
        ],
    )

    make_box_diagram(
        bp_path,
        [
            {
                "rect": (40, 135, 410, 410),
                "fill": (232, 242, 255),
                "lines": ["静态特征输入", "客流/气象/环境/时间"],
            },
            {
                "rect": (610, 135, 930, 410),
                "fill": (241, 236, 255),
                "lines": ["隐藏层1", "20个神经元"],
            },
            {
                "rect": (1130, 135, 1450, 410),
                "fill": (241, 236, 255),
                "lines": ["隐藏层2", "10个神经元"],
            },
            {
                "rect": (1660, 135, 2130, 410),
                "fill": (231, 248, 239),
                "lines": ["输出层", "总冷负荷/kW"],
            },
        ],
    )
    return {"lstm": lstm_path, "bp": bp_path}


def find_figure_dir(root: Path) -> Path:
    for child in root.iterdir():
        candidate = child / "figures_png" / "Fig_1_1_research_route.png"
        if candidate.exists():
            return child / "figures_png"
    raise FileNotFoundError("figures_png directory not found")


def move_caption_and_table_after(doc: Document, caption_prefix: str, target_prefix: str) -> Table:
    caption = find_paragraph(doc, caption_prefix)
    table_el = find_next_table_element(caption)
    if table_el is None:
        raise ValueError(f"table not found after {caption_prefix}")
    target = find_paragraph(doc, target_prefix)
    parent = caption._parent
    caption_el = caption._p
    body = caption_el.getparent()
    body.remove(table_el)
    body.remove(caption_el)
    target._p.addnext(caption_el)
    caption_el.addnext(table_el)
    return Table(table_el, parent)


def add_table_explanation(doc: Document, caption_prefix: str, explanation: str) -> None:
    caption = find_paragraph(doc, caption_prefix)
    table_el = find_next_table_element(caption)
    if table_el is None:
        return
    nxt = table_el.getnext()
    if nxt is not None and nxt.tag == qn("w:p"):
        next_para = Paragraph(nxt, caption._parent)
        if text_of(next_para).startswith(explanation[:10]):
            return
    para = insert_paragraph_after_element(table_el, caption._parent, explanation)
    format_body(para)


def add_caption_explanation(doc: Document, caption_prefix: str, explanation: str) -> None:
    caption = find_paragraph(doc, caption_prefix)
    nxt = next_nonempty_paragraph(caption)
    if nxt is not None and text_of(nxt).startswith(explanation[:10]):
        return
    para = insert_paragraph_after(caption, explanation)
    format_body(para)


def main() -> None:
    if len(sys.argv) != 3:
        print("usage: revise_draft06.py <input.docx> <output.docx>")
        raise SystemExit(2)

    src = Path(sys.argv[1])
    dst = Path(sys.argv[2])
    root = src.parent
    dst.parent.mkdir(parents=True, exist_ok=True)
    shutil.copyfile(src, dst)

    asset_dir = root / "output" / "doc_revision_work" / "draft06_assets"
    generated = generate_assets(asset_dir)
    fig_dir = find_figure_dir(root)

    doc = Document(str(dst))

    replacements = {
        "图3-4 负荷分项占比图": "图3-4 负荷分项占比图",
        "表3-1 关键影响因素相关性排序": "表3-1 关键影响因素相关性排序",
        "表3-3 不同聚类数的轮廓系数": "表3-3 不同聚类数的轮廓系数",
        "表5-2 经济性计算参数": "表5-2 经济性计算参数",
        "图4-3 LSTM预测结果图": "图4-4 LSTM 测试集预测结果图",
        "图4-4 BP预测结果图": "图4-5 BP 测试集预测结果图",
        "图4-5 模型RMSE对比图": "图4-6 LSTM 与 BP 预测误差指标对比图",
        "图5-1 NSGA-II容量优化流程图": "图5-1 NSGA-II 与 TOPSIS 容量优化流程图",
    }
    for old_prefix, new_text in replacements.items():
        paragraph = find_paragraph(doc, old_prefix)
        set_paragraph_text(paragraph, new_text)
        format_caption(paragraph)

    placeholder = find_paragraph_exact(doc, "章节之间需要过度文字")
    if placeholder:
        set_paragraph_text(
            placeholder,
            "第4章已经证明 LSTM 预测模型能够较好刻画站台区域负荷变化，本章进一步把预测结果转化为设备容量配置问题。容量优化的核心不再是单纯比较预测误差，而是要在满足设计负荷和工程约束的前提下，确定冷源、风机、水泵和空气处理机组的容量及台数组合，并在经济性与冗余控制之间取得平衡。",
        )
        format_body(placeholder)

    p = find_paragraph_exact(doc, "1）写成一段分号连接；；；；；")
    if p:
        set_paragraph_text(
            p,
            "具体而言，决策变量包括冷水机组单机容量与台数、风机单机容量与台数、水泵单机容量与台数、空气处理机组单台额定风量与台数，分别用于描述设备规格、装机规模和系统处理能力。",
        )
        format_body(p)
        for txt in [
            "1. 冷水机组单机容量；",
            "2. 冷水机组台数；",
            "3. 风机单机容量；",
            "4. 风机台数；",
            "5. 水泵单机容量；",
            "6. 水泵台数；",
            "7. 空气处理机组单台额定风量；",
            "8. 空气处理机组台数。",
        ]:
            target = find_paragraph_exact(doc, txt)
            if target:
                remove_paragraph(target)

    p = find_paragraph_exact(doc, "1）")
    if p:
        remove_paragraph(p)

    p = find_paragraph_exact(doc, "1）；；；；；；；；；；；；；；；；；；；；")
    if p:
        set_paragraph_text(p, "上述流程可概括为以下八个步骤。")
        format_body(p)

    p = find_paragraph_exact(doc, "模型结构如下：   增加模型框图")
    if p:
        set_paragraph_text(p, "结合图4-2，本文 LSTM 模型结构可概括为以下五个层次：")
        format_body(p)

    add_transition_after_heading(
        doc,
        "2  地铁车站通风空调系统负荷特性与容量配置理论基础",
        "第1章已经明确本文以负荷特性分析、负荷预测和容量配置优化为主线。为了使后续模型具有明确的工程对象，本章先从地铁车站通风空调系统组成、站台区域负荷构成和容量优化基本理论入手，说明各类设备与负荷需求之间的对应关系。",
    )
    add_transition_after_heading(
        doc,
        "3  数据预处理与负荷特性分析",
        "第2章给出了通风空调系统负荷构成和容量配置的理论基础，但工程优化仍需要可靠的数据支撑。本章转入实测运行数据层面，先完成数据清洗和特征构造，再通过相关性分析、分项分解和典型日聚类识别站台区域负荷的主要变化规律。",
    )
    add_transition_after_heading(
        doc,
        "4  基于 LSTM 的地铁车站空调负荷预测模型",
        "第3章已经筛选出影响站台区域总冷负荷的关键因素，并识别出不同典型日负荷模式。本章在此基础上构建负荷预测模型，将历史负荷、客流、环境和时间特征组织为序列输入，并通过 LSTM 与 BP 对比分析验证时序建模的必要性。",
    )
    add_transition_after_heading(
        doc,
        "6  优化结果分析",
        "第5章已经得到基于 NSGA-II 和 TOPSIS 的推荐容量配置方案。本章进一步从工程应用角度对推荐方案进行解释，重点比较其相对于传统基准方案在容量冗余、生命周期成本、年运行能耗、设备匹配性和改造适配性方面的变化。",
    )
    add_transition_after_heading(
        doc,
        "7  结论",
        "前六章依次完成了理论分析、数据处理、负荷预测、容量优化和工程评价。为归纳本文研究工作，本章对主要结论进行总结，并说明该方法对地铁车站通风空调系统容量配置和节能改造的参考意义。",
    )

    # Replace or insert figure assets.
    fig1_caption = find_paragraph(doc, "图1-1 技术路线图")
    replace_paragraph_with_picture(fig1_caption, fig_dir / "Fig_1_1_research_route.png", 5.9)
    route_caption = next_nonempty_paragraph(fig1_caption)
    set_paragraph_text(route_caption, "图1-1 技术路线图")
    format_caption(route_caption)

    for caption_prefix, image_name, width in [
        ("图3-2 缺失值统计图", "Fig_2_8_missing_values.png", 5.5),
        ("图3-3 Pearson相关性排序图", "Fig_3_1_pearson_ranking.png", 5.5),
        ("图3-4 负荷分项占比图", "Fig_2_6_component_ratio.png", 5.3),
        ("图3-5 典型日负荷曲线聚类图", "Fig_3_4_cluster_centers.png", 5.6),
        ("图4-1 预测样本构造示意图", "Fig_4_1_lstm_sequence.png", 5.6),
        ("图4-4 LSTM 测试集预测结果图", "Fig_4_4_lstm_prediction.png", 5.7),
        ("图4-5 BP 测试集预测结果图", "Fig_4_5_bp_prediction.png", 5.7),
        ("图4-6 LSTM 与 BP 预测误差指标对比图", "Fig_4_6_model_metrics.png", 5.2),
        ("图5-1 NSGA-II 与 TOPSIS 容量优化流程图", "Fig_5_4_optimization_flow.png", 5.8),
        ("图5-2 Pareto前沿图", "Fig_5_5_pareto_front.png", 5.3),
        ("图6-1 基准方案与优化方案对比图", "Fig_5_7_scheme_comparison.png", 5.2),
    ]:
        caption = find_paragraph(doc, caption_prefix)
        replace_previous_drawing(caption, fig_dir / image_name, width)

    lstm_para = find_paragraph(doc, "图4-2 LSTM网络结构图")
    replace_paragraph_with_picture(lstm_para, generated["lstm"], 5.9)
    lstm_caption = insert_paragraph_after(lstm_para, "图4-2 LSTM 网络结构图")
    format_caption(lstm_caption)
    lstm_exp = insert_paragraph_after(
        lstm_caption,
        "图4-2 展示了本文 LSTM 预测模型的数据流向：16 个时间步的多变量序列首先进入第一层 LSTM 以提取连续时序特征，再由第二层 LSTM 汇聚为最后时刻的状态表达，最终经全连接层输出下一时刻总冷负荷。两层隐藏单元数分别为 96 和 48，既保留了地铁负荷的短时惯性，又控制了模型复杂度。",
    )
    format_body(lstm_exp)

    bp_insert_after = find_paragraph(doc, "本文 BP 模型采用两层隐藏层结构")
    bp_exp = insert_figure_after(
        bp_insert_after,
        generated["bp"],
        "图4-3 BP 神经网络结构图",
        "图4-3 给出了 BP 对比模型的前馈结构。该模型以同一时刻的客流、气象、站内环境和时间特征为输入，经过 20 个神经元和 10 个神经元的两层隐藏层后输出总冷负荷预测值。与 LSTM 不同，BP 模型不显式保留历史状态，因此主要用于检验静态非线性映射与时序记忆结构之间的预测差异。",
        5.6,
    )

    # Move misplaced result/parameter tables to a more natural position.
    move_caption_and_table_after(doc, "表5-2 经济性计算参数", "运行电费根据年运行能耗")
    move_caption_and_table_after(doc, "表5-3 TOPSIS排序结果", "TOPSIS 贴近度可表示为")

    figure_explanations = {
        "图1-1 技术路线图": "图1-1 将本文研究过程概括为数据预处理、负荷特性分析、典型日聚类、负荷预测、容量优化和工程评价六个连续环节。该路线体现了从运行数据到设备容量配置的递进关系，后续各章均围绕这一链路展开。",
        "图2-1 地铁车站通风空调系统组成示意图": "图2-1 说明了冷源侧、风系统、水系统和空气处理设备之间的连接关系。后续容量优化并不是只调整单一冷水机组容量，而是需要同时考虑冷量产生、输配和空气处理能力的匹配。",
        "图3-1 站台总冷负荷时间序列图": "图3-1 反映了全年站台总冷负荷的连续变化特征，可用于观察负荷峰谷、季节差异和异常波动。该图说明负荷预测模型需要同时处理周期性变化和局部随机扰动。",
        "图3-2 缺失值统计图": "图3-2 用于检查各关键变量的缺失分布。缺失值数量较少或集中于个别变量时，可采用插值和邻近值补齐；若缺失集中在连续时段，则需要在后续建模中保留质量标记，避免对相关性和预测结果造成偏差。",
        "图3-3 Pearson相关性排序图": "图3-3 展示了候选特征与总冷负荷之间的线性相关强度。滞后负荷和客流变量排序靠前，说明站台负荷具有明显时间连续性和客流驱动特征，这为第4章采用序列预测模型提供了依据。",
        "图3-4 负荷分项占比图": "图3-4 直观展示了人员、新风、围护结构和设备散热四类负荷的平均占比。围护结构和设备散热占比较高，说明车站存在稳定基础负荷，容量配置不能只依据客流峰值进行估算。",
        "图3-5 典型日负荷曲线聚类图": "图3-5 给出了不同典型日负荷曲线的形态差异。各类曲线在峰值时刻、峰谷差和持续高负荷时间上存在差别，说明优化方案需要兼顾高峰响应能力和中低负荷运行效率。",
        "图4-1 预测样本构造示意图": "图4-1 表明单个预测样本由连续 16 个 15 min 时间步组成，对应过去 4 h 的运行状态。该结构使模型能够利用负荷滞后、客流变化和环境条件的连续信息预测下一时刻负荷。",
        "图4-4 LSTM 测试集预测结果图": "图4-4 对比了测试集真实负荷与 LSTM 预测负荷。两条曲线在高峰和低谷区间均保持较好一致性，说明 LSTM 能够较稳定地跟踪站台负荷的日内波动和峰值变化。",
        "图4-5 BP 测试集预测结果图": "图4-5 显示 BP 模型能够拟合总体负荷水平，但在负荷快速上升或下降区间更容易出现滞后和偏差。这与 BP 模型缺少显式时序记忆结构有关。",
        "图4-6 LSTM 与 BP 预测误差指标对比图": "图4-6 从 RMSE、MAE 和 MAPE 三类误差指标对两种模型进行比较。LSTM 各项误差均低于 BP，说明引入序列结构后模型对负荷波动和峰值偏差的刻画能力更强。",
        "图5-1 NSGA-II 与 TOPSIS 容量优化流程图": "图5-1 展示了容量优化从设备编码到综合排序的计算流程。NSGA-II 用于搜索低成本、低冗余的 Pareto 候选方案，TOPSIS 用于在候选方案中选出综合表现较好的推荐配置。",
        "图5-2 Pareto前沿图": "图5-2 反映了全生命周期成本与综合容量冗余率之间的权衡关系。前沿上的方案均为非支配解，红色推荐点表示在当前权重和约束条件下兼顾经济性与冗余控制的综合最优方案。",
        "图6-1 基准方案与优化方案对比图": "图6-1 从冷源容量、生命周期成本、综合冗余率和年运行能耗等指标展示优化前后的差异。结果表明优化方案并非单纯压缩容量，而是在满足约束的同时降低过度配置和长期运行成本。",
    }
    for prefix, explanation in figure_explanations.items():
        add_caption_explanation(doc, prefix, explanation)

    table_explanations = {
        "表2-1 站台区域负荷组成及影响因素": "表2-1 将站台区域负荷按来源、影响因素和容量配置意义进行归纳，说明不同负荷分项对峰值需求、基础负荷和空气处理能力的作用不同。",
        "表2-2 容量配置优化对象": "表2-2 明确了容量优化涉及的主要子系统、决策内容和约束条件，为后文建立离散变量组合优化模型提供对象范围。",
        "表3-1 关键影响因素相关性排序": "表3-1 给出了关键影响因素的相关系数排序，数值越接近 1 表示与总冷负荷的同步变化越强。该表用于确定后续预测模型的核心输入变量。",
        "表3-2 负荷分项平均占比": "表3-2 用平均负荷和占比量化各分项贡献，其中围护结构负荷和设备负荷占比较高，是解释站台基础冷负荷的重要数据依据。",
        "表3-3 不同聚类数的轮廓系数": "表3-3 用平均轮廓系数比较不同 K 值的聚类质量。K = 4 时指标最高，且与工程场景解释相符，因此作为本文典型日模式划分结果。",
        "表4-1 LSTM模型主要训练参数": "表4-1 汇总了 LSTM 模型的输入窗口、隐藏单元数、训练轮数、学习率和数据划分比例。这些参数共同决定模型对历史信息长度、非线性表达能力和训练稳定性的控制。",
        "表4-2 负荷预测模型评价指标": "表4-2 从绝对误差、相对误差和拟合优度三个角度比较 LSTM 与 BP 模型。LSTM 在 RMSE、MAE、MAPE 和 R² 上均表现更优。",
        "表5-1 设备候选容量与台数范围": "表5-1 给出了各类设备可选容量和台数范围，说明优化模型采用工程可选规格进行离散搜索，避免得到无法直接选型的连续容量结果。",
        "表5-2 经济性计算参数": "表5-2 列出了生命周期成本计算所需的电价、寿命、折现率、设备单位造价和维护费率。后续 LCC 结果均基于这些参数进行方案间相对比较。",
        "表5-3 TOPSIS排序结果": "表5-3 展示了 Pareto 候选方案的综合排序结果。排名第 1 的方案同时具有较低全生命周期成本和较低综合容量冗余率，因此被选为推荐优化方案。",
        "表5-4 基准方案与优化方案主要指标": "表5-4 直接比较基准方案和优化方案的核心指标，说明优化方案在冷源总容量、生命周期成本和综合冗余率方面均有下降。",
        "表6-1 设备配置明细表": "表6-1 给出了基准方案和优化方案的具体设备组合，便于观察容量下降来自冷水机组、风机、水泵和空气处理机组的协同调整，而不是单一设备削减。",
        "表6-2 基准方案与优化方案综合对比": "表6-2 从投资、成本、冗余率、能耗、负荷率和 COP 等指标综合比较两种方案，说明优化方案同时改善了容量合理性和运行经济性。",
    }
    for prefix, explanation in table_explanations.items():
        add_table_explanation(doc, prefix, explanation)

    # Add a short lead sentence before the first visual/table in section 6.1.
    sec61 = find_paragraph(doc, "6.1 基准方案与优化方案对比")
    nxt = next_nonempty_paragraph(sec61)
    lead = "为直观呈现优化前后的配置差异，本文先从总体指标和设备明细两个层面进行对比，再结合容量冗余、成本和能耗指标展开分析。"
    if nxt is not None and not text_of(nxt).startswith(lead[:10]):
        para = insert_paragraph_after(sec61, lead)
        format_body(para)

    # Keep only real short captions centered; explanatory paragraphs often also
    # begin with "图x-x" or "表x-x" and must remain body text.
    for paragraph in doc.paragraphs:
        txt = text_of(paragraph)
        is_short_caption = (txt.startswith("图") or txt.startswith("表")) and len(txt) <= 38 and "。" not in txt
        if is_short_caption:
            format_caption(paragraph)

    doc.save(str(dst))
    print(dst)


if __name__ == "__main__":
    main()
