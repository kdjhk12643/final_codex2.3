from __future__ import annotations

import math
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from docx.text.paragraph import Paragraph
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
TABLE_DIR = ROOT / "output" / "tables"
FIGURE_DIR = ROOT / "output" / "figures"
INPUT_DOCX = ROOT / "毕业设计论文初稿04.docx"
OUTPUT_DOCX = ROOT / "毕业设计论文初稿04_规范与结果修订最终版.docx"
DIAGRAM_DIR = ROOT / "output" / "doc_revision_work" / "generated_diagrams"


def set_run_font(run, east="宋体", west="Times New Roman", size=None, bold=None):
    run.font.name = west
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), east)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), west)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), west)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def replace_paragraph(paragraph, text, style=None):
    paragraph.clear()
    if style is not None:
        paragraph.style = style
    run = paragraph.add_run(text)
    return run


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def clear_table(table):
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)


def ensure_rows(table, row_count):
    while len(table.rows) < row_count:
        table.add_row()
    while len(table.rows) > row_count:
        table._tbl.remove(table.rows[-1]._tr)


def set_cell_border(cell, top=None, bottom=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = tcBorders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tcBorders.append(element)
        element.set(qn("w:val"), "nil")
    for edge, value in (("top", top), ("bottom", bottom)):
        if not value:
            continue
        element = tcBorders.find(qn("w:" + edge))
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), str(value))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_data(table, rows):
    ensure_rows(table, len(rows))
    for r_idx, row_data in enumerate(rows):
        for c_idx, value in enumerate(row_data):
            cell = table.cell(r_idx, c_idx)
            cell.text = str(value)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            for p in cell.paragraphs:
                if c_idx == 0 and r_idx != 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                for run in p.runs:
                    set_run_font(run, size=10.5, bold=(r_idx == 0))
            if r_idx == 0:
                set_cell_border(cell, top=12, bottom=8)
            elif r_idx == len(rows) - 1:
                set_cell_border(cell, bottom=12)
            else:
                set_cell_border(cell)


def percent(value, digits=2):
    return f"{value * 100:.{digits}f}%"


def money(value):
    return f"{int(round(value)):,}"


def replace_picture(paragraph, image_path):
    if not image_path.exists():
        return
    blips = paragraph._p.xpath(".//a:blip")
    if not blips:
        return
    r_id = blips[0].get(qn("r:embed"))
    image_part = paragraph.part.related_parts[r_id]
    image_part._blob = image_path.read_bytes()


def add_picture_to_paragraph(paragraph, image_path, width_inches=5.35):
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))


def load_cn_font(size=30, bold=False):
    candidates = [
        Path(r"C:\Windows\Fonts\msyhbd.ttc" if bold else r"C:\Windows\Fonts\msyh.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\simsun.ttc"),
    ]
    for font in candidates:
        if font.exists():
            return ImageFont.truetype(str(font), size)
    return ImageFont.load_default()


def draw_centered(draw, xy, text, font, fill=(30, 30, 30)):
    x0, y0, x1, y1 = xy
    bbox = draw.multiline_textbbox((0, 0), text, font=font, spacing=6, align="center")
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    draw.multiline_text(((x0 + x1 - tw) / 2, (y0 + y1 - th) / 2), text, font=font, fill=fill, spacing=6, align="center")


def draw_box(draw, xy, text, font, fill=(245, 248, 252), outline=(54, 83, 111)):
    draw.rounded_rectangle(xy, radius=14, fill=fill, outline=outline, width=3)
    draw_centered(draw, xy, text, font)


def draw_arrow(draw, start, end, fill=(60, 60, 60), width=4):
    draw.line([start, end], fill=fill, width=width)
    angle = math.atan2(end[1] - start[1], end[0] - start[0])
    length = 14
    left = (end[0] - length * math.cos(angle - math.pi / 6), end[1] - length * math.sin(angle - math.pi / 6))
    right = (end[0] - length * math.cos(angle + math.pi / 6), end[1] - length * math.sin(angle + math.pi / 6))
    draw.polygon([end, left, right], fill=fill)


def make_horizontal_flow(path, labels, width=1400, height=260):
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_cn_font(28, bold=True)
    font = load_cn_font(25)
    margin = 45
    y0, y1 = 85, 170
    gap = 28
    box_w = (width - 2 * margin - gap * (len(labels) - 1)) / len(labels)
    for i, label in enumerate(labels):
        x0 = margin + i * (box_w + gap)
        x1 = x0 + box_w
        draw_box(draw, (x0, y0, x1, y1), label, font)
        if i < len(labels) - 1:
            draw_arrow(draw, (x1 + 4, (y0 + y1) / 2), (x1 + gap - 6, (y0 + y1) / 2))
    img.save(path)


def make_system_diagram(path):
    img = Image.new("RGB", (1400, 560), "white")
    draw = ImageDraw.Draw(img)
    font = load_cn_font(28)
    small = load_cn_font(24)
    boxes = {
        "冷水机组": (80, 95, 280, 180),
        "冷冻水泵": (360, 95, 560, 180),
        "空气处理机组": (640, 95, 900, 180),
        "站台送风机": (980, 95, 1200, 180),
        "站台区域": (1030, 285, 1250, 390),
        "站台排风机": (700, 365, 920, 450),
        "冷却水泵": (360, 365, 560, 450),
    }
    for label, xy in boxes.items():
        draw_box(draw, xy, label, font if len(label) < 6 else small)
    draw_arrow(draw, (280, 137), (360, 137))
    draw_arrow(draw, (560, 137), (640, 137))
    draw_arrow(draw, (900, 137), (980, 137))
    draw_arrow(draw, (1090, 180), (1120, 285))
    draw_arrow(draw, (1030, 350), (920, 407))
    draw_arrow(draw, (700, 407), (560, 407))
    draw_arrow(draw, (460, 365), (460, 180))
    draw_centered(draw, (65, 245, 1260, 280), "冷源供冷、空气处理、送排风与站台热湿负荷形成闭环匹配关系", small, fill=(80, 80, 80))
    img.save(path)


def make_diagrams():
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)
    make_system_diagram(DIAGRAM_DIR / "fig2_1_system.png")
    make_horizontal_flow(
        DIAGRAM_DIR / "fig4_1_sample.png",
        ["连续16个\n时间步特征", "序列输入\n矩阵", "LSTM\n时序建模", "下一时刻\n总冷负荷"],
        width=1200,
    )
    make_horizontal_flow(
        DIAGRAM_DIR / "fig5_1_nsga.png",
        ["初始化\n种群", "解码\n设备方案", "约束\n校核", "目标函数\n计算", "非支配\n排序", "选择交叉\n变异", "Pareto\n解集"],
        width=1550,
    )


def replace_placeholder_with_diagram(doc, marker, image_path, caption, width_inches=5.35):
    for p in doc.paragraphs:
        if marker in p.text:
            add_picture_to_paragraph(p, image_path, width_inches=width_inches)
            return insert_paragraph_after(p, caption)
    return None


def style_caption(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(6)
    for run in paragraph.runs:
        set_run_font(run, size=10.5)


def style_formula(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = None
    for run in paragraph.runs:
        set_run_font(run, west="Times New Roman", size=12)


def apply_document_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    normal.font.size = Pt(12)

    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(3.0)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)
        section.footer_distance = Cm(1.5)
        sect_pr = section._sectPr
        doc_grid = sect_pr.find(qn("w:docGrid"))
        if doc_grid is None:
            doc_grid = OxmlElement("w:docGrid")
            sect_pr.append(doc_grid)
        doc_grid.set(qn("w:type"), "linesAndChars")
        doc_grid.set(qn("w:linePitch"), "452")
        doc_grid.set(qn("w:charSpace"), "512")

    for p in doc.paragraphs:
        text = p.text.strip()
        style_name = p.style.name if p.style else ""
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)

        if style_name == "Heading 1":
            p.paragraph_format.page_break_before = True
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if text[:1].isdigit() else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = None
            for run in p.runs:
                set_run_font(run, size=16, bold=True)
        elif style_name == "Heading 2":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = None
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            for run in p.runs:
                set_run_font(run, size=15, bold=True)
        elif style_name == "Heading 3":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.first_line_indent = None
            for run in p.runs:
                set_run_font(run, size=14, bold=True)
        elif text.startswith(("图", "表")) and len(text) < 40:
            style_caption(p)
        elif text.startswith(("Q_", "C_", "R_", "LCC", "f1", "f2", "RMSE", "MAE", "MAPE", "R²", "J =", "s(i)", "Loss")) or "=" in text and len(text) < 90 and any(sym in text for sym in ["Σ", "sqrt", ">=", "min", "sin", "cos"]):
            style_formula(p)
        else:
            if p.runs:
                for run in p.runs:
                    set_run_font(run, size=12)
            if text and not text.startswith(("关键词", "Key words", "Abstract", "Student", "学", "题目", "桂林理工大学", "GUILIN", "本科毕业", "设计(", "指 导 教 师")):
                p.paragraph_format.first_line_indent = Pt(24)


def main():
    make_diagrams()
    doc = Document(INPUT_DOCX)

    pred = pd.read_csv(TABLE_DIR / "step3_prediction_metrics.csv")
    comp = pd.read_csv(TABLE_DIR / "step2_load_component_ratio.csv")
    pearson = pd.read_csv(TABLE_DIR / "step2_pearson_features.csv")
    silhouette = pd.read_csv(TABLE_DIR / "step2_cluster_silhouette.csv")
    scheme = pd.read_csv(TABLE_DIR / "step4_scheme_evaluation.csv")
    redundancy = pd.read_csv(TABLE_DIR / "step4_subsystem_redundancy.csv")
    topsis = pd.read_csv(TABLE_DIR / "step4_topsis_ranking.csv").sort_values("TOPSIS得分", ascending=False)

    lstm = pred[pred["模型"] == "LSTM"].iloc[0]
    bp = pred[pred["模型"] == "BP"].iloc[0]
    base = scheme[scheme["方案"] == "基准方案"].iloc[0]
    opt = scheme[scheme["方案"] == "优化方案"].iloc[0]

    cost_drop = 1 - opt["生命周期成本"] / base["生命周期成本"]
    energy_drop = 1 - opt["年能耗_kWh"] / base["年能耗_kWh"]
    invest_drop = 1 - opt["初投资"] / base["初投资"]
    redundancy_drop = 1 - opt["综合冗余率"] / base["综合冗余率"]
    energy_save = base["年能耗_kWh"] - opt["年能耗_kWh"]
    annual_fee_save = energy_save * 0.85

    replacements = {
        36: (
            "本文以地铁车站站台区域通风空调系统为研究对象，针对传统容量配置中设备容量偏大、系统长期低负荷运行和全生命周期成本较高等问题，构建了基于负荷特性的容量配置优化方法。研究基于福州地铁东街口站 2025 年全年 15 min 粒度运行数据，首先完成时间对齐、缺失值处理、标准化和滞后特征构造；随后采用 Pearson 相关系数法筛选关键影响因素，并将总冷负荷分解为人员负荷、新风负荷、围护结构负荷和设备散热负荷；在此基础上利用 K-Means 聚类识别典型日负荷模式，建立两层 LSTM 神经网络负荷预测模型，并与 BP 神经网络进行对比。最后，以全生命周期成本最小和综合容量冗余率最小为目标，采用 NSGA-II 求解 Pareto 候选方案，并通过 TOPSIS 方法确定推荐容量配置方案。结果表明，LSTM 模型在测试集上的 RMSE、MAE、MAPE 和 R² 分别为 "
            f"{lstm['均方根误差_kW']:.2f} kW、{lstm['平均绝对误差_kW']:.2f} kW、{lstm['平均绝对百分比误差_百分比']:.2f}% 和 {lstm['决定系数R2']:.4f}，预测精度优于 BP 模型；推荐优化方案使冷源总容量由 {int(base['总制冷容量_kW'])} kW 降至 {int(opt['总制冷容量_kW'])} kW，综合容量冗余率由 {percent(base['综合冗余率'])} 降至 {percent(opt['综合冗余率'])}，全生命周期成本降低 {percent(cost_drop)}，年运行能耗降低 {percent(energy_drop)}。研究结果可为地铁车站通风空调系统容量配置、节能设计和既有车站改造方案比选提供参考。"
        ),
        42: (
            "This thesis studies the ventilation and air-conditioning system of a metro station platform area and proposes a load-characteristic-based capacity configuration optimization method. Using a 15-minute annual operating dataset of Fuzhou Metro Dongjiekou Station in 2025, the study completes data preprocessing, Pearson correlation analysis, load component decomposition and K-Means daily load clustering. A two-layer LSTM model is then established for cooling-load prediction and compared with a BP neural network. Based on the predicted load demand, a multi-objective capacity optimization model is built with lifecycle cost and composite capacity redundancy as objectives. NSGA-II is used to generate Pareto candidate schemes, and TOPSIS is used to select the recommended configuration. The results show that the LSTM model achieves "
            f"RMSE = {lstm['均方根误差_kW']:.2f} kW, MAE = {lstm['平均绝对误差_kW']:.2f} kW, MAPE = {lstm['平均绝对百分比误差_百分比']:.2f}% and R² = {lstm['决定系数R2']:.4f} on the test set. The recommended scheme reduces total cooling capacity from {int(base['总制冷容量_kW'])} kW to {int(opt['总制冷容量_kW'])} kW, composite redundancy from {percent(base['综合冗余率'])} to {percent(opt['综合冗余率'])}, lifecycle cost by {percent(cost_drop)}, and annual operating energy consumption by {percent(energy_drop)}. The proposed method can support energy-saving design and retrofit decision-making for metro station HVAC systems."
        ),
        49: "1.1 研究背景",
        54: "1.2 研究目的与意义",
        62: "1.3 国内外研究现状",
        67: "1.4 研究内容与技术路线",
        64: "在负荷预测方面，传统方法包括线性回归、时间序列模型、灰色预测和 BP 神经网络等。这类方法结构相对简单、实现方便，但在处理非线性、多变量和长短期时序依赖时存在一定局限。近年来，LSTM、GRU、Transformer 等深度学习模型逐渐用于建筑和交通场站空调负荷预测，其中 LSTM 通过门控结构保存历史状态，适合描述负荷惯性、客流高峰持续和环境参数滞后影响等问题[1]。",
        65: "在容量配置和节能优化方面，传统工程设计通常依据设计峰值负荷、远期客流和经验安全裕度确定设备容量，方法可靠但容易造成装机容量偏大。多目标优化方法能够同时考虑经济性、冗余率和工程约束，NSGA-II 可输出 Pareto 非劣解集[2]，TOPSIS 可进一步根据成本和冗余等指标进行综合排序[3]，适合用于设备容量组合方案比选。",
        87: "地铁车站通风空调系统主要用于维持车站公共区、设备管理用房及相关区域的热湿环境和空气品质，其系统配置应满足地铁工程和暖通空调设计规范的基本要求[4-5]。本文研究对象为地铁车站站台区域通风空调系统，重点关注与站台区域冷负荷承担和空气处理相关的主要设备，包括冷源系统、送排风系统、水系统和空气处理机组等。地铁车站通风空调系统组成示意图如下图所示：",
        181: "经过上述处理后，数据集中的缺失值和明显异常值均已完成修复，最终形成覆盖 2025 年全年、共 35040 条 15 min 粒度样本的建模数据集。该数据集既保留了负荷峰谷变化、季节变化和日周期结构，也降低了采集噪声对后续分析结果的干扰。",
        183: "图3-3 Pearson相关性排序图",
        184: "为了识别站台区域总冷负荷的主要驱动因素，本文在完成时间对齐和标准化后，对候选特征与目标负荷进行 Pearson 相关分析[6]。Pearson 相关系数用于衡量两个连续变量之间的线性相关强度，其计算公式为：",
        190: f"从结果可以看出，历史负荷项与当前总冷负荷的相关性最高，其中 load_lag_1 的相关系数达到 {pearson.iloc[0]['Pearson相关系数']:.4f}，说明相邻 15 min 内负荷变化具有明显连续性。这一结果与地铁车站围护结构蓄热、设备运行惯性和空调系统控制滞后相一致，也为后续 LSTM 模型引入序列窗口提供了统计依据。",
        192: "气象和站内环境变量中，站台温度、太阳辐射等变量与总冷负荷具有较高相关性，表明外部气象边界和站内热湿环境共同影响空调系统负荷。相比之下，周末标识和部分周期编码变量的单变量相关性较弱，但这并不意味着其无价值；周期特征更多用于帮助模型识别日内运行阶段和工作日/周末差异，其作用可能体现在与客流、气象变量的非线性组合中。",
        193: "综合统计排序和工程解释性，本文选取历史负荷、进出站客流、站台人数、二氧化碳浓度、站台温度和太阳辐射等变量作为后续预测模型的核心输入。同时保留周期编码特征和聚类标签作为辅助输入，用于增强模型对不同运行模式的识别能力。",
        197: "图3-4 负荷分项占比图",
        202: "结果表明，围护结构负荷在总负荷中占比最高，达到 50.00%，说明站台区域冷负荷主要受围护结构传热和地下环境耦合作用影响。设备负荷占比为 31.96%，是第二大负荷来源，表明站台照明、电扶梯、屏蔽门及其他机电设备运行会持续产生大量热量。人员负荷占比为 6.46%，新风负荷占比为 5.21%，二者虽然相对较小，但与客流高峰和空气品质要求直接相关，对负荷波动和通风能力校核具有重要影响。",
        206: "图3-5 典型日负荷曲线聚类图",
        209: "K-Means 聚类通过最小化簇内平方误差来划分样本[7]，其目标函数为：",
        212: "聚类数 K 的选择直接影响模式划分效果。若 K 过小，典型场景会被合并，难以区分工作日和周末、峰值日和低负荷日；若 K 过大，则可能把偶然波动识别为独立模式，降低工程解释性。本文在 K = 2 至 K = 4 范围内搜索，并采用平均轮廓系数评价聚类质量。单个样本的轮廓系数为：",
        216: f"从统计指标看，K = 4 时平均轮廓系数最高，为 {silhouette[silhouette['K值'] == 4].iloc[0]['平均轮廓系数']:.4f}，说明四类划分既能保持较好的簇间分离度，也便于识别工作日双峰高负荷、工作日平稳中负荷、周末午后单峰和低客流运行等差异。因此，本文采用“统计质量满足要求、工程解释性优先”的原则确定 K = 4。",
        230: "本章基于福州地铁东街口站 2025 年全年 15 min 粒度运行数据，对站台区域通风空调系统负荷进行了系统预处理和特性分析。首先完成时间戳统一、等间隔对齐、缺失值分级填补、异常值校核、Z-Score 标准化、周期编码和滞后特征构造，形成覆盖全年、共 35040 条样本的建模数据集。随后利用 Pearson 相关系数并结合工程机理筛选关键影响因素，结果表明历史负荷、客流、二氧化碳浓度、站台温度和太阳辐射均与总冷负荷密切相关。进一步的负荷分项分解表明，围护结构负荷和设备散热负荷是站台区域基础负荷的主要来源。最后，采用 K-Means 对日负荷曲线进行聚类，并综合轮廓系数与工程解释性确定 K = 4，识别出工作日双峰高负荷、工作日平稳中负荷、周末午后单峰和低客流低负荷四类典型模式，为下一章负荷预测模型输入设计和后续容量配置优化提供了更充分的数据依据。",
        251: "图4-2 LSTM网络结构图：Sequence Input → LSTM(96) → LSTM(48) → Fully Connected → Regression Output",
        260: "2. 第一层 LSTM：隐藏单元数为 96，输出完整序列；",
        261: "3. 第二层 LSTM：隐藏单元数为 48，输出最后一个时间步状态；",
        264: "训练过程中采用 Adam 优化算法，最大训练轮数为 120，批量大小为 32，初始学习率为 0.0008。为避免梯度爆炸，设置梯度阈值为 1；为提高训练稳定性，学习率采用分段下降策略，验证集若连续多轮未改善则提前停止训练。",
        289: "图4-3 LSTM预测结果图",
        291: "图4-4 BP预测结果图",
        293: "图4-5 模型RMSE对比图",
        295: "LSTM 模型和 BP 模型在测试集上的评价结果如表4-2所示。",
        296: "表4-2 负荷预测模型评价指标",
        298: f"由表4-2可知，LSTM 模型在各项评价指标上均优于 BP 神经网络。LSTM 的 RMSE 为 {lstm['均方根误差_kW']:.2f} kW，低于 BP 模型的 {bp['均方根误差_kW']:.2f} kW；LSTM 的 MAE 为 {lstm['平均绝对误差_kW']:.2f} kW，低于 BP 模型的 {bp['平均绝对误差_kW']:.2f} kW；LSTM 的 MAPE 为 {lstm['平均绝对百分比误差_百分比']:.2f}%，低于 BP 模型的 {bp['平均绝对百分比误差_百分比']:.2f}%；LSTM 的 R² 为 {lstm['决定系数R2']:.4f}，高于 BP 模型的 {bp['决定系数R2']:.4f}。",
        299: f"从误差降低幅度来看，LSTM 相比 BP 模型的 RMSE 降低约 {percent(1 - lstm['均方根误差_kW'] / bp['均方根误差_kW'])}，MAE 降低约 {percent(1 - lstm['平均绝对误差_kW'] / bp['平均绝对误差_kW'])}，MAPE 降低约 {percent(1 - lstm['平均绝对百分比误差_百分比'] / bp['平均绝对百分比误差_百分比'])}。在本文数据集和特征构造条件下，LSTM 对站台区域空调负荷变化具有更好的拟合能力和泛化能力。由于容量优化更关注高峰负荷和典型运行区间，RMSE 的明显降低表明 LSTM 对大偏差样本的控制能力更强，可为第5章设计负荷确定和容量约束设置提供更可靠的输入基础。",
        302: "同时需要看到，负荷预测误差仍可能来自三类因素：一是构造年度数据集仍需进一步用真实 BMS、AFC 和气象实测数据校准；二是列车运行间隔、屏蔽门启闭、设备启停策略等变量未完全纳入输入特征；三是传感器采样延迟和局部异常值会影响历史负荷项的稳定性。因此，本文在第5章确定设计负荷时采用高分位负荷、预测误差裕量和工程安全系数联合校核，而不是直接以单一预测峰值作为设备选型依据。",
        304: f"本章基于第三章筛选得到的关键影响因素，建立了地铁车站站台区域空调负荷预测模型。首先构建了包含客流、气象、站内环境、时间特征和历史负荷项的预测输入矩阵，并按 7:2:1 的比例划分训练集、验证集和测试集。随后建立两层 LSTM 神经网络模型，并采用 BP 神经网络作为对比模型。测试结果表明，LSTM 模型的 RMSE 为 {lstm['均方根误差_kW']:.2f} kW，MAPE 为 {lstm['平均绝对百分比误差_百分比']:.2f}%，R² 为 {lstm['决定系数R2']:.4f}，整体预测精度优于 BP 模型。说明 LSTM 能够较好捕捉地铁车站空调负荷的时序变化规律，可作为后续容量配置优化的负荷输入基础。",
        327: "容量配置优化需要首先确定设计校核负荷。本文以第4章 LSTM 负荷预测结果为基础，将负荷需求划分为典型情景、峰值情景和极端情景，并采用 P99 极端分位数作为容量优化的主要设计边界。为降低预测误差导致的欠配风险，在冷机 P99 需求基础上叠加 LSTM 测试集 RMSE 作为预测误差裕量。该处理方式体现了“预测驱动、极端校核、误差修正”的原则：预测结果用于反映运行趋势，高分位负荷用于约束高峰需求，误差裕量用于覆盖模型可能低估的峰值工况。",
        329: "Q_d = Q_{0.99} + RMSE_LSTM",
        330: "式中，Q_d 为容量优化设计负荷，Q_{0.99} 为预测负荷对应的 99% 分位需求，RMSE_LSTM 为 LSTM 模型测试集均方根误差。本文计算得到冷机 P99 需求为 413.28 kW，叠加 9.58 kW 预测误差裕量后，容量优化采用的冷源校核需求为 422.86 kW。",
        331: "在本文计算中，设计负荷用于约束冷水机组总容量、风机总容量、水泵总容量和空气处理机组总处理能力，使优化方案能够满足高负荷工况下的运行需求。对应的极端情景需求为：总冷负荷 403.72 kW、冷机需求 413.28 kW、风机需求 27.17 kW、水泵需求 24.83 kW、AHU 风量需求 120308.07 m³/h。",
        388: "5.6 基于 TOPSIS 的综合最优方案选择",
        407: "根据 NSGA-II 和 TOPSIS 计算结果，本文得到一组 Pareto 候选方案，并从中选出综合最优容量配置方案。基准方案与优化方案的主要结果如表5-4所示。",
        408: "表5-4 基准方案与优化方案主要指标",
        410: f"由表5-4可知，优化方案的冷源总容量由基准方案的 {int(base['总制冷容量_kW'])} kW 降低至 {int(opt['总制冷容量_kW'])} kW，综合容量冗余率由 {percent(base['综合冗余率'])} 降低至 {percent(opt['综合冗余率'])}。同时，全生命周期成本由 {base['生命周期成本'] / 10000:.2f} 万元降低至 {opt['生命周期成本'] / 10000:.2f} 万元。说明本文建立的容量配置优化模型能够在满足负荷需求和工程约束的前提下，有效降低设备容量冗余和生命周期经济成本。",
        411: f"从优化逻辑看，容量由 {int(base['总制冷容量_kW'])} kW 降至 {int(opt['总制冷容量_kW'])} kW 并非简单削减设备规模，而是在预测负荷、设计负荷、安全系数和子系统匹配约束共同作用下，对传统基准方案中不必要冗余的再分配。综合容量冗余率降至 {percent(opt['综合冗余率'])} 表明优化方案仍保留一定安全裕度，但已显著减少长期低负荷运行风险。全生命周期成本降低则同时来自设备初投资减少和部分负荷运行效率改善。",
        416: f"本章在第4章负荷预测结果基础上，构建了地铁车站通风空调系统容量配置多目标优化模型。首先明确了冷水机组、风机、水泵和空气处理机组的容量及台数组合为优化对象，并建立了以全生命周期成本最小和综合容量冗余率最小为目标的优化函数。随后，结合负荷满足、设备容量边界、台数组合和系统匹配等工程约束，采用 NSGA-II 算法求解 Pareto 解集，并利用 TOPSIS 方法选取综合最优方案。计算结果表明，优化方案相较基准方案能够将冷源总容量由 {int(base['总制冷容量_kW'])} kW 降至 {int(opt['总制冷容量_kW'])} kW，将综合容量冗余率由 {percent(base['综合冗余率'])} 降至 {percent(opt['综合冗余率'])}，并使全生命周期成本降低 {percent(cost_drop)}，为后续工程评价提供了基础。",
        422: "表6-1 设备配置明细表",
        425: "基准方案与优化方案的主要评价指标如表6-2所示。",
        426: "表6-2 基准方案与优化方案综合对比",
        427: f"由表6-2可知，优化方案在冷源总容量、初投资、全生命周期成本、容量冗余率和年运行能耗等方面均优于基准方案。优化方案冷源总容量由 {int(base['总制冷容量_kW'])} kW 降低至 {int(opt['总制冷容量_kW'])} kW，说明在基于负荷预测和系统匹配约束的前提下，可以减少传统设计中不必要的容量冗余。与此同时，优化方案全生命周期成本降低 {percent(cost_drop)}，年运行能耗降低 {percent(energy_drop)}，表明容量配置优化不仅能够降低设备投资，也能够改善系统部分负荷运行效率。",
        429: f"容量冗余率是评价设备容量配置合理性的重要指标。基准方案的综合容量冗余率为 {percent(base['综合冗余率'])}，说明传统配置方案中设备容量明显高于实际负荷需求。较高的容量冗余虽然能够提高峰值工况下的安全裕度，但也会导致设备长期运行在较低负荷率区间，降低系统运行效率。",
        430: f"优化方案的综合容量冗余率为 {percent(opt['综合冗余率'])}，相比基准方案下降 {percent(redundancy_drop)}。该结果表明，本文优化模型能够在满足设计负荷和安全系数约束的基础上，有效压缩过度裕量，使设备配置更加接近实际负荷需求。",
        431: f"从工程角度看，容量冗余率并非越低越好。过低的冗余率可能导致系统在极端高温、高客流或设备性能衰减条件下缺乏必要安全裕度。本文优化方案保留约 {percent(opt['综合冗余率'], 1)} 的综合容量冗余，既避免了传统方案中 30% 以上的过度冗余，又为实际运行中的负荷波动和设备调节保留了一定余量，具有较好的工程可接受性。",
        434: f"从初投资看，基准方案初投资为 {money(base['初投资'])} 元，优化方案初投资为 {money(opt['初投资'])} 元，降低 {money(base['初投资'] - opt['初投资'])} 元，降幅为 {percent(invest_drop)}。初投资降低主要来自设备总容量减少和台数组合优化。由于优化方案减少了过度配置，冷水机组、风机、水泵和空气处理机组的容量选型更加贴近负荷需求，设备购置成本随之下降。",
        435: f"从全生命周期成本看，基准方案为 {money(base['生命周期成本'])} 元，优化方案为 {money(opt['生命周期成本'])} 元，降低 {money(base['生命周期成本'] - opt['生命周期成本'])} 元，降幅为 {percent(cost_drop)}。全生命周期成本降幅低于初投资降幅，说明优化方案的主要经济收益来自设备配置收缩和部分负荷效率改善，运行期能耗收益仍需结合更精细的逐时效率模型进一步校核。",
        438: f"按照年节电量 {energy_save:.1f} kWh 和电价 0.85 元/kWh 估算，优化方案每年可节约运行电费约 {annual_fee_save / 10000:.2f} 万元。若进一步考虑设备维护费用随容量下降而减少，则优化方案的长期经济收益会更加明显。该结果说明，容量配置优化不仅体现在一次性初投资下降，也会在系统长期运营阶段持续产生节能和降费效果。",
        440: f"计算结果表明，基准方案年运行能耗为 {base['年能耗_kWh']:.1f} kWh，优化方案年运行能耗为 {opt['年能耗_kWh']:.1f} kWh，年节电量为 {energy_save:.1f} kWh，节能率为 {percent(energy_drop)}。",
        450: f"从冷源侧看，优化方案冷源总容量由 {int(base['总制冷容量_kW'])} kW 降低至 {int(opt['总制冷容量_kW'])} kW，仍能够满足设计负荷与安全系数要求，说明基准方案存在一定冷源容量过剩。",
        458: f"此外，优化方案保留约 {percent(opt['综合冗余率'], 1)} 的综合容量冗余，为客流短时波动、气象条件变化和设备性能衰减提供了一定缓冲空间。因此，从运行适应性角度看，优化方案兼顾了节能性和安全裕度。",
        464: "6.8 本章小结",
        465: f"本章对基准方案和优化方案进行了对比分析，并从容量冗余率、全生命周期成本、年运行能耗、设备匹配性、调节灵活性和既有车站改造适配性等方面评价了优化效果。结果表明，优化方案冷源总容量由 {int(base['总制冷容量_kW'])} kW 降低至 {int(opt['总制冷容量_kW'])} kW，综合容量冗余率由 {percent(base['综合冗余率'])} 降低至 {percent(opt['综合冗余率'])}，全生命周期成本降低 {percent(cost_drop)}，年运行能耗降低 {percent(energy_drop)}。说明基于负荷特性的容量配置优化方法能够在满足工程约束的前提下，有效降低设备容量冗余、减少生命周期成本并提升系统节能潜力。与此同时，本文也指出当前能耗评价仍存在简化，后续应结合全年运行数据和设备效率曲线进一步验证。",
        468: "本文以地铁车站站台区域通风空调系统为研究对象，围绕容量配置偏大、系统长期低负荷运行和运行能耗较高等问题，构建了“数据预处理、负荷特性分析、负荷预测建模、容量配置优化、工程评价”的研究流程。基于福州地铁东街口站 2025 年全年 15 分钟粒度运行数据，采用 Pearson 相关分析、负荷分项分解、K-Means 聚类、LSTM 神经网络、NSGA-II 多目标优化和 TOPSIS 综合评价等方法，开展了地铁车站通风空调系统容量配置优化研究。主要结论如下。",
        469: f"第一，地铁车站站台区域空调负荷受历史负荷、客流和环境参数共同影响，负荷序列表现出明显的时序连续性。Pearson 相关分析结果表明，load_lag_1 与总冷负荷的相关系数达到 {pearson.iloc[0]['Pearson相关系数']:.4f}，说明前一时段负荷对当前负荷具有很强指示作用。站台人数、进站客流和出站客流与总冷负荷的相关系数分别为 {pearson[pearson['特征'] == '站台人数'].iloc[0]['Pearson相关系数']:.4f}、{pearson[pearson['特征'] == '进站客流'].iloc[0]['Pearson相关系数']:.4f} 和 {pearson[pearson['特征'] == '出站客流'].iloc[0]['Pearson相关系数']:.4f}，表明客流变化是站台区域负荷波动的重要驱动因素。二氧化碳浓度、站台温度和太阳辐射等环境变量也与负荷具有较高相关性，可作为负荷预测模型的重要输入。",
        470: "第二，负荷分项结果表明，站台区域总冷负荷并非仅由人员活动决定，而是由围护结构传热、设备散热、新风处理和人员散热共同构成。其中，围护结构负荷平均占比为 50.00%，设备负荷平均占比为 31.96%，人员负荷平均占比为 6.46%，新风负荷平均占比为 5.21%。这说明地铁车站站台区域负荷具有较强的地下空间和机电设备运行特征，容量配置优化需要综合考虑多类负荷来源。",
        471: "第三，基于 K-Means 的典型负荷曲线聚类能够有效识别站台区域负荷模式。本文在 K = 2 至 K = 4 范围内进行聚类分析，综合考虑轮廓系数和工程解释性，最终选取 K = 4 作为典型负荷模式划分结果。聚类结果可概括为工作日双峰高负荷型、工作日平稳中负荷型、周末午后单峰型和低客流低负荷型。该结果说明地铁车站空调负荷具有明显的日内周期和运行场景差异，为负荷预测和容量配置提供了模式识别依据。",
        472: f"第四，LSTM 神经网络能够较好预测地铁车站站台区域空调负荷。本文以客流、气象、站内环境、时间特征和历史负荷项作为输入，构建两层 LSTM 负荷预测模型，并与 BP 神经网络进行对比。测试结果表明，LSTM 模型 RMSE 为 {lstm['均方根误差_kW']:.2f} kW，MAE 为 {lstm['平均绝对误差_kW']:.2f} kW，MAPE 为 {lstm['平均绝对百分比误差_百分比']:.2f}%，R² 为 {lstm['决定系数R2']:.4f}，整体预测精度优于 BP 模型。说明 LSTM 对地铁车站负荷时序特征具有较强表达能力，适合作为容量配置优化的负荷预测工具。",
        473: f"第五，基于 NSGA-II 和 TOPSIS 的容量配置优化方法能够有效降低系统容量冗余和全生命周期成本。本文以全生命周期成本最小和综合容量冗余率最小为目标，建立冷水机组、风机、水泵和空气处理机组的离散容量组合优化模型，并通过 TOPSIS 方法从 Pareto 解集中选取综合最优方案。结果表明，优化方案冷源总容量由基准方案的 {int(base['总制冷容量_kW'])} kW 降低至 {int(opt['总制冷容量_kW'])} kW，综合容量冗余率由 {percent(base['综合冗余率'])} 降低至 {percent(opt['综合冗余率'])}，全生命周期成本由 {base['生命周期成本'] / 10000:.2f} 万元降低至 {opt['生命周期成本'] / 10000:.2f} 万元。",
        474: f"第六，优化方案具有较好的经济性、节能性和工程适用性。与基准方案相比，优化方案初投资降低 {percent(invest_drop)}，全生命周期成本降低 {percent(cost_drop)}，年运行能耗降低 {percent(energy_drop)}。同时，优化方案保留约 {percent(opt['综合冗余率'], 1)} 的综合容量冗余，在减少过度配置的同时仍具备一定安全裕度。设备容量和台数组合均基于工程常用规格选取，具有较好的设备匹配性和既有车站改造适配性。",
    }

    for index, text in replacements.items():
        replace_paragraph(doc.paragraphs[index], text)

    set_table_data(
        doc.tables[1],
        [["特征", "Pearson R"]]
        + [[row["特征"], f"{row['Pearson相关系数']:.4f}"] for _, row in pearson.head(10).iterrows()],
    )
    set_table_data(
        doc.tables[2],
        [["分项", "平均负荷/kW", "占比"]]
        + [[row["分项负荷"], f"{row['平均负荷_kW']:.2f}", percent(row["占比"])] for _, row in comp.iterrows()],
    )
    set_table_data(
        doc.tables[3],
        [["K", "平均轮廓系数"]]
        + [[int(row["K值"]), f"{row['平均轮廓系数']:.4f}"] for _, row in silhouette.iterrows()],
    )
    set_table_data(
        doc.tables[4],
        [
            ["参数", "取值", "说明"],
            ["输入序列长度", "16", "对应4小时历史信息"],
            ["第一层LSTM单元数", "96", "提取主要时序特征"],
            ["第二层LSTM单元数", "48", "输出最终时序状态"],
            ["最大训练轮数", "120", "控制训练上限"],
            ["批量大小", "32", "小批量训练"],
            ["初始学习率", "0.0008", "Adam优化器参数"],
            ["梯度阈值", "1", "抑制梯度爆炸"],
            ["验证耐心", "50", "连续未改善时提前停止"],
            ["训练/验证/测试比例", "7:2:1", "按时间顺序划分"],
        ],
    )
    set_table_data(
        doc.tables[5],
        [["模型", "RMSE/kW", "MAE/kW", "MAPE/%", "R²"]]
        + [
            [
                row["模型"],
                f"{row['均方根误差_kW']:.2f}",
                f"{row['平均绝对误差_kW']:.2f}",
                f"{row['平均绝对百分比误差_百分比']:.2f}",
                f"{row['决定系数R2']:.4f}",
            ]
            for _, row in pred.iterrows()
        ],
    )
    set_table_data(
        doc.tables[6],
        [
            ["设备类型", "候选容量或风量", "台数范围"],
            ["冷水机组", "100、120、140、150、160、180、200、220、240、250、280、300、320、350、380 kW", "1-4"],
            ["风机", "16、18、20、22、24、26、28、30、32、35、40、45、50 kW", "1-6"],
            ["水泵", "14、16、18、20、22、24、26、28、30、32、36、40 kW", "1-6"],
            ["空气处理机组", "25000、30000、35000、40000、43000、45000、48000、50000、55000、60000、70000 m³/h", "1-4"],
        ],
    )
    set_table_data(
        doc.tables[7],
        [
            ["参数", "取值", "说明"],
            ["电价", "0.85 元/kWh", "运行电费估算"],
            ["使用年限", "15 年", "生命周期计算周期"],
            ["折现率", "5%", "折现计算"],
            ["制冷季天数", "120 天", "年运行能耗估算"],
            ["冷水机组单位造价", "1100 元/kW", "设备初投资估算"],
            ["风机单位造价", "650 元/kW", "设备初投资估算"],
            ["水泵单位造价", "520 元/kW", "设备初投资估算"],
            ["AHU单位风量造价", "0.9 元/(m³/h)", "设备初投资估算"],
            ["维护费率", "3.5%", "按初投资比例估算"],
        ],
    )
    set_table_data(
        doc.tables[8],
        [["排名", "全生命周期成本/元", "综合容量冗余率", "TOPSIS得分"]]
        + [
            [idx + 1, money(row["生命周期成本"]), percent(row["综合冗余率"]), f"{row['TOPSIS得分']:.4f}"]
            for idx, (_, row) in enumerate(topsis.head(5).iterrows())
        ],
    )
    set_table_data(
        doc.tables[9],
        [["方案", "冷源总容量/kW", "全生命周期成本/元", "综合容量冗余率"]]
        + [
            ["基准方案", int(base["总制冷容量_kW"]), money(base["生命周期成本"]), percent(base["综合冗余率"])],
            ["优化方案", int(opt["总制冷容量_kW"]), money(opt["生命周期成本"]), percent(opt["综合冗余率"])],
        ],
    )
    set_table_data(
        doc.tables[10],
        [
            ["方案", "冷水机组", "风机", "水泵", "空气处理机组"],
            ["基准方案", "280 kW × 2 台", "18 kW × 2 台", "36 kW × 1 台", "40000 m³/h × 4 台"],
            ["优化方案", "120 kW × 4 台", "32 kW × 1 台", "30 kW × 1 台", "45000 m³/h × 3 台"],
        ],
    )
    set_table_data(
        doc.tables[11],
        [
            ["指标", "基准方案", "优化方案", "变化幅度"],
            ["冷源总容量/kW", int(base["总制冷容量_kW"]), int(opt["总制冷容量_kW"]), percent(opt["总制冷容量_kW"] / base["总制冷容量_kW"] - 1)],
            ["初投资/元", money(base["初投资"]), money(opt["初投资"]), f"-{percent(invest_drop)}"],
            ["全生命周期成本/元", money(base["生命周期成本"]), money(opt["生命周期成本"]), f"-{percent(cost_drop)}"],
            ["综合容量冗余率", percent(base["综合冗余率"]), percent(opt["综合冗余率"]), f"-{percent(redundancy_drop)}"],
            ["年运行能耗/kWh", f"{base['年能耗_kWh']:.1f}", f"{opt['年能耗_kWh']:.1f}", f"-{percent(energy_drop)}"],
            ["平均冷机负荷率", f"{base['平均冷机负荷率']:.3f}", f"{opt['平均冷机负荷率']:.3f}", percent(opt["平均冷机负荷率"] / base["平均冷机负荷率"] - 1)],
            ["平均COP", f"{base['平均COP']:.3f}", f"{opt['平均COP']:.3f}", percent(opt["平均COP"] / base["平均COP"] - 1)],
        ],
    )

    image_map = {
        169: "step1_cooling_load_timeseries.png",
        176: "step1_missing_value_summary.png",
        196: "step2_load_component_ratio.png",
        205: "step2_daily_load_clusters.png",
        288: "step3_lstm_prediction.png",
        290: "step3_bp_prediction.png",
        292: "step3_model_rmse_comparison.png",
        404: "step4_pareto_front.png",
        420: "step4_scheme_comparison.png",
    }
    for index, filename in image_map.items():
        replace_picture(doc.paragraphs[index], FIGURE_DIR / filename)

    # Turn the prior Pearson placeholder into a real figure and add a proper caption.
    add_picture_to_paragraph(doc.paragraphs[183], FIGURE_DIR / "step2_pearson_feature_ranking.png")
    inserted_caption = insert_paragraph_after(doc.paragraphs[183], "图3-3 Pearson相关性排序图")

    replace_placeholder_with_diagram(
        doc,
        "地铁车站通风空调系统组成示意图",
        DIAGRAM_DIR / "fig2_1_system.png",
        "图2-1 地铁车站通风空调系统组成示意图",
    )
    replace_placeholder_with_diagram(
        doc,
        "预测样本构造示意图",
        DIAGRAM_DIR / "fig4_1_sample.png",
        "图4-1 预测样本构造示意图",
    )
    replace_placeholder_with_diagram(
        doc,
        "NSGA-II容量优化流程图",
        DIAGRAM_DIR / "fig5_1_nsga.png",
        "图5-1 NSGA-II容量优化流程图",
        width_inches=5.7,
    )
    for p in list(doc.paragraphs):
        stripped = p.text.strip()
        if "\u3010" in stripped or "\u3011" in stripped or stripped.startswith("图 2-1"):
            p._element.getparent().remove(p._element)

    # Replace the empty reference section with GB/T 7714-style entries.
    references = [
        "[1] Hochreiter S, Schmidhuber J. Long short-term memory[J]. Neural Computation, 1997, 9(8):1735-1780.",
        "[2] Deb K, Pratap A, Agarwal S, Meyarivan T. A fast and elitist multiobjective genetic algorithm: NSGA-II[J]. IEEE Transactions on Evolutionary Computation, 2002, 6(2):182-197.",
        "[3] Hwang C L, Yoon K. Multiple Attribute Decision Making: Methods and Applications[M]. Berlin: Springer-Verlag, 1981.",
        "[4] 中华人民共和国住房和城乡建设部. 地铁设计规范: GB 50157-2013[S]. 北京: 中国建筑工业出版社, 2014.",
        "[5] 中华人民共和国住房和城乡建设部. 民用建筑供暖通风与空气调节设计规范: GB 50736-2012[S]. 北京: 中国建筑工业出版社, 2012.",
        "[6] Pearson K. Notes on regression and inheritance in the case of two parents[J]. Proceedings of the Royal Society of London, 1895, 58:240-242.",
        "[7] MacQueen J. Some methods for classification and analysis of multivariate observations[C]//Proceedings of the Fifth Berkeley Symposium on Mathematical Statistics and Probability. Berkeley: University of California Press, 1967:281-297.",
    ]
    # Find the reference heading by text after insertions.
    ref_idx = next(i for i, p in enumerate(doc.paragraphs) if p.text.strip() == "参考文献")
    insert_after = doc.paragraphs[ref_idx]
    for ref in references:
        insert_after = insert_paragraph_after(insert_after, ref)
    # Remove old blank paragraphs immediately following the inserted references when possible.
    for p in list(doc.paragraphs):
        if p.text.strip() != "":
            continue
        prev_el = p._p.getprevious()
        next_el = p._p.getnext()
        if prev_el is None or prev_el.tag != qn("w:p"):
            continue
        prev_text = Paragraph(prev_el, p._parent).text or ""
        next_text = ""
        if next_el is not None and next_el.tag == qn("w:p"):
            next_text = Paragraph(next_el, p._parent).text or ""
        if prev_text.strip().startswith("[") and (next_text.strip() == "" or next_text.strip().startswith("附录")):
            p._element.getparent().remove(p._element)

    apply_document_styles(doc)
    style_caption(inserted_caption)

    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
