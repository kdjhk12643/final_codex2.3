# -*- coding: utf-8 -*-
from __future__ import annotations

import csv
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SRC = next(p for p in ROOT.glob("*.docx") if "03" in p.name and "修改稿" not in p.name)
OUT = ROOT / "毕业设计论文初稿03_修改稿.docx"


def clear_para(p: Paragraph) -> None:
    for r in list(p.runs):
        r._element.getparent().remove(r._element)


def set_para(
    p: Paragraph,
    text: str,
    style=None,
    bold: bool = False,
    align: WD_ALIGN_PARAGRAPH | None = None,
) -> Paragraph:
    clear_para(p)
    if style:
        p.style = style
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(11)
    run.bold = bold
    return p


def find_para(doc: Document, predicate):
    for p in doc.paragraphs:
        if predicate(p.text.strip()):
            return p
    return None


def find_para_index(doc: Document, predicate) -> int:
    for i, p in enumerate(doc.paragraphs):
        if predicate(p.text.strip()):
            return i
    return -1


def insert_paragraph_after(paragraph: Paragraph, text: str, style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    set_para(new_para, text)
    return new_para


def insert_paragraph_before(paragraph: Paragraph) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    return Paragraph(new_p, paragraph._parent)


def delete_paragraph(paragraph: Paragraph) -> None:
    el = paragraph._element
    el.getparent().remove(el)
    paragraph._p = paragraph._element = None


def replace_between(doc: Document, start_pred, end_pred, new_paras: list[str]) -> None:
    start_i = find_para_index(doc, start_pred)
    end_i = find_para_index(doc, end_pred)
    if start_i < 0 or end_i < 0 or end_i <= start_i:
        raise RuntimeError(f"bad paragraph range: {start_i}, {end_i}")
    start_p = doc.paragraphs[start_i]
    for p in list(doc.paragraphs[start_i + 1 : end_i]):
        delete_paragraph(p)
    cur = start_p
    for text in new_paras:
        cur = insert_paragraph_after(cur, text)


def replace_all(doc: Document, old: str, new: str) -> None:
    for p in doc.paragraphs:
        if old in p.text:
            set_para(p, p.text.replace(old, new), style=p.style)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if old in p.text:
                        set_para(p, p.text.replace(old, new), style=p.style)


def set_cell(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(str(text)) <= 18 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.font.name = "宋体"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def rewrite_table(table, rows: list[list[str]]) -> None:
    while len(table.rows) < len(rows):
        table.add_row()
    while len(table.rows) > len(rows):
        tr = table.rows[-1]._tr
        tr.getparent().remove(tr)
    max_cols = max(len(r) for r in rows)
    if len(table.columns) != max_cols:
        # Existing thesis tables already have the intended column count except
        # the appended ablation table. Avoid unsafe column surgery on user DOCX.
        max_cols = min(max_cols, len(table.columns))
    for i, row in enumerate(rows):
        for j, val in enumerate(row[:max_cols]):
            set_cell(table.cell(i, j), val, bold=(i == 0))


def has_drawing(paragraph: Paragraph) -> bool:
    return bool(paragraph._p.xpath(".//w:drawing") or paragraph._p.xpath(".//w:pict"))


def previous_paragraph(doc: Document, paragraph: Paragraph):
    for i, p in enumerate(doc.paragraphs):
        if p._p is paragraph._p:
            return doc.paragraphs[i - 1] if i > 0 else None
    return None


def insert_picture_before_caption(doc: Document, caption: str, image_path: Path) -> None:
    target = find_para(doc, lambda t: t == caption)
    if not target or not image_path.exists():
        return
    prev = previous_paragraph(doc, target)
    if prev is not None and has_drawing(prev):
        return
    pic_p = insert_paragraph_before(target)
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pic_p.add_run()
    run.add_picture(str(image_path), width=Cm(13.5))


def read_csv(name: str) -> list[dict[str, str]]:
    with open(ROOT / "output" / "tables" / name, encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def pct(x: str | float, digits: int = 2) -> str:
    return f"{float(x) * 100:.{digits}f}%"


def num(x: str | float, digits: int = 2) -> str:
    return f"{float(x):.{digits}f}"


def update_cover_and_abstracts(doc: Document) -> None:
    for p in doc.paragraphs:
        if p.text.strip().startswith("题目："):
            set_para(p, "题目：    基于负荷特性的地铁车站通风空调系统")
            break
    idx = find_para_index(doc, lambda t: t == "题目：    基于负荷特性的地铁车站通风空调系统")
    if idx >= 0 and idx + 1 < len(doc.paragraphs):
        set_para(doc.paragraphs[idx + 1], "容量配置优化研究")
    for p in doc.paragraphs:
        if "指 导 教 师 签 名" in p.text and "2025年" in p.text:
            set_para(p, p.text.replace("2025年", "2026年"))

    abstract_cn = (
        "本文以地铁车站站台区域通风空调系统为研究对象，针对传统容量配置中设备容量偏大、系统长期低负荷运行和全生命周期成本较高等问题，构建了基于负荷特性的容量配置优化方法。"
        "研究基于福州地铁东街口站 2025 年全年 15 min 粒度运行数据，首先完成时间对齐、缺失值处理、标准化和滞后特征构造；随后采用 Pearson 相关系数法筛选关键影响因素，"
        "并将总冷负荷分解为人员负荷、新风负荷、围护结构负荷和设备散热负荷；在此基础上利用 K-Means 聚类识别典型日负荷模式，建立两层 LSTM 神经网络负荷预测模型，并与 BP 神经网络进行对比。"
        "最后，以全生命周期成本最小和综合容量冗余率最小为目标，采用 NSGA-II 求解 Pareto 候选方案，并通过 TOPSIS 方法确定推荐容量配置方案。结果表明，LSTM 模型在测试集上的 RMSE、MAE、MAPE 和 R² 分别为 9.47 kW、7.14 kW、3.81% 和 0.9715，预测精度优于 BP 模型；"
        "推荐优化方案使冷源总容量由 560 kW 降至 480 kW，综合容量冗余率由 34.51% 降至 13.06%，全生命周期成本降低 11.10%，年运行能耗降低 7.56%。研究结果可为地铁车站通风空调系统容量配置、节能设计和既有车站改造方案比选提供参考。"
    )
    p = find_para(
        doc,
        lambda t: t.startswith("本文以地铁车站站台区域通风空调系统为研究对象") and "Pearson" in t,
    )
    if p:
        set_para(p, abstract_cn)

    english_i = find_para_index(doc, lambda t: t.startswith("Design and Implementation of a Self-Balancing"))
    if english_i >= 0:
        set_para(
            doc.paragraphs[english_i],
            "Capacity Configuration Optimization of Metro Station Ventilation and Air-Conditioning System Based on Load Characteristics",
            align=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True,
        )
        set_para(
            doc.paragraphs[english_i + 1],
            "Student: WEN Mingdong    Supervisor: WU Dong",
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        english_abs = (
            "Abstract: This thesis studies the ventilation and air-conditioning system of a metro station platform area and proposes a load-characteristic-based capacity configuration optimization method. "
            "Using a 15-minute operating dataset of Fuzhou Metro Dongjiekou Station in 2025, the study completes data preprocessing, Pearson correlation analysis, load component decomposition and K-Means daily load clustering. "
            "A two-layer LSTM model is then established for cooling-load prediction and compared with a BP neural network. Based on the predicted load demand, a multi-objective capacity optimization model is built with lifecycle cost and composite capacity redundancy as objectives. NSGA-II is used to generate Pareto candidate schemes, and TOPSIS is used to select the recommended configuration. "
            "The results show that the LSTM model achieves RMSE = 9.47 kW, MAE = 7.14 kW, MAPE = 3.81% and R² = 0.9715 on the test set. The recommended scheme reduces total cooling capacity from 560 kW to 480 kW, composite redundancy from 34.51% to 13.06%, lifecycle cost by 11.10%, and annual operating energy consumption by 7.56%. The proposed method can support energy-saving design and retrofit decision-making for metro station HVAC systems."
        )
        set_para(doc.paragraphs[english_i + 2], english_abs)
        set_para(
            doc.paragraphs[english_i + 3],
            "Key words: Metro station; ventilation and air-conditioning system; load prediction; LSTM; NSGA-II; TOPSIS; capacity optimization",
        )


def update_research_status(doc: Document) -> None:
    research_status = [
        "地铁车站通风空调系统负荷受室外气象、客流强度、列车运行、站内设备散热、围护结构传热以及新风需求等多因素共同影响。与普通公共建筑相比，地铁车站处于地下空间，受太阳辐射直接影响较弱，但客流时变性、列车活塞风、站内设备长期运行等因素使其负荷呈现明显的时段性和运行场景差异。已有研究普遍认为，环控系统是城市轨道交通运营能耗的重要组成部分，节能优化既要关注运行控制，也要关注设计容量是否与实际需求匹配。",
        "在负荷预测方面，传统方法包括线性回归、时间序列模型、灰色预测和 BP 神经网络等。这类方法结构相对简单、实现方便，但在处理非线性、多变量和长短期时序依赖时存在一定局限。近年来，LSTM、GRU、Transformer 等深度学习模型逐渐用于建筑和交通场站空调负荷预测，其中 LSTM 通过门控结构保存历史状态，适合描述负荷惯性、客流高峰持续和环境参数滞后影响等问题。",
        "在容量配置和节能优化方面，传统工程设计通常依据设计峰值负荷、远期客流和经验安全裕度确定设备容量，方法可靠但容易造成装机容量偏大。多目标优化方法能够同时考虑经济性、冗余率和工程约束，NSGA-II 可输出 Pareto 非劣解集，TOPSIS 可进一步根据成本和冗余等指标进行综合排序，适合用于设备容量组合方案比选。",
        "综合来看，现有研究已在地铁车站负荷特性、空调负荷预测和系统节能优化方面取得进展，但仍存在预测结果与容量配置衔接不足、容量冗余率和全生命周期成本考虑不充分、工程边界说明不够明确等问题。因此，本文将负荷特性分析、LSTM 预测和 NSGA-II/TOPSIS 容量优化串联起来，形成面向站台区域通风空调系统的完整方法链。",
    ]
    replace_between(doc, lambda t: t == "国内外研究现状", lambda t: t == "研究内容与技术路线", research_status)


def update_captions_and_static_text(doc: Document) -> None:
    caption_map = {
        "图1-1 技术路线图": "图 1-1 技术路线图",
        "【图2-1 地铁车站通风空调系统组成示意图：冷水机组、冷冻水泵、冷却水泵、空气处理机组、站台送风机、站台排风机、站台区域】": "图 2-1 地铁车站通风空调系统组成示意图",
        "【图3-2 缺失值统计图】": "图 3-2 缺失值统计图",
        "【图3-3 站台总冷负荷时间序列图】": "图 3-3 站台总冷负荷时间序列图",
        "【图3-4 Pearson相关性排序图】": "图 3-4 Pearson 相关性排序图",
        "【图3-5 负荷分项占比图】": "图 3-5 负荷分项占比图",
        "【图3-6 典型日负荷曲线聚类图】": "图 3-6 典型日负荷曲线聚类图",
        "【图4-1 预测样本构造示意图：连续16个时间步多维输入 → 下一时刻总冷负荷】": "图 4-1 预测样本构造示意图",
        "【图4-2 LSTM网络结构图：Sequence Input → LSTM(64) → LSTM(32) → Fully Connected → Regression Output】": "图 4-2 LSTM 网络结构图",
        "【图4-3 LSTM预测结果图】": "图 4-3 LSTM 预测结果图",
        "【图4-4 BP预测结果图】": "图 4-4 BP 预测结果图",
        "【图4-5 模型RMSE对比图】": "图 4-5 模型 RMSE 对比图",
        "【图5-1 NSGA-II容量优化流程图：初始化种群 → 解码设备方案 → 约束校核 → 目标函数计算 → 非支配排序 → 拥挤距离计算 → 选择交叉变异 → Pareto解集】": "图 5-1 NSGA-II 容量优化流程图",
        "【图5-2 Pareto前沿图】": "图 5-2 Pareto 前沿图",
        "【图6-1 基准方案与优化方案对比图】": "图 6-1 基准方案与优化方案对比图",
    }
    for old, new in caption_map.items():
        replace_all(doc, old, new)
    for old, new in {
        "表2-1": "表 2-1",
        "表2-2": "表 2-2",
        "表4-1": "表 4-1",
        "表5-1": "表 5-1",
        "表5-2": "表 5-2",
        "表5-3": "表 5-3",
        "表6-2": "表 6-2",
    }.items():
        replace_all(doc, old, new)
    replace_all(doc, "2025 年 7 月", "2025 年全年")
    replace_all(doc, "2025年7月", "2025年全年")
    replace_all(doc, "数据时间范围为 2025 年全年", "数据时间范围为 2025 年 1 月 1 日至 2025 年 12 月 31 日")
    replace_all(doc, "K = 2 至 K = 6", "K = 2 至 K = 4")
    replace_all(doc, "站台在岗人数", "站台人数")

    for p in doc.paragraphs:
        t = p.text.strip()
        if t == "6.4 调节灵活性与运行适应性分析":
            set_para(p, "6.6 调节灵活性与运行适应性分析", style=p.style)
        elif t == "6.5 既有车站改造适配性分析":
            set_para(p, "6.7 既有车站改造适配性分析", style=p.style)
        elif t == "6.6本章小结":
            set_para(p, "6.8 本章小结", style=p.style)

    figure_dir = ROOT / "output" / "figures"
    for caption, image_name in {
        "图 3-2 缺失值统计图": "step1_missing_value_summary.png",
        "图 3-3 站台总冷负荷时间序列图": "step1_cooling_load_timeseries.png",
        "图 3-4 Pearson 相关性排序图": "step2_pearson_feature_ranking.png",
        "图 3-5 负荷分项占比图": "step2_load_component_ratio.png",
        "图 3-6 典型日负荷曲线聚类图": "step2_daily_load_clusters.png",
        "图 4-3 LSTM 预测结果图": "step3_lstm_prediction.png",
        "图 4-4 BP 预测结果图": "step3_bp_prediction.png",
        "图 4-5 模型 RMSE 对比图": "step3_model_rmse_comparison.png",
        "图 5-2 Pareto 前沿图": "step4_pareto_front.png",
        "图 6-1 基准方案与优化方案对比图": "step4_scheme_comparison.png",
    }.items():
        insert_picture_before_caption(doc, caption, figure_dir / image_name)


def update_key_paragraphs(doc: Document) -> None:
    replacements_by_start = {
        "第一，地铁车站站台区域空调负荷受历史负荷": "第一，地铁车站站台区域空调负荷受历史负荷、客流和环境参数共同影响，负荷序列表现出明显的时序连续性。Pearson 相关分析结果表明，冷负荷滞后1步与总冷负荷的相关系数为 0.9861，冷负荷滞后4步为 0.9387，说明历史负荷对当前负荷具有较强指示作用。站台人数、进站客流和出站客流与总冷负荷的相关系数分别为 0.8339、0.8266 和 0.8155，表明客流变化是站台区域负荷波动的重要驱动因素。二氧化碳浓度、站台温度、太阳辐射和室外温度等环境变量也与负荷具有较高相关性，可作为负荷预测模型的重要输入。",
        "第二，负荷分项结果表明": "第二，负荷分项结果表明，站台区域总冷负荷并非仅由人员活动决定，而是由围护结构传热、设备散热、新风处理和人员散热共同构成。其中，围护结构负荷平均占比约为 50.00%，设备散热负荷平均占比约为 31.96%，人员负荷平均占比约为 6.46%，新风负荷平均占比约为 5.21%。这说明地铁车站站台区域负荷具有较强的地下空间和机电设备运行特征，容量配置优化需要综合考虑基础负荷和客流波动负荷。",
        "第三，基于 K-Means 的典型负荷曲线聚类": "第三，基于 K-Means 的典型负荷曲线聚类能够有效识别站台区域负荷模式。本文在 K = 2、3、4 范围内进行聚类分析，K=4 时平均轮廓系数最高，为 0.6797，因此选取 K=4 作为典型负荷模式划分结果。聚类结果可用于解释不同客流强度和日负荷形态下的运行场景差异，并通过聚类标签消融实验进一步检验其对预测模型的贡献。",
        "第四，LSTM 神经网络能够较好预测": "第四，LSTM 神经网络能够较好预测地铁车站站台区域空调负荷。本文以客流、气象、站内环境、时间特征、历史负荷项和聚类标签作为输入，构建两层 LSTM 负荷预测模型，并与 BP 神经网络进行对比。测试结果表明，LSTM 模型 RMSE 为 9.47 kW，MAE 为 7.14 kW，MAPE 为 3.81%，R² 为 0.9715，整体预测精度优于 BP 模型。说明 LSTM 对地铁车站负荷时序特征具有较强表达能力，适合作为容量配置优化的负荷预测工具。",
        "第五，基于 NSGA-II 和 TOPSIS 的容量配置优化方法": "第五，基于 NSGA-II 和 TOPSIS 的容量配置优化方法能够有效降低系统容量冗余和全生命周期成本。本文以全生命周期成本最小和综合容量冗余率最小为目标，建立冷水机组、风机、水泵和空气处理机组的离散容量组合优化模型，并通过 TOPSIS 方法从 Pareto 解集中选取综合最优方案。结果表明，优化方案冷源总容量由基准方案的 560 kW 降低至 480 kW，综合容量冗余率由 34.51% 降低至 13.06%，全生命周期成本由 217.86 万元降低至 193.67 万元。",
        "第六，优化方案具有较好的经济性": "第六，优化方案具有较好的经济性、节能性和工程适用性。与基准方案相比，优化方案初投资由 80.21 万元降低至 68.49 万元，全生命周期成本降低 11.10%，年运行能耗由 122987 kWh 降低至 113691 kWh，节能率为 7.56%。同时，优化方案冷机平均部分负荷率由 0.517 提高至 0.750，平均 COP 由 4.699 提高至 5.037，说明容量配置优化可以改善设备部分负荷运行状态。",
    }
    for p in doc.paragraphs:
        t = p.text.strip()
        for start, new in replacements_by_start.items():
            if t.startswith(start):
                set_para(p, new, style=p.style)

    # Update detailed LSTM model settings paragraph and forecast/optimization result paragraphs.
    paragraph_replacements = {
        "2. 第一层 LSTM：隐藏单元数为 64，输出完整序列；": "2. 第一层 LSTM：隐藏单元数为 96，输出完整序列；",
        "3. 第二层 LSTM：隐藏单元数为 32，输出最后一个时间步状态；": "3. 第二层 LSTM：隐藏单元数为 48，输出最后一个时间步状态；",
        "训练过程中采用 Adam 优化算法，最大训练轮数为 60，批量大小为 64，初始学习率为 0.001。为避免梯度爆炸，设置梯度阈值为 1；为提高训练稳定性，学习率采用分段下降策略，验证集若连续多轮未改善则提前停止训练。": "训练过程中采用 Adam 优化算法，最大训练轮数为 120，批量大小为 32，初始学习率为 0.0008。为避免梯度爆炸，设置梯度阈值为 1；为提高训练稳定性，学习率采用分段下降策略，验证集若连续多轮未改善则提前停止训练。",
        "由表 4-1可知，LSTM 模型在各项评价指标上均优于 BP 神经网络。LSTM 的 RMSE 为 13.28 kW，低于 BP 模型的 17.87 kW；LSTM 的 MAE 为 10.69 kW，低于 BP 模型的 14.33 kW；LSTM 的 MAPE 为 5.69%，低于 BP 模型的 7.64%；LSTM 的 R² 为 0.9557，高于 BP 模型的 0.9194。": "由表 4-1 可知，LSTM 模型在各项评价指标上均优于 BP 神经网络。LSTM 的 RMSE 为 9.47 kW，低于 BP 模型的 22.21 kW；LSTM 的 MAE 为 7.14 kW，低于 BP 模型的 15.86 kW；LSTM 的 MAPE 为 3.81%，低于 BP 模型的 8.59%；LSTM 的 R² 为 0.9715，高于 BP 模型的 0.8434。",
        "本章基于第三章筛选得到的关键影响因素，建立了地铁车站站台区域空调负荷预测模型。首先构建了包含客流、气象、站内环境、时间特征和历史负荷项的预测输入矩阵，并按 7:2:1 的比例划分训练集、验证集和测试集。随后建立两层 LSTM 神经网络模型，并采用 BP 神经网络作为对比模型。测试结果表明，LSTM 模型的 RMSE 为 13.28 kW，MAPE 为 5.69%，R² 为 0.9557，整体预测精度优于 BP 模型。说明 LSTM 能够较好捕捉地铁车站空调负荷的时序变化规律，可作为后续容量配置优化的负荷输入基础。": "本章基于第三章筛选得到的关键影响因素，建立了地铁车站站台区域空调负荷预测模型。首先构建了包含客流、气象、站内环境、时间特征和历史负荷项的预测输入矩阵，并按 7:2:1 的比例划分训练集、验证集和测试集。随后建立两层 LSTM 神经网络模型，并采用 BP 神经网络作为对比模型。测试结果表明，LSTM 模型的 RMSE 为 9.47 kW，MAPE 为 3.81%，R² 为 0.9715，整体预测精度优于 BP 模型。说明 LSTM 能够较好捕捉地铁车站空调负荷的时序变化规律，可作为后续容量配置优化的负荷输入基础。",
        "由表 5-1可知，优化方案的冷源总容量由基准方案的 440 kW 降低至 360 kW，综合容量冗余率由 33.40% 降低至 9.72%。同时，全生命周期成本由 471.19 万元降低至 371.15 万元。说明本文建立的容量配置优化模型能够在满足负荷需求和工程约束的前提下，有效降低设备容量冗余和生命周期经济成本。": "由表 5-1 可知，优化方案的冷源总容量由基准方案的 560 kW 降低至 480 kW，综合容量冗余率由 34.51% 降低至 13.06%。同时，全生命周期成本由 217.86 万元降低至 193.67 万元。说明本文建立的容量配置优化模型能够在满足负荷需求和工程约束的前提下，有效降低设备容量冗余和生命周期经济成本。",
        "从优化逻辑看，容量由 440 kW 降至 360 kW 并非简单削减设备规模，而是在预测负荷、设计负荷、安全系数和子系统匹配约束共同作用下，对传统基准方案中不必要冗余的再分配。综合容量冗余率降至 9.72% 表明优化方案仍保留一定安全裕度，但已显著减少长期低负荷运行风险。全生命周期成本降低则同时来自设备初投资减少和部分负荷运行效率改善。": "从优化逻辑看，容量由 560 kW 降至 480 kW 并非简单削减设备规模，而是在预测负荷、设计负荷、安全系数和子系统匹配约束共同作用下，对传统基准方案中不必要冗余的再分配。综合容量冗余率降至 13.06% 表明优化方案仍保留一定安全裕度，但已显著减少长期低负荷运行风险。全生命周期成本降低则同时来自设备初投资减少和部分负荷运行效率改善。",
        "由表6-1可知，优化方案在冷源总容量、初投资、全生命周期成本、容量冗余率和年运行能耗等方面均优于基准方案。优化方案冷源总容量由 440 kW 降低至 360 kW，说明在基于负荷预测和系统匹配约束的前提下，可以减少传统设计中不必要的容量冗余。与此同时，优化方案全生命周期成本降低 21.23%，年运行能耗降低 21.93%，表明容量配置优化不仅能够降低设备投资，也能够改善系统部分负荷运行效率。": "由表 6-1 可知，优化方案在冷源总容量、初投资、全生命周期成本、容量冗余率和年运行能耗等方面均优于基准方案。优化方案冷源总容量由 560 kW 降低至 480 kW，说明在基于负荷预测和系统匹配约束的前提下，可以减少传统设计中不必要的容量冗余。与此同时，优化方案全生命周期成本降低 11.10%，年运行能耗降低 7.56%，表明容量配置优化不仅能够降低设备投资，也能够改善系统部分负荷运行效率。",
        "容量冗余率是评价设备容量配置合理性的重要指标。基准方案的综合容量冗余率为 33.40%，说明传统配置方案中设备容量明显高于实际负荷需求。较高的容量冗余虽然能够提高峰值工况下的安全裕度，但也会导致设备长期运行在较低负荷率区间，降低系统运行效率。": "容量冗余率是评价设备容量配置合理性的重要指标。基准方案的综合容量冗余率为 34.51%，说明传统配置方案中设备容量明显高于实际负荷需求。较高的容量冗余虽然能够提高峰值工况下的安全裕度，但也会导致设备长期运行在较低负荷率区间，降低系统运行效率。",
        "优化方案的综合容量冗余率为 9.72%，相比基准方案下降 70.90%。该结果表明，本文优化模型能够在满足设计负荷和安全系数约束的基础上，有效压缩过度裕量，使设备配置更加接近实际负荷需求。": "优化方案的综合容量冗余率为 13.06%，相比基准方案下降 62.15%。该结果表明，本文优化模型能够在满足设计负荷和安全系数约束的基础上，有效压缩过度裕量，使设备配置更加接近实际负荷需求。",
        "从全生命周期成本看，基准方案为 4711887 元，优化方案为 3711458 元，降低 1000429 元，降幅为 21.23%。全生命周期成本降幅高于初投资降幅，说明优化方案不仅降低了设备购置成本，也在长期运行过程中减少了能耗费用和维护成本。": "从全生命周期成本看，基准方案为 2178599 元，优化方案为 1936720 元，降低 241879 元，降幅为 11.10%。全生命周期成本降幅高于初投资降幅，说明优化方案不仅降低了设备购置成本，也在长期运行过程中减少了能耗费用和维护成本。",
        "计算结果表明，基准方案年运行能耗为 444407.7 kWh，优化方案年运行能耗为 346955.5 kWh，年节电量为 97452.2 kWh，节能率为 21.93%。": "计算结果表明，基准方案年运行能耗为 122987 kWh，优化方案年运行能耗为 113690.6 kWh，年节电量为 9296.4 kWh，节能率为 7.56%。",
        "从冷源侧看，优化方案冷源总容量由 440 kW 降低至 360 kW，仍能够满足设计负荷与安全系数要求，说明基准方案存在一定冷源容量过剩。": "从冷源侧看，优化方案冷源总容量由 560 kW 降低至 480 kW，仍能够满足设计负荷与安全系数要求，说明基准方案存在一定冷源容量过剩。",
        "本章对基准方案和优化方案进行了对比分析，并从容量冗余率、全生命周期成本、年运行能耗、设备匹配性、调节灵活性和既有车站改造适配性等方面评价了优化效果。结果表明，优化方案冷源总容量由 440 kW 降低至 360 kW，综合容量冗余率由 33.40% 降低至 9.72%，全生命周期成本降低 21.23%，年运行能耗降低 21.93%。说明基于负荷特性的容量配置优化方法能够在满足工程约束的前提下，有效降低设备容量冗余、减少生命周期成本并提升系统节能潜力。与此同时，本文也指出当前能耗评价仍存在简化，后续应结合全年运行数据和设备效率曲线进一步验证。": "本章对基准方案和优化方案进行了对比分析，并从容量冗余率、全生命周期成本、年运行能耗、设备匹配性、调节灵活性和既有车站改造适配性等方面评价了优化效果。结果表明，优化方案冷源总容量由 560 kW 降低至 480 kW，综合容量冗余率由 34.51% 降低至 13.06%，全生命周期成本降低 11.10%，年运行能耗降低 7.56%。说明基于负荷特性的容量配置优化方法能够在满足工程约束的前提下，有效降低设备容量冗余、减少生命周期成本并提升系统节能潜力。与此同时，本文也指出当前能耗评价仍存在简化，后续应结合全年运行数据和设备效率曲线进一步验证。",
    }
    for p in doc.paragraphs:
        t = p.text.strip()
        for old, new in paragraph_replacements.items():
            if t == old:
                set_para(p, new, style=p.style)
                break

    # Add a concise note about the cluster-label ablation experiment after the
    # main prediction comparison paragraph.
    p = find_para(doc, lambda t: t.startswith("由表 4-1 可知"))
    if p and not any("消融实验" in x.text for x in doc.paragraphs):
        insert_paragraph_after(
            p,
            "聚类标签消融实验结果表明，含聚类标签的 LSTM 测试集 RMSE 为 8.87 kW，略优于不含聚类标签的 9.06 kW，R² 由 0.9739 提升至 0.9750。说明典型日负荷模式对预测具有一定辅助作用，但提升幅度有限，因此本文将聚类结果主要用于运行场景解释和情景分析。",
        )

    summary = find_para(doc, lambda t: t.startswith("综上，本文提出的基于负荷特性的"))
    if summary and not any("施工图阶段最终设备选型" in p.text for p in doc.paragraphs):
        insert_paragraph_after(
            summary,
            "需要说明的是，本文优化方案主要用于本科毕设层面的容量配置方法验证和方案相对比较。当前模型已完成负荷预测误差裕量、典型工况负荷率、子系统容量比例等校核，但在极端工况显式裕量、设备厂家性能曲线、风水系统阻力计算、舒适性闭环控制和真实 BMS/AFC 数据校准方面仍需进一步复核，不能直接等同于施工图阶段最终设备选型。",
        )


def update_tables(doc: Document) -> None:
    pearson = read_csv("step2_pearson_features.csv")[:10]
    ratio = read_csv("step2_load_component_ratio.csv")
    silh = read_csv("step2_cluster_silhouette.csv")
    metrics = read_csv("step3_prediction_metrics.csv")
    topsis = read_csv("step4_topsis_ranking.csv")[:5]
    eval_rows = read_csv("step4_scheme_evaluation.csv")

    rewrite_table(doc.tables[1], [["特征", "Pearson R"]] + [[r["特征"], num(r["Pearson相关系数"], 4)] for r in pearson])
    rewrite_table(doc.tables[2], [["分项", "平均负荷/kW", "占比"]] + [[r["分项负荷"], num(r["平均负荷_kW"], 2), pct(r["占比"], 2)] for r in ratio])
    rewrite_table(doc.tables[3], [["K", "平均轮廓系数"]] + [[r["K值"], num(r["平均轮廓系数"], 4)] for r in silh])
    rewrite_table(
        doc.tables[4],
        [
            ["参数", "取值", "说明"],
            ["输入序列长度", "16", "对应 4 小时历史信息"],
            ["第一层 LSTM 单元数", "96", "提取主要时序特征"],
            ["第二层 LSTM 单元数", "48", "输出最终时序状态"],
            ["最大训练轮数", "120", "控制训练上限"],
            ["批量大小", "32", "小批量训练"],
            ["初始学习率", "0.0008", "训练初始步长"],
            ["训练/验证/测试比例", "7:2:1", "按时间顺序划分"],
        ],
    )
    rewrite_table(
        doc.tables[5],
        [["模型", "RMSE/kW", "MAE/kW", "MAPE/%", "R²"]]
        + [
            [
                r["模型"],
                num(r["均方根误差_kW"], 2),
                num(r["平均绝对误差_kW"], 2),
                num(r["平均绝对百分比误差_百分比"], 2),
                num(r["决定系数R2"], 4),
            ]
            for r in metrics
        ],
    )
    rewrite_table(
        doc.tables[6],
        [
            ["设备类型", "候选容量或风量", "台数范围"],
            ["冷水机组", "100、120、140、150、160、180、200、220、240、250、280、300、320、350、380 kW", "1-4"],
            ["风机", "16、18、20、22、24、26、28、30、32、35、40、45、50 kW", "1-6"],
            ["水泵", "14、16、18、20、22、24、26、28、30、32、36、40 kW", "1-6"],
            ["空气处理机组", "25000、30000、35000、40000、43000、45000、48000、50000、55000、60000、70000 m³/h", "1-4"],
        ],
    )
    rewrite_table(
        doc.tables[7],
        [
            ["参数", "取值", "说明"],
            ["电价", "0.85 元/kWh", "运行电费估算"],
            ["使用年限", "15 年", "生命周期计算周期"],
            ["折现率", "5%", "折现计算"],
            ["制冷季天数", "120 天", "年运行能耗估算"],
            ["维护费率", "3.5%", "按初投资比例估算"],
        ],
    )
    rewrite_table(
        doc.tables[8],
        [["排名", "全生命周期成本/元", "综合容量冗余率", "TOPSIS 得分"]]
        + [[str(i + 1), num(r["生命周期成本"], 0), pct(r["综合冗余率"], 2), num(r["TOPSIS得分"], 4)] for i, r in enumerate(topsis)],
    )
    rewrite_table(
        doc.tables[9],
        [["方案", "冷源总容量/kW", "全生命周期成本/元", "综合容量冗余率"]]
        + [[r["方案"], num(r["总制冷容量_kW"], 0), num(r["生命周期成本"], 0), pct(r["综合冗余率"], 2)] for r in eval_rows],
    )
    rewrite_table(
        doc.tables[10],
        [
            ["方案", "冷水机组", "风机", "水泵", "空气处理机组"],
            ["基准方案", "280 kW × 2 台", "18 kW × 2 台", "36 kW × 1 台", "40000 m³/h × 4 台"],
            ["优化方案", "120 kW × 4 台", "32 kW × 1 台", "28 kW × 1 台", "45000 m³/h × 3 台"],
        ],
    )
    base, opt = eval_rows[0], eval_rows[1]
    rewrite_table(
        doc.tables[11],
        [
            ["指标", "基准方案", "优化方案", "变化幅度"],
            ["冷源总容量/kW", num(base["总制冷容量_kW"], 0), num(opt["总制冷容量_kW"], 0), "-14.29%"],
            ["初投资/元", num(base["初投资"], 0), num(opt["初投资"], 0), "-14.62%"],
            ["全生命周期成本/元", num(base["生命周期成本"], 0), num(opt["生命周期成本"], 0), "-11.10%"],
            ["综合容量冗余率", pct(base["综合冗余率"], 2), pct(opt["综合冗余率"], 2), "-62.15%"],
            ["年运行能耗/kWh", num(base["年能耗_kWh"], 0), num(opt["年能耗_kWh"], 0), "-7.56%"],
            ["平均冷机负荷率", num(base["平均冷机负荷率"], 3), num(opt["平均冷机负荷率"], 3), "+45.01%"],
            ["平均 COP", num(base["平均COP"], 3), num(opt["平均COP"], 3), "+7.19%"],
        ],
    )


def update_references_and_appendix(doc: Document) -> None:
    ack = find_para(doc, lambda t: t.replace(" ", "") == "致谢")
    ref = find_para(doc, lambda t: t == "参考文献")
    if ack and ref:
        ai = find_para_index(doc, lambda t: t.replace(" ", "") == "致谢")
        ri = find_para_index(doc, lambda t: t == "参考文献")
        for p in list(doc.paragraphs[ai + 1 : ri]):
            delete_paragraph(p)
        insert_paragraph_after(
            ack,
            "在本课题研究和论文撰写过程中，感谢指导教师在研究方向、技术路线和论文结构方面给予的指导；感谢学院提供毕业设计平台和学习条件；感谢同学在数据整理、模型调试和论文排版过程中给予的帮助。本文在写作过程中参考了通风空调、负荷预测和多目标优化等领域的相关研究成果，在此一并致谢。",
        )

    references = [
        "中华人民共和国住房和城乡建设部. 地铁设计规范: GB 50157-2013[S]. 北京: 中国计划出版社, 2013.",
        "中华人民共和国住房和城乡建设部. 民用建筑供暖通风与空气调节设计规范: GB 50736-2012[S]. 北京: 中国建筑工业出版社, 2012.",
        "中华人民共和国住房和城乡建设部. 城市轨道交通通风空气调节与供暖设计标准: GB/T 51357-2019[S]. 北京: 中国计划出版社, 2019.",
        "ASHRAE. ASHRAE Handbook: HVAC Applications[M]. Atlanta: ASHRAE, 2019.",
        "尚聪. 计及客流特性的地铁环控系统节能策略研究[D]. 2024.",
        "Hochreiter S, Schmidhuber J. Long short-term memory[J]. Neural Computation, 1997, 9(8): 1735-1780.",
        "Deb K, Pratap A, Agarwal S, et al. A fast and elitist multiobjective genetic algorithm: NSGA-II[J]. IEEE Transactions on Evolutionary Computation, 2002, 6(2): 182-197.",
        "Hwang C L, Yoon K. Multiple Attribute Decision Making: Methods and Applications[M]. Berlin: Springer, 1981.",
        "MacQueen J. Some methods for classification and analysis of multivariate observations[C]//Proceedings of the Fifth Berkeley Symposium on Mathematical Statistics and Probability. Berkeley: University of California Press, 1967: 281-297.",
        "MATLAB. Deep Learning Toolbox User’s Guide[Z]. Natick: The MathWorks, Inc., 2022.",
    ]
    ref_i = find_para_index(doc, lambda t: t == "参考文献")
    if ref_i >= 0:
        for p in list(doc.paragraphs[ref_i + 1 :]):
            delete_paragraph(p)
        cur = doc.paragraphs[ref_i]
        for i, item in enumerate(references, 1):
            cur = insert_paragraph_after(cur, f"[{i}] {item}")
        cur = insert_paragraph_after(cur, "附录 A 主要代码文件说明", style="Heading 1")
        for line in [
            "main.m：项目主入口，依次调用数据预处理、影响因素分析与聚类、负荷预测和容量优化四个步骤。",
            "config.m：集中管理数据路径、模型超参数、设备候选容量、经济性假设和工程约束参数。",
            "step1_data_prepare.m：完成时间戳统一、缺失值处理、异常校核、标准化和建模数据表输出。",
            "step2_analysis_cluster.m：完成 Pearson 相关性排序、负荷分项统计和 K-Means 典型日负荷聚类。",
            "step3_load_prediction.m：完成 LSTM 与 BP 负荷预测、评价指标计算、聚类标签消融实验和子系统需求情景输出。",
            "step4_capacity_optimization.m：完成基准方案构造、NSGA-II/Pareto 候选方案搜索、TOPSIS 排序、工程约束校核和方案评价。",
        ]:
            cur = insert_paragraph_after(cur, line)
        cur = insert_paragraph_after(cur, "附录 B 结果文件与复现实验说明", style="Heading 1")
        for line in [
            "复现实验时，在 MATLAB R2022b 环境下运行 main.m，可重新生成 output/tables、output/figures 和 output/models 中的主要结果文件。",
            "关键结果表包括 step2_pearson_features.csv、step2_load_component_ratio.csv、step2_cluster_silhouette.csv、step3_prediction_metrics.csv、step3_cluster_ablation_metrics.csv、step4_topsis_ranking.csv 和 step4_scheme_evaluation.csv。",
            "关键图件包括 step1_cooling_load_timeseries.png、step2_pearson_feature_ranking.png、step2_daily_load_clusters.png、step3_lstm_prediction.png、step3_model_rmse_comparison.png、step4_pareto_front.png 和 step4_scheme_comparison.png。",
            "本文结果应以 output/tables 中当前 CSV 文件为准；若修改 config.m 中的安全系数、TOPSIS 权重或设备候选规格，容量优化结果需重新运行并同步更新论文表格。",
        ]:
            cur = insert_paragraph_after(cur, line)


def main() -> None:
    shutil.copy2(SRC, OUT)
    doc = Document(str(OUT))
    update_cover_and_abstracts(doc)
    update_research_status(doc)
    update_captions_and_static_text(doc)
    update_key_paragraphs(doc)
    update_tables(doc)
    update_references_and_appendix(doc)
    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    main()
