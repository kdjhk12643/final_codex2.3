from __future__ import annotations

import copy
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path.cwd()
INPUT = next(p for p in ROOT.glob("*10*.docx") if not p.name.startswith("~$"))
OUTPUT = ROOT / "毕业设计论文初稿11_终稿优先问题修订稿.docx"


def set_paragraph_text(paragraph, text: str, font_name: str | None = None, font_size: float | None = None) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    if font_name:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if font_size:
        run.font.size = Pt(font_size)


def insert_paragraph_after(paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    new_para.style = paragraph.style
    set_paragraph_text(new_para, text)
    return new_para


def find_para(doc: Document, contains: str) -> Paragraph:
    for p in doc.paragraphs:
        if contains in p.text:
            return p
    raise ValueError(f"paragraph not found: {contains}")


def replace_first(doc: Document, old: str, new: str) -> None:
    for p in doc.paragraphs:
        if p.text == old:
            set_paragraph_text(p, new)
            return
    raise ValueError(f"exact paragraph not found: {old[:40]}")


def replace_contains(doc: Document, needle: str, new: str) -> None:
    p = find_para(doc, needle)
    set_paragraph_text(p, new)


def set_cell_text(cell, text: str, size: float | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    if size:
        run.font.size = Pt(size)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)


def add_data_scope_row(doc: Document) -> None:
    table = doc.tables[2]
    existing = "\n".join(cell.text for row in table.rows for cell in row.cells)
    if "数据口径" in existing:
        return
    row = table.add_row()
    values = [
        "数据口径",
        "案例化运行数据",
        "变量范围和变化规律依据公开文献、客流规律及环控系统机理构建；不等同于运营单位原始BMS/AFC实测数据",
    ]
    for cell, value in zip(row.cells, values):
        set_cell_text(cell, value, 9.5)


def add_cant_split_to_table(table) -> None:
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = tr_pr.find(qn("w:cantSplit"))
        if cant_split is None:
            tr_pr.append(OxmlElement("w:cantSplit"))


def compact_table(table, font_size: float = 10.0) -> None:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    run.font.size = Pt(font_size)


def improve_formula_paragraphs(doc: Document) -> None:
    replacements = {
        ("Q = Qp + Qf + Qe + Qm", "（2-1）"): "Q = Q_p + Q_f + Q_e + Q_m\t（2-1）",
        ("r = cov(X,Y) / (sigma_X sigma_Y)", "（2-2）"): "r_{XY} = cov(X,Y) / (σ_X σ_Y)\t（2-2）",
        ("J = sum sum ||xi - mu_k||^2", "（2-3）"): "J = Σ_{k=1}^{K} Σ_{x_i∈C_k} ||x_i - μ_k||²\t（2-3）",
        ("R = (C - Qmax) / Qmax × 100%", "（2-4）"): "R = (C - Q_max) / Q_max × 100%\t（2-4）",
        ("f1 = min LCC", "（2-5）"): "f_1 = min LCC\t（2-5）",
        ("f2 = min R", "（2-6）"): "f_2 = min R\t（2-6）",
        ("RMSE = sqrt(1/n sum(yi - yhat_i)^2)", "（4-4）"): "RMSE = √[(1/n) Σ_{i=1}^{n}(y_i - ŷ_i)²]\t（4-4）",
        ("MAE = 1/n sum |yi - yhat_i|", "（4-5）"): "MAE = (1/n) Σ_{i=1}^{n}|y_i - ŷ_i|\t（4-5）",
        ("MAPE = 1/n sum |(yi - yhat_i) / yi| × 100%", "（4-6）"): "MAPE = (1/n) Σ_{i=1}^{n}|(y_i - ŷ_i)/y_i| × 100%\t（4-6）",
        ("R² = 1 - sum(yi - yhat_i)^2 / sum(yi - ybar)^2", "（4-7）"): "R² = 1 - [Σ_{i=1}^{n}(y_i - ŷ_i)²] / [Σ_{i=1}^{n}(y_i - ȳ)²]\t（4-7）",
        ("Q_d = Q_{0.99} + RMSE_LSTM", "（5-1）"): "Q_d = Q_{0.99} + RMSE_LSTM\t（5-1）",
        ("LCC = C_0 + sum_{t=1}^{N} C_{o,t}/(1+r)^t + sum_{t=1}^{N} C_{m,t}/(1+r)^t", "（5-2）"): "LCC = C_0 + Σ_{t=1}^{N} C_{o,t}/(1+r)^t + Σ_{t=1}^{N} C_{m,t}/(1+r)^t\t（5-2）",
        ("C_0 = C_ch + C_f + C_p + C_ahu = sum_k n_k q_k u_k", "（5-3）"): "C_0 = C_ch + C_f + C_p + C_ahu = Σ_k n_k q_k u_k\t（5-3）",
        ("f1 = min LCC", "（5-4）"): "f_1 = min LCC\t（5-4）",
        ("R_c = (C_c - alpha Q_d) / (alpha Q_d)", "（5-5）"): "R_c = (C_c - αQ_d) / (αQ_d)\t（5-5）",
        ("R = w_c R_c + w_f R_f + w_p R_p + w_ahu R_ahu", "（5-6）"): "R = w_cR_c + w_fR_f + w_pR_p + w_ahuR_ahu\t（5-6）",
        ("f2 = min R", "（5-7）"): "f_2 = min R\t（5-7）",
        ("C_c >= alpha Q_d", "（5-8）"): "C_c ≥ αQ_d\t（5-8）",
        ("C_f >= alpha D_f", "（5-9）"): "C_f ≥ αD_f\t（5-9）",
        ("C_p >= alpha D_p", "（5-10）"): "C_p ≥ αD_p\t（5-10）",
        ("C_ahu >= alpha D_ahu", "（5-11）"): "C_ahu ≥ αD_ahu\t（5-11）",
    }

    for p in doc.paragraphs:
        stripped = p.text.strip()
        number_match = re.search(r"（\d+-\d+）", stripped)
        if not number_match:
            continue
        num = number_match.group(0)
        key = stripped.replace("\t", " ")
        key = re.sub(r"\s+（\d+-\d+）$", "", key).strip()
        if (key, num) in replacements:
            text = "\t" + replacements[(key, num)]
            set_paragraph_text(p, text, font_name="Cambria Math", font_size=12)


def insert_page_break_before_caption(doc: Document, caption_text: str) -> None:
    p = find_para(doc, caption_text)
    if p.text.startswith("\f"):
        return
    old = p.text
    p.clear()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)
    p.add_run(old)


def shrink_pareto_figure(doc: Document) -> None:
    if len(doc.inline_shapes) >= 15:
        fig = doc.inline_shapes[14]
        fig.width = Inches(4.9)
        fig.height = Inches(3.28)


def main() -> None:
    shutil.copyfile(INPUT, OUTPUT)
    doc = Document(str(OUTPUT))

    # 1. 文献引用：补齐正文与参考文献的对应关系，尤其是绿色低碳与气流组织类文献。
    replace_contains(
        doc,
        "近年的研究进一步关注环控系统能耗诊断",
        "在节能控制和运行诊断方面，国内研究主要集中在控制策略优化、控制参数调整、用能诊断和既有车站改造适配等方向。相关综述认为，地铁站通风空调控制系统的节能优化可从控制策略、控制参数和智能控制方法三个层面展开[8]；热舒适与节能策略研究指出，站内环境控制不能只追求低温，还应兼顾乘客热舒适、空气品质和能耗水平[9]。近年的研究进一步关注环控系统能耗诊断、控制方案对比、前馈控制、智能环控系统、绿色低碳技术、气流组织模拟以及制式改造可行性，为既有车站节能改造和新建车站系统设计提供了工程参考[10-18,21-22]。",
    )

    # 2. 数据来源与真实性口径：摘要、数据章节、结论统一改为“案例化运行数据”。
    replace_contains(
        doc,
        "研究基于福州地铁东街口站 2025 年全年 15 min 粒度运行数据",
        "研究以福州地铁东街口站为案例背景，使用依据公开文献、客流变化规律和环控系统热工机理构建的 2025 年全年 15 min 粒度案例化运行数据，首先完成时间对齐、缺失值处理、标准化和滞后特征构造；随后采用 Pearson 相关系数法筛选关键影响因素，并将总冷负荷分解为人员负荷、新风负荷、围护结构负荷和设备散热负荷；在此基础上利用 K-Means 聚类识别典型日负荷模式，建立两层 LSTM 神经网络负荷预测模型，并与 BP 神经网络进行对比。最后，以全生命周期成本最小和综合容量冗余率最小为目标，采用 NSGA-II 求解 Pareto 候选方案，并通过 TOPSIS 方法确定推荐容量配置方案。",
    )
    replace_contains(
        doc,
        "结果表明，LSTM 模型在测试集上的 RMSE",
        "结果表明，在本文序列输入 LSTM 与静态 BP 基准模型的对比条件下，LSTM 模型在测试集上的 RMSE、MAE、MAPE 和 R² 分别为 9.58 kW、7.20 kW、3.86% 和 0.9709，预测误差低于 BP 基准模型；推荐优化方案使冷源总容量由 560 kW 降至 480 kW，综合容量冗余率由 33.35% 降至 14.05%，全生命周期成本降低 10.81%，年运行能耗降低 7.09%。研究结果可为地铁车站通风空调系统容量配置、节能设计和既有车站改造方案比选提供参考。",
    )
    replace_contains(
        doc,
        "Using a 15-minute annual operating dataset",
        "This thesis studies the ventilation and air-conditioning system of a metro station platform area and proposes a load-characteristic-based capacity configuration optimization method. Using a 15-minute annual case dataset constructed for Fuzhou Metro Dongjiekou Station in 2025 based on public literature, passenger-flow patterns and HVAC thermal mechanisms, the study completes data preprocessing, Pearson correlation analysis, load component decomposition and K-Means daily load clustering. A two-layer LSTM model is then established for cooling-load prediction and compared with a static BP baseline model. Based on the predicted load demand, a multi-objective capacity optimization model is built with lifecycle cost and composite capacity redundancy as objectives. NSGA-II is used to generate Pareto candidate schemes, and TOPSIS is used to select the recommended configuration. Under this comparison setting, the LSTM model achieves RMSE = 9.58 kW, MAE = 7.20 kW, MAPE = 3.86% and R² = 0.9709 on the test set. The recommended scheme reduces total cooling capacity from 560 kW to 480 kW, composite redundancy from 33.35% to 14.05%, lifecycle cost by 10.81%, and annual operating energy consumption by 7.09%. The proposed method can support energy-saving design and retrofit decision-making for metro station HVAC systems.",
    )
    replace_contains(
        doc,
        "为增强数据来源和复现口径的可追溯性",
        "为增强数据来源和复现口径的可追溯性，本文进一步对原始案例数据集和建模数据集的基本情况进行汇总。案例数据覆盖 2025 年全年 15 min 时间尺度，经过时间对齐、缺失修复和特征构造后形成建模数据，具体如表3-1所示。",
    )
    insert_paragraph_after(
        find_para(doc, "案例数据覆盖 2025 年全年 15 min 时间尺度"),
        "需要说明的是，本文使用的数据集用于方法链验证和容量配置案例分析，变量范围、日内变化和季节变化主要依据公开文献中地铁车站负荷规律、客流变化特征及通风空调系统热工机理构建，并不等同于运营单位 BMS、AFC 或传感器系统直接导出的原始实测数据。后续工程应用时，应以具体车站的实测 BMS 运行数据、AFC 客流数据和气象监测数据进行校准。",
    )
    replace_contains(
        doc,
        "经过上述处理后，数据集中的缺失值和明显异常值均已完成修复",
        "经过上述处理后，案例数据集中的缺失值和明显异常值均已完成修复，最终形成覆盖全年、共 35040 条 15 min 粒度样本的建模数据集。该数据集既保留了负荷峰谷变化、季节变化和日周期结构，也降低了单点异常值对后续分析结果的干扰。",
    )
    replace_contains(
        doc,
        "本章基于福州地铁东街口站 2025 年全年 15 min 粒度运行数据",
        "本章基于以福州地铁东街口站为案例背景构建的 2025 年全年 15 min 粒度案例化运行数据，对站台区域通风空调系统负荷进行了系统预处理和特性分析。首先完成时间戳统一、等间隔对齐、缺失值分级填补、异常值校核、Z-Score 标准化、周期编码和滞后特征构造，形成覆盖全年、共 35040 条样本的建模数据集。随后利用 Pearson 相关系数并结合工程机理筛选关键影响因素，结果表明历史负荷、客流、二氧化碳浓度、站台温度和太阳辐射均与总冷负荷密切相关。进一步的负荷分项分解表明，围护结构负荷和设备散热负荷是站台区域基础负荷的主要来源。最后，采用 K-Means 对日负荷曲线进行聚类，并综合轮廓系数与工程解释性确定 K = 4，识别出工作日双峰高负荷、工作日平稳中负荷、周末午后单峰和低客流低负荷四类典型模式，为下一章负荷预测模型输入设计和后续容量配置优化提供了更充分的数据依据。",
    )
    replace_contains(
        doc,
        "基于福州地铁东街口站 2025 年全年 15 分钟粒度运行数据",
        "基于以福州地铁东街口站为案例背景构建的 2025 年全年 15 分钟粒度案例化运行数据，采用 Pearson 相关分析、负荷分项分解、K-Means 聚类、LSTM 神经网络、NSGA-II 多目标优化和 TOPSIS 综合评价等方法，开展了地铁车站通风空调系统容量配置优化研究。主要结论如下。",
    )
    add_data_scope_row(doc)

    # 3. 公式：把程序式写法改为规范数学表达。
    improve_formula_paragraphs(doc)

    # 4. LSTM 与 BP 结论：强调“当前输入设置下”的比较，避免绝对化。
    replace_contains(
        doc,
        "图4-6 从 RMSE、MAE 和 MAPE",
        "图4-6 从 RMSE、MAE 和 MAPE 三类误差指标对两种模型进行比较。在本文“序列输入 LSTM”与“静态特征 BP 基准模型”的对比设置下，LSTM 各项误差均低于 BP，说明将历史负荷与多变量特征组织为序列输入后，模型对负荷波动和峰值偏差的刻画能力更强。",
    )
    replace_contains(
        doc,
        "从预测曲线角度看，LSTM 模型对总体变化趋势",
        "从预测曲线角度看，LSTM 模型对总体变化趋势和高负荷区间的跟踪能力更强，预测曲线与真实负荷曲线的偏差较小。BP 模型虽然能够反映负荷总体水平，但由于本文将其定位为静态基准模型，未采用显式序列记忆结构，因此在负荷快速上升或下降阶段更容易出现滞后。该结果说明时序结构和滞后信息组织对地铁车站负荷预测具有重要作用，但不应简单理解为两类模型在任意输入条件下的绝对优劣。",
    )
    replace_contains(
        doc,
        "表4-2 从绝对误差、相对误差和拟合优度",
        "表4-2 从绝对误差、相对误差和拟合优度三个角度比较 LSTM 与 BP 模型。LSTM 在当前序列输入设置下的 RMSE、MAE、MAPE 和 R² 表现均优于静态 BP 基准模型。",
    )
    replace_contains(
        doc,
        "由表4-2可知，LSTM 模型在各项评价指标上均优于 BP",
        "由表4-2可知，在本文设定的输入特征和数据划分条件下，LSTM 模型各项评价指标均优于静态 BP 基准模型。LSTM 的 RMSE 为 9.58 kW，低于 BP 模型的 22.41 kW；LSTM 的 MAE 为 7.20 kW，低于 BP 模型的 16.03 kW；LSTM 的 MAPE 为 3.86%，低于 BP 模型的 8.67%；LSTM 的 R² 为 0.9709，高于 BP 模型的 0.8406。",
    )
    replace_contains(
        doc,
        "从误差降低幅度来看，LSTM 相比 BP 模型的 RMSE",
        "从误差降低幅度来看，LSTM 相比静态 BP 基准模型的 RMSE 降低约 57.26%，MAE 降低约 55.08%，MAPE 降低约 55.47%。该结果主要说明在地铁车站负荷具有明显连续性和滞后性的条件下，序列输入和门控记忆结构能够提升预测表现。若 BP 模型同样引入充分滞后特征或其他时序特征，其性能可能会接近或改善，因此本文结论更强调输入组织方式和时序建模对容量优化前置负荷预测的价值。",
    )
    replace_contains(
        doc,
        "BP 神经网络虽然能够拟合非线性映射关系",
        "BP 神经网络虽然能够拟合非线性映射关系，但本文中的 BP 输入为静态特征，难以充分表达连续时序中的状态演化。因此，在负荷快速变化或高峰转换时段，BP 模型更容易出现滞后或偏差。相比之下，LSTM 模型能够更好地跟踪负荷变化趋势，对后续容量配置优化具有更高适用性。需要注意的是，该判断限定于本文实验设定，并不排除含滞后项 BP 或其他机器学习模型在补充特征后取得更好效果的可能。",
    )
    replace_contains(
        doc,
        "测试结果表明，LSTM 模型的 RMSE 为 9.58 kW",
        "测试结果表明，LSTM 模型的 RMSE 为 9.58 kW，MAPE 为 3.86%，R² 为 0.9709，在当前序列输入与静态 BP 基准对比条件下整体预测精度更优。说明 LSTM 能够较好捕捉地铁车站空调负荷的时序变化规律，可作为后续容量配置优化的负荷输入基础；同时，后续研究仍可进一步比较含滞后特征 BP、GRU 或 Transformer 等模型，以检验不同输入信息量下的预测差异。",
    )
    replace_contains(
        doc,
        "第四，LSTM 神经网络能够较好预测地铁车站站台区域空调负荷",
        "第四，LSTM 神经网络在本文设定的序列输入条件下能够较好预测地铁车站站台区域空调负荷。本文以客流、气象、站内环境、时间特征和历史负荷项作为输入，构建两层 LSTM 负荷预测模型，并与静态 BP 神经网络基准模型进行对比。测试结果表明，LSTM 模型 RMSE 为 9.58 kW，MAE 为 7.20 kW，MAPE 为 3.86%，R² 为 0.9709，整体预测误差低于 BP 基准模型。该结果说明时序特征组织对地铁车站负荷预测具有重要作用，可作为容量配置优化的负荷输入基础。",
    )

    # 5. 表5-5跨页：压缩图5-2和表5-5，并禁止表格行内分页；若空间仍不足，渲染后再处理。
    shrink_pareto_figure(doc)
    compact_table(doc.tables[13], 10.0)
    add_cant_split_to_table(doc.tables[13])

    doc.save(str(OUTPUT))
    print(OUTPUT)


if __name__ == "__main__":
    main()
