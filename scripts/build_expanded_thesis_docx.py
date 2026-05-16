from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"D:\桌面\codex2.3")
SRC = ROOT / "docs" / "thesis_chapter1_2.md"
OUT = ROOT / "output" / "基于负荷特性的地铁车站通风空调系统容量配置优化研究_本科论文终稿.docx"


def read_source() -> str:
    text = SRC.read_text(encoding="utf-8")
    ref = re.search(r"\n## 本阶段参考资料\n(?P<body>.*?)(?=\n## 第3章)", text, flags=re.S)
    refs = ""
    if ref:
        refs = ref.group("body").strip()
        text = text[: ref.start()] + text[ref.end() - len("\n## 第3章") :]
    text = text.strip()
    if refs:
        text += "\n\n## 参考文献\n\n" + refs
    return text


def inject_expansions(text: str) -> str:
    insertions = {
        "### 1.3 国内外研究现状": """
在正式梳理研究现状之前，有必要说明本文文献综述的组织逻辑。本文研究对象既涉及地铁车站通风空调系统，也涉及负荷预测和容量优化方法，因此研究现状不宜仅按单一学科分类展开。本文将相关研究归纳为系统节能、负荷特性、负荷预测和容量优化四个方面，并在每一类研究之后总结其对本文工作的启示。
""",
        "### 1.4 研究内容与技术路线": """
#### 1.4.1 本文研究思路

本文研究思路可以概括为“先认识负荷，再预测负荷，最后优化容量”。其中，负荷特性分析用于回答“系统负荷由什么决定”，负荷预测模型用于回答“未来或典型工况下负荷是多少”，容量配置优化用于回答“设备容量和台数组合应如何确定”。三个环节之间具有递进关系：若缺少负荷特性分析，预测模型输入缺乏依据；若缺少负荷预测，容量优化只能依赖保守峰值；若缺少容量优化，负荷预测结果难以转化为工程设计方案。

#### 1.4.2 技术路线

【图1-1 技术路线图：数据收集与预处理 → 负荷影响因素筛选 → 负荷分项分解 → 典型负荷曲线聚类 → LSTM负荷预测 → NSGA-II容量优化 → TOPSIS方案排序 → 工程评价】

#### 1.4.3 本文主要创新点

相较于单独开展负荷预测或单独进行设备容量估算的研究，本文的特点主要体现在以下三个方面。第一，构建了负荷特性分析、负荷预测和容量优化相衔接的方法链，使数据分析结果能够直接服务于设备选型。第二，将 LSTM 时序预测结果引入容量配置优化过程，减少单纯依赖经验峰值带来的过度保守。第三，在容量优化中同时考虑全生命周期成本和综合容量冗余率，并结合 TOPSIS 方法输出可用于工程比选的推荐方案。
""",
        "### 2.1 地铁车站通风空调系统组成": """
【图2-1 地铁车站通风空调系统组成示意图：冷水机组、冷冻水泵、冷却水泵、空气处理机组、站台送风机、站台排风机、站台区域】

为便于后续容量优化建模，本文将系统设备按功能划分为四类：冷源侧设备、空气侧设备、水系统设备和空气处理设备。该划分方式既符合地铁环控系统工程设计习惯，也便于建立设备容量、台数和系统负荷之间的对应关系。
""",
        "### 2.2 站台区域空调负荷构成": """
在实际工程计算中，人员负荷、新风负荷、围护结构负荷和设备散热负荷可分别按下列思路估算。人员负荷通常与站台人数、人体显热和潜热指标有关；新风负荷与新风量及室内外空气焓差有关；围护结构负荷与传热系数、传热面积和温差有关；设备散热负荷与设备功率及同时使用系数有关。上述负荷共同决定站台区域总冷负荷，也是后续容量配置的基础。

表2-1 站台区域负荷组成及影响因素

| 负荷类型 | 主要来源 | 主要影响因素 | 对容量配置的意义 |
| --- | --- | --- | --- |
| 人员负荷 | 乘客显热、潜热 | 进出站客流、站台滞留人数、停留时间 | 决定高峰时段负荷波动 |
| 新风负荷 | 室外新风处理 | 新风量、室外温湿度、站内空气品质要求 | 影响空调处理能力和除湿负担 |
| 围护结构负荷 | 地下结构传热 | 土壤温度、结构形式、温差 | 形成相对稳定的基础负荷 |
| 设备散热负荷 | 照明、电扶梯、屏蔽门等 | 设备功率、运行时间、同时使用系数 | 构成长期持续负荷 |
""",
        "### 2.5 容量配置优化问题描述": """
表2-2 容量配置优化对象

| 子系统 | 优化对象 | 决策内容 | 主要约束 |
| --- | --- | --- | --- |
| 冷源系统 | 冷水机组 | 单机容量、台数、总装机容量 | 满足设计冷负荷和安全裕度 |
| 风系统 | 送排风机 | 单机容量、台数、总风量或功率 | 满足站台送排风需求 |
| 水系统 | 冷冻水泵、冷却水泵 | 单机容量、台数、总输配能力 | 与冷源容量和管网阻力匹配 |
| 空气处理系统 | 空气处理机组 | 单台风量、台数、总处理能力 | 满足空气处理和新风需求 |

容量配置优化的难点在于各类设备之间存在耦合关系。例如，冷水机组容量降低后，水泵和空气处理机组的能力也应相应匹配；若仅减少冷源容量而不调整风机、水泵容量，系统仍可能存在辅助设备偏大的问题。因此，本文采用综合冗余率描述系统整体配置水平，而不是只评价冷水机组容量。
""",
        "### 3.1 数据来源与变量说明": """
表3-1 数据变量说明表

| 变量类别 | 变量名称 | 含义 | 单位 | 是否用于建模 |
| --- | --- | --- | --- | --- |
| 客流特征 | entry_flow | 进站客流 | 人/15min | 是 |
| 客流特征 | exit_flow | 出站客流 | 人/15min | 是 |
| 客流特征 | platform_passengers | 站台人数 | 人 | 是 |
| 气象特征 | outdoor_temp | 室外温度 | ℃ | 是 |
| 气象特征 | outdoor_rh | 室外相对湿度 | % | 是 |
| 气象特征 | solar_radiation | 太阳辐射 | W/m² | 是 |
| 站内环境 | platform_temp | 站台温度 | ℃ | 是 |
| 站内环境 | platform_rh | 站台相对湿度 | % | 是 |
| 站内环境 | co2 | 二氧化碳浓度 | ppm | 是 |
| 负荷变量 | total_cooling_load_kw | 站台总冷负荷 | kW | 预测目标 |

【图3-1 数据预处理流程图：原始数据 → 时间排序 → 缺失值填补 → 时间特征构造 → 滞后特征构造 → Z-Score标准化 → 建模数据集】
""",
        "### 3.2 数据清洗与预处理": """
【图3-2 缺失值统计图：对应 output/figures/step1_missing_value_summary.png】

【图3-3 站台总冷负荷时间序列图：对应 output/figures/step1_cooling_load_timeseries.png】

从负荷时间序列可以进一步观察日内峰谷变化。一般而言，站台区域总冷负荷在客流高峰和环境温度较高时段出现上升，在深夜和低客流时段下降。该变化规律说明负荷序列同时具有周期性、滞后性和随机波动性，因此后续预测模型需要同时引入时间特征、客流特征和历史负荷特征。
""",
        "### 3.3 基于 Pearson 相关系数的关键影响因素筛选": """
【图3-4 Pearson相关性排序图：对应 output/figures/step2_pearson_feature_ranking.png】

需要注意的是，Pearson 相关系数主要反映变量之间的线性相关程度，并不能完全代表因果关系。因此，本文在选择预测输入时不仅依据相关系数大小，还结合地铁车站通风空调系统的工程机理进行判断。例如，客流变量与负荷高度相关，既有统计依据，也符合人员散热和新风需求随客流增加而上升的工程规律。
""",
        "### 3.4 负荷分项分解": """
【图3-5 负荷分项占比图：对应 output/figures/step2_load_component_ratio.png】

负荷分项结果对容量配置具有直接启示。围护结构负荷和设备散热负荷占比较高，说明站台区域存在较大的基础负荷，即使在客流较低时段，系统仍需维持一定供冷和通风能力。人员负荷占比虽然相对较小，但其波动性强，是导致峰值负荷变化的重要因素。因此，容量配置既要保证基础负荷需求，也要考虑客流峰值带来的短时波动。
""",
        "### 3.5 基于 K-Means 的典型负荷曲线聚类": """
【图3-6 典型日负荷曲线聚类图：对应 output/figures/step2_daily_load_clusters.png】

聚类结果还可用于解释不同运行场景下的设备调节需求。工作日双峰高负荷型需要系统具备较强的高峰响应能力；工作日平稳中负荷型更关注部分负荷效率；周末午后单峰型需要设备具备灵活启停能力；低客流低负荷型则要求系统避免大容量设备长时间低效运行。这些差异为后续容量配置优化提供了工程依据。
""",
        "### 4.2 预测输入特征构建": """
【图4-1 预测样本构造示意图：连续16个时间步多维输入 → 下一时刻总冷负荷】

本文选取 16 个时间步作为 LSTM 输入序列长度，对应 4 小时历史窗口。该设置可以覆盖地铁车站早晚高峰变化前后的局部动态过程，同时避免序列过长导致模型训练复杂度过高。若后续引入全年数据，可进一步通过验证集比较不同序列长度的预测效果。
""",
        "### 4.3 LSTM 神经网络模型构建": """
【图4-2 LSTM网络结构图：Sequence Input → LSTM(64) → LSTM(32) → Fully Connected → Regression Output】

表4-1 LSTM模型主要训练参数

| 参数 | 取值 | 说明 |
| --- | ---: | --- |
| 输入序列长度 | 16 | 对应4小时历史信息 |
| 第一层LSTM单元数 | 64 | 提取主要时序特征 |
| 第二层LSTM单元数 | 32 | 输出最终时序状态 |
| 最大训练轮数 | 60 | 控制训练上限 |
| 批量大小 | 64 | 小批量训练 |
| 初始学习率 | 0.001 | Adam优化器参数 |
| 训练/验证/测试比例 | 7:2:1 | 按时间顺序划分 |
""",
        "### 4.6 预测结果与对比分析": """
【图4-3 LSTM预测结果图：对应 output/figures/step3_lstm_prediction.png】

【图4-4 BP预测结果图：对应 output/figures/step3_bp_prediction.png】

【图4-5 模型RMSE对比图：对应 output/figures/step3_model_rmse_comparison.png】

从预测曲线角度看，LSTM 模型对总体变化趋势和高负荷区间的跟踪能力更强，预测曲线与真实负荷曲线的偏差较小。BP 模型虽然能够反映负荷总体水平，但在负荷快速上升或下降阶段更容易出现滞后。对于容量配置优化而言，高峰负荷和典型负荷区间的预测准确性尤为重要，因为它们直接影响设备容量边界和生命周期能耗估算。
""",
        "### 5.1 容量配置优化对象与决策变量": """
表5-1 设备候选容量与台数范围

| 设备类型 | 候选容量或风量 | 台数范围 |
| --- | --- | --- |
| 冷水机组 | 120、150、180、200、220、250、280、300、320、350、380 kW | 1-4 |
| 风机 | 18、22、26、30、35、40、45、50 kW | 1-6 |
| 水泵 | 16、20、24、28、32、36、40 kW | 1-6 |
| 空气处理机组 | 25000、30000、35000、40000、45000、50000、55000、60000 m³/h | 1-4 |

表5-2 经济性计算参数

| 参数 | 取值 | 说明 |
| --- | ---: | --- |
| 电价 | 0.85 元/kWh | 运行电费估算 |
| 使用年限 | 15 年 | 生命周期计算周期 |
| 折现率 | 5% | 折现计算 |
| 制冷季天数 | 120 天 | 年运行能耗估算 |
| 维护费率 | 3.5% | 按初投资比例估算 |
""",
        "### 5.5 NSGA-II 多目标优化模型": """
【图5-1 NSGA-II容量优化流程图：初始化种群 → 解码设备方案 → 约束校核 → 目标函数计算 → 非支配排序 → 拥挤距离计算 → 选择交叉变异 → Pareto解集】
""",
        "### 5.6 基于 TOPSIS 的综合最优方案选择": """
表5-3 TOPSIS排序结果

| 排名 | 全生命周期成本/元 | 综合容量冗余率 | TOPSIS得分 |
| ---: | ---: | ---: | ---: |
| 1 | 3711458 | 9.72% | 1.0000 |
| 2 | 3862830 | 9.72% | 0.9511 |
| 3 | 3788934 | 12.44% | 0.8169 |
| 4 | 3940306 | 12.44% | 0.8038 |
| 5 | 3805742 | 12.91% | 0.7852 |
""",
        "### 5.7 优化结果概述": """
【图5-2 Pareto前沿图：对应 output/figures/step4_pareto_front.png】

Pareto 前沿能够直观反映全生命周期成本与容量冗余率之间的权衡关系。一般而言，容量冗余率越低，设备配置越接近负荷需求，但部分方案可能需要更复杂的设备组合；生命周期成本越低，经济性越好，但也需要保证安全裕度和系统匹配。TOPSIS 方法的作用就是在这两类指标之间进行综合折中，选择更适合工程应用的方案。
""",
        "### 6.1 基准方案与优化方案对比": """
【图6-1 基准方案与优化方案对比图：对应 output/figures/step4_scheme_comparison.png】

表6-2 设备配置明细表

| 方案 | 冷水机组 | 风机 | 水泵 | 空气处理机组 |
| --- | --- | --- | --- | --- |
| 基准方案 | 220 kW × 2 台 | 50 kW × 1 台 | 36 kW × 1 台 | 50000 m³/h × 1 台 |
| 优化方案 | 120 kW × 3 台 | 40 kW × 1 台 | 28 kW × 1 台 | 45000 m³/h × 1 台 |
""",
        "### 6.4 年运行能耗与节能效果分析": """
按照年节电量 97452.2 kWh 和电价 0.85 元/kWh 估算，优化方案每年可节约运行电费约 8.28 万元。若进一步考虑设备维护费用随容量下降而减少，则优化方案的长期经济收益会更加明显。该结果说明，容量配置优化不仅体现在一次性初投资下降，也会在系统长期运营阶段持续产生节能和降费效果。
""",
        "### 7.2 研究不足": """
上述不足并不影响本文方法链的完整性和结果的相对比较意义，但提示后续研究应进一步增强数据覆盖范围和工程校核深度。在本科毕业设计阶段，本文重点完成了方法构建、模型计算和案例验证，已经能够支撑容量配置优化的基本结论。
""",
    }
    for marker, addition in insertions.items():
        text = text.replace(marker, marker + "\n" + addition.strip() + "\n", 1)
    return text


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) < 18 else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table_from_markdown(doc: Document, block: list[str]):
    rows = []
    for line in block:
        if not line.strip().startswith("|"):
            continue
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=max(len(r) for r in rows))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            set_cell_text(cell, val, bold=(i == 0))
            if i == 0:
                set_cell_shading(cell, "D9EAF7")
    doc.add_paragraph()


def add_paragraph(doc: Document, text: str):
    if not text.strip():
        return
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text.strip())
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(11)


def add_placeholder(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(80, 80, 80)


def add_captioned_picture(doc: Document, text: str, image_path: Path):
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image_path), width=Cm(13.5))
    caption = re.sub(r"：对应\s+output/figures/[^】]+", "", text)
    add_placeholder(doc, caption)


def add_abstract(doc: Document):
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = heading.add_run("摘要")
    r.bold = True
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(16)

    abstract = (
        "本文以地铁车站站台区域通风空调系统为研究对象，针对传统容量配置中设备容量偏大、"
        "系统长期低负荷运行和全生命周期成本较高等问题，构建了基于负荷特性的容量配置优化方法。"
        "研究首先对车站客流、气象、站内环境和系统运行数据进行预处理，采用 Pearson 相关系数法筛选关键影响因素，"
        "并将总冷负荷分解为人员负荷、新风负荷、围护结构负荷和设备散热负荷。随后，利用 K-Means 聚类识别典型日负荷模式，"
        "建立 LSTM 神经网络负荷预测模型，并与 BP 神经网络进行对比。最后，以全生命周期成本最小和容量冗余率最小为目标，"
        "采用 NSGA-II 算法求解 Pareto 解集，并通过 TOPSIS 方法确定推荐容量配置方案。结果表明，LSTM 模型测试集 RMSE 为 13.28 kW，"
        "R2 为 0.9557，预测效果优于 BP 模型；优化方案使冷源总容量由 440 kW 降至 360 kW，综合容量冗余率由 33.40% 降至 9.72%，"
        "全生命周期成本降低 21.23%，年运行能耗降低 21.93%。研究结果可为地铁车站通风空调系统节能设计和既有车站改造提供参考。"
    )
    add_paragraph(doc, abstract)
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run("关键词：地铁车站；通风空调系统；负荷预测；LSTM；NSGA-II；容量配置优化")
    r.bold = True
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(11)
    doc.add_page_break()


def configure_doc(doc: Document):
    sec = doc.sections[0]
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.5)
    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(11)
    for name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = styles[name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True


def add_cover(doc: Document):
    for _ in range(2):
        doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("基于负荷特性的地铁车站通风空调系统容量配置优化研究")
    run.bold = True
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("本科毕业论文终稿")
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(16)

    for label in ["学生姓名：__________", "专业班级：__________", "指导教师：__________", "完成日期：2026年5月"]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(18)
        r = p.add_run(label)
        r.font.name = "宋体"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        r.font.size = Pt(14)
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("目录")
    r.bold = True
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(16)
    add_placeholder(doc, "【目录说明：提交前可在 Word 中右键更新域或按学校模板生成目录】")
    doc.add_page_break()
    add_abstract(doc)


def build_docx(md: str):
    doc = Document()
    configure_doc(doc)
    add_cover(doc)

    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i])
                i += 1
            add_table_from_markdown(doc, block)
            continue
        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
        elif line.startswith("【") and line.endswith("】"):
            image_match = re.search(r"对应\s+(output/figures/[^】]+)", line)
            if image_match:
                add_captioned_picture(doc, line, ROOT / image_match.group(1))
            else:
                add_placeholder(doc, line)
        elif re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style=None)
            p.paragraph_format.left_indent = Pt(18)
            p.paragraph_format.line_spacing = 1.5
            r = p.add_run(line.strip())
            r.font.name = "宋体"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            r.font.size = Pt(11)
        else:
            add_paragraph(doc, line)
        i += 1

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)


if __name__ == "__main__":
    md = inject_expansions(read_source())
    (ROOT / "output" / "thesis_expanded_intermediate.md").write_text(md, encoding="utf-8")
    build_docx(md)
    print(OUT)
