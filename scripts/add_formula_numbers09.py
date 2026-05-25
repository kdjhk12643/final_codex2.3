from __future__ import annotations

import argparse
import shutil
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


CN_FONT = "宋体"
EN_FONT = "Times New Roman"


PURE_FORMULAS = {
    # Chapter 2
    "Q = Qp + Qf + Qe + Qm": "（2-1）",
    "r = cov(X,Y) / (sigma_X sigma_Y)": "（2-2）",
    "J = sum sum ||xi - mu_k||^2": "（2-3）",
    "R = (C - Qmax) / Qmax × 100%": "（2-4）",
    # Chapter 3
    "r_{X,Y} = Σ_{i=1}^{n}(x_i - x_bar)(y_i - y_bar) / sqrt[Σ_{i=1}^{n}(x_i - x_bar)^2 · Σ_{i=1}^{n}(y_i - y_bar)^2]": "（3-3）",
    "J = Σ_{k=1}^{K} Σ_{q'_d ∈ C_k} ||q'_d - μ_k||_2^2": "（3-4）",
    "s(i) = [b(i) - a(i)] / max{a(i), b(i)}": "（3-5）",
    # Chapter 4
    "RMSE = sqrt(1/n sum(yi - yhat_i)^2)": "（4-4）",
    "MAE = 1/n sum |yi - yhat_i|": "（4-5）",
    "MAPE = 1/n sum |(yi - yhat_i) / yi| × 100%": "（4-6）",
    "R² = 1 - sum(yi - yhat_i)^2 / sum(yi - ybar)^2": "（4-7）",
    # Chapter 5
    "Q_d = Q_{0.99} + RMSE_LSTM": "（5-1）",
    "LCC = C_0 + sum_{t=1}^{N} C_{o,t}/(1+r)^t + sum_{t=1}^{N} C_{m,t}/(1+r)^t": "（5-2）",
    "C_0 = C_ch + C_f + C_p + C_ahu = sum_k n_k q_k u_k": "（5-3）",
    "R_c = (C_c - alpha Q_d) / (alpha Q_d)": "（5-5）",
    "R = w_c R_c + w_f R_f + w_p R_p + w_ahu R_ahu": "（5-6）",
    "C_c >= alpha Q_d": "（5-8）",
    "C_f >= alpha D_f": "（5-9）",
    "C_p >= alpha D_p": "（5-10）",
    "C_ahu >= alpha D_ahu": "（5-11）",
}


# These two formulas occur once in Chapter 2 and once in Chapter 5.
DUPLICATE_FORMULA_NUMBERS = {
    "f1 = min LCC": ["（2-5）", "（5-4）"],
    "f2 = min R": ["（2-6）", "（5-7）"],
}


INLINE_REPLACEMENTS = [
    (
        "X = {x_t | t = t_0, t_0 + Δt, ..., t_n}",
        "X = {x_t | t = t_0, t_0 + Δt, ..., t_n}（3-1）",
    ),
    ("z_i = (x_i - μ_x)/σ_x", "z_i = (x_i - μ_x)/σ_x（3-2）"),
    (
        "X_t = [x_{t-T+1}, x_{t-T+2}, ..., x_t]，预测下一时刻站台区域总冷负荷 y_{t+1}",
        "X_t = [x_{t-T+1}, x_{t-T+2}, ..., x_t]，预测下一时刻站台区域总冷负荷 y_{t+1}（4-1）",
    ),
    ("x'_i = (x_i - μ_x) / σ_x", "x'_i = (x_i - μ_x) / σ_x（4-2）"),
    (
        "h_t = o_t ⊙ tanh(C_t)",
        "h_t = o_t ⊙ tanh(C_t)（4-3）",
    ),
    (
        "S_i = D_i^- / (D_i^+ + D_i^-)",
        "S_i = D_i^- / (D_i^+ + D_i^-)（5-12）",
    ),
]


MULTILINE_FORMULAS = {
    "r_{X,Y} = Σ_{i=1}^{n}(x_i - x_bar)(y_i - y_bar) / sqrt[Σ_{i=1}^{n}(x_i - x_bar)^2 · Σ_{i=1}^{n}(y_i - y_bar)^2]": [
        "r_{X,Y} = Σ_{i=1}^{n}(x_i - x_bar)(y_i - y_bar) /",
        "sqrt[Σ_{i=1}^{n}(x_i - x_bar)^2 · Σ_{i=1}^{n}(y_i - y_bar)^2]",
    ],
}


def set_run_font(run, size_pt: float = 11.5, bold: bool = False) -> None:
    run.font.name = EN_FONT
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)


def set_cell_margins(cell, top: int = 40, start: int = 80, bottom: int = 40, end: int = 80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def clear_paragraph_runs(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag in {qn("w:r"), qn("w:hyperlink")}:
            paragraph._p.remove(child)


def clear_tabs(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    tabs = p_pr.find(qn("w:tabs"))
    if tabs is not None:
        p_pr.remove(tabs)


def format_formula_paragraph(paragraph, formula: str, number: str) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = paragraph.paragraph_format
    pf.first_line_indent = Pt(0)
    pf.left_indent = Pt(0)
    pf.right_indent = Pt(0)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.5
    clear_tabs(paragraph)
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(3.15), WD_TAB_ALIGNMENT.CENTER
    )
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.25), WD_TAB_ALIGNMENT.RIGHT
    )
    clear_paragraph_runs(paragraph)
    lines = MULTILINE_FORMULAS.get(formula, [formula])
    size = 9.0 if len(formula) > 100 else (10.0 if len(formula) > 70 else 11.5)
    for line_idx, line in enumerate(lines):
        set_run_font(paragraph.add_run("\t"), size)
        set_run_font(paragraph.add_run(line), size)
        if line_idx < len(lines) - 1:
            paragraph.add_run().add_break()
    set_run_font(paragraph.add_run("\t"), size)
    set_run_font(paragraph.add_run(number), 11.5)


def compact_pearson_table(doc: Document) -> None:
    """Keep Table 3-4 from spilling into the next page header in LibreOffice render."""
    if len(doc.tables) <= 3:
        return
    table = doc.tables[3]
    if not table.rows or table.cell(0, 0).text.strip() != "特征":
        return
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pf = paragraph.paragraph_format
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                pf.line_spacing = 1.0
                for run in paragraph.runs:
                    set_run_font(run, 9.0, bool(run.bold))


def replace_preserving_runs(paragraph, old: str, new: str) -> bool:
    runs = list(paragraph.runs)
    text = "".join(run.text for run in runs)
    start = text.find(old)
    if start < 0:
        return False
    end = start + len(old)

    positions: list[tuple[int, int, object]] = []
    cursor = 0
    for run in runs:
        run_start = cursor
        run_end = cursor + len(run.text)
        positions.append((run_start, run_end, run))
        cursor = run_end

    first_idx = last_idx = None
    for idx, (run_start, run_end, _) in enumerate(positions):
        if first_idx is None and run_end > start:
            first_idx = idx
        if run_start < end <= run_end:
            last_idx = idx
            break
    if first_idx is None or last_idx is None:
        return False

    first_start, first_end, first_run = positions[first_idx]
    last_start, last_end, last_run = positions[last_idx]
    prefix = first_run.text[: start - first_start]
    suffix = last_run.text[end - last_start :]

    first_run.text = prefix + new + suffix
    for idx in range(first_idx + 1, last_idx + 1):
        positions[idx][2].text = ""
    return True


def add_formula_numbers(input_path: Path, output_path: Path) -> dict[str, int]:
    if input_path.resolve() != output_path.resolve():
        shutil.copyfile(input_path, output_path)

    doc = Document(output_path)
    paragraphs = list(doc.paragraphs)
    pure_count = 0
    inline_count = 0
    unchanged_formulas: list[str] = []
    duplicate_seen = {key: 0 for key in DUPLICATE_FORMULA_NUMBERS}

    for idx, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue

        if text in DUPLICATE_FORMULA_NUMBERS:
            seq = DUPLICATE_FORMULA_NUMBERS[text]
            number = seq[min(duplicate_seen[text], len(seq) - 1)]
            duplicate_seen[text] += 1
        else:
            number = PURE_FORMULAS.get(text)
        if number:
            format_formula_paragraph(paragraph, text, number)
            pure_count += 1
            continue

        for old, new in INLINE_REPLACEMENTS:
            if old in paragraph.text and new not in paragraph.text:
                if not replace_preserving_runs(paragraph, old, new):
                    unchanged_formulas.append(old)
                else:
                    inline_count += 1

    compact_pearson_table(doc)
    doc.save(output_path)
    return {
        "pure_count": pure_count,
        "inline_count": inline_count,
        "unchanged_count": len(unchanged_formulas),
    }


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("input")
    parser.add_argument("output")
    args = parser.parse_args()

    stats = add_formula_numbers(Path(args.input), Path(args.output))
    print(f"Wrote: {args.output}")
    print(
        "Formula numbers inserted: "
        f"pure={stats['pure_count']}, inline={stats['inline_count']}, "
        f"unchanged={stats['unchanged_count']}"
    )


if __name__ == "__main__":
    main()
