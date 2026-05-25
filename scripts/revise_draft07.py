from __future__ import annotations

import shutil
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from docx.text.paragraph import Paragraph


def text_of(paragraph: Paragraph) -> str:
    return "".join(run.text for run in paragraph.runs).strip()


def set_run_font(run, size_pt: float = 12, bold: bool = False) -> None:
    run.bold = bold
    run.font.size = Pt(size_pt)
    run.font.name = "Times New Roman"
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), "宋体")
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")


def clear_paragraph(paragraph: Paragraph) -> None:
    p_pr = paragraph._p.pPr
    for child in list(paragraph._p):
        if child is not p_pr:
            paragraph._p.remove(child)


def set_paragraph_text(paragraph: Paragraph, text: str, size: float = 12, bold: bool = False) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_font(run, size, bold)


def format_body(paragraph: Paragraph) -> None:
    paragraph.style = "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Pt(24)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        set_run_font(run, 12)


def find_paragraph(doc: Document, target: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if text_of(paragraph) == target:
            return paragraph
    raise ValueError(f"paragraph not found: {target}")


def find_paragraph_starts(doc: Document, prefix: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if text_of(paragraph).startswith(prefix):
            return paragraph
    raise ValueError(f"paragraph not found: {prefix}")


def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        run = new_para.add_run(text)
        set_run_font(run, 12)
    return new_para


def remove_following_empty_paragraphs(paragraph: Paragraph, limit: int = 4) -> None:
    node = paragraph._p.getnext()
    removed = 0
    while node is not None and removed < limit:
        next_node = node.getnext()
        if node.tag == qn("w:p"):
            candidate = Paragraph(node, paragraph._parent)
            if not text_of(candidate) and not node.findall(".//" + qn("w:sectPr")):
                node.getparent().remove(node)
                removed += 1
                node = next_node
                continue
        break


def create_toc_paragraph(text: str, page: str, level: int, doc: Document) -> OxmlElement:
    p = OxmlElement("w:p")
    para = Paragraph(p, doc)
    para.style = "Normal"
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.paragraph_format.line_spacing = 1.0
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.space_before = Pt(0)
    if level == 1:
        para.paragraph_format.left_indent = Inches(0.0)
        font_size = 12
        bold = True
    elif level == 2:
        para.paragraph_format.left_indent = Inches(0.32)
        font_size = 11.5
        bold = False
    else:
        para.paragraph_format.left_indent = Inches(0.64)
        font_size = 11
        bold = False
    para.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.35), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS
    )
    run = para.add_run(text)
    set_run_font(run, font_size, bold)
    run = para.add_run("\t" + page)
    set_run_font(run, font_size, bold)
    return p


def create_toc_title(doc: Document) -> OxmlElement:
    p = OxmlElement("w:p")
    para = Paragraph(p, doc)
    para.style = "Normal"
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after = Pt(12)
    run = para.add_run("目    次")
    set_run_font(run, 18, True)
    return p


def create_page_break_paragraph(doc: Document) -> OxmlElement:
    p = OxmlElement("w:p")
    para = Paragraph(p, doc)
    para.add_run().add_break(WD_BREAK.PAGE)
    return p


def replace_toc(doc: Document) -> None:
    toc_sdt = None
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:sdt") and "目" in "".join(child.itertext()):
            toc_sdt = child
            break
    if toc_sdt is None:
        raise ValueError("TOC sdt not found")

    entries_page_1 = [
        ("摘    要", "I", 1),
        ("Abstract:", "II", 1),
        ("1  绪论", "1", 1),
        ("1.1 研究背景", "1", 2),
        ("1.2 研究目的与意义", "1", 2),
        ("1.3 国内外研究现状", "2", 2),
        ("1.4 研究内容与技术路线", "3", 2),
        ("1.5 本章小结", "5", 2),
        ("2  地铁车站通风空调系统负荷特性与容量配置理论基础", "6", 1),
        ("2.1 地铁车站通风空调系统组成", "6", 2),
        ("2.2 站台区域空调负荷构成", "7", 2),
        ("2.3 负荷影响因素分析", "9", 2),
        ("2.4 典型负荷模式识别理论", "9", 2),
        ("2.5 容量配置优化问题描述", "10", 2),
        ("2.6 多目标优化与综合评价方法", "12", 2),
        ("2.7 本章小结", "13", 2),
        ("3  数据预处理与负荷特性分析", "14", 1),
        ("3.1 数据清洗与预处理", "14", 2),
        ("3.2 基于 Pearson 相关系数的关键影响因素筛选", "17", 2),
        ("3.3 负荷分项分解", "19", 2),
        ("3.4 基于 K-Means 的典型负荷曲线聚类", "21", 2),
        ("3.5 本章小结", "23", 2),
        ("4  基于 LSTM 的地铁车站空调负荷预测模型", "24", 1),
        ("4.1 负荷预测问题描述", "24", 2),
        ("4.2 预测输入特征构建", "25", 2),
        ("4.3 LSTM 神经网络模型构建", "26", 2),
    ]
    entries_page_2 = [
        ("4.4 BP 神经网络对比模型", "28", 2),
        ("4.5 模型评价指标", "29", 2),
        ("4.6 预测结果与对比分析", "30", 2),
        ("4.7 本章小结", "33", 2),
        ("5  地铁车站通风空调系统容量配置多目标优化", "34", 1),
        ("5.1 容量配置优化对象与决策变量", "34", 2),
        ("5.2 设计负荷确定", "36", 2),
        ("5.3 优化目标函数", "36", 2),
        ("5.3.1 全生命周期成本目标", "36", 3),
        ("5.3.2 容量冗余率目标", "38", 3),
        ("5.4 工程约束条件", "38", 2),
        ("5.5 NSGA-II 多目标优化模型", "40", 2),
        ("5.6 基于 TOPSIS 的综合最优方案选择", "41", 2),
        ("5.7 优化结果概述", "42", 2),
        ("5.8 本章小结", "44", 2),
        ("6  优化结果分析", "45", 1),
        ("6.1 基准方案与优化方案对比", "45", 2),
        ("6.2 容量冗余率分析", "47", 2),
        ("6.3 全生命周期成本分析", "47", 2),
        ("6.4 年运行能耗与节能效果分析", "48", 2),
        ("6.5 设备匹配性分析", "48", 2),
        ("6.6 调节灵活性与运行适应性分析", "49", 2),
        ("6.7 既有车站改造适配性分析", "49", 2),
        ("6.8 本章小结", "50", 2),
        ("7  结论", "51", 1),
        ("致  谢", "53", 1),
        ("参考文献", "54", 1),
        ("附录1  主要程序文件与数据字段说明", "57", 1),
    ]

    parent = toc_sdt.getparent()
    insert_before = toc_sdt
    for element in [create_toc_title(doc)] + [
        create_toc_paragraph(text, page, level, doc) for text, page, level in entries_page_1
    ] + [create_page_break_paragraph(doc)] + [
        create_toc_paragraph(text, page, level, doc) for text, page, level in entries_page_2
    ]:
        insert_before.addprevious(element)
    parent.remove(toc_sdt)


def set_table_widths(table, widths_cm: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)


def format_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing = 1.2
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    set_run_font(run, 10.5, False)
    for cell in table.rows[0].cells:
        for p in cell.paragraphs:
            for run in p.runs:
                set_run_font(run, 10.5, True)


def append_caption_after_table(table, text: str) -> Paragraph:
    tbl = table._tbl
    p = OxmlElement("w:p")
    tbl.addnext(p)
    para = Paragraph(p, table._parent)
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(3)
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text)
    set_run_font(run, 10.5)
    return para


def fill_appendix(doc: Document) -> None:
    heading = find_paragraph(doc, "附录1")
    set_paragraph_text(heading, "附录1  主要程序文件与数据字段说明", 14, True)
    heading.style = "Heading 1"
    remove_following_empty_paragraphs(heading)

    p = insert_paragraph_after(
        heading,
        "为便于复现实验过程，本附录对本文建模与优化过程中涉及的主要程序文件和关键数据字段进行说明。程序文件主要承担数据预处理、负荷特性分析、预测建模、容量优化和结果汇总等功能；数据字段则对应客流、环境、负荷和模型辅助特征。",
    )
    format_body(p)

    caption1 = insert_paragraph_after(p, "表A-1 主要程序文件及功能说明")
    caption1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption1.runs:
        set_run_font(run, 10.5)

    table1 = doc.add_table(rows=1, cols=3)
    caption1._p.addnext(table1._tbl)
    headers = ["程序文件", "主要功能", "对应章节"]
    rows = [
        ["config.m", "统一设置数据路径、候选容量、经济参数和优化约束", "第3-5章"],
        ["step1_data_prepare.m", "完成时间对齐、缺失值处理、异常值校核和特征构造", "第3.1节"],
        ["step2_analysis_cluster.m", "开展 Pearson 相关分析、负荷分项统计和典型日聚类", "第3.2-3.4节"],
        ["step3_load_prediction.m", "构建 LSTM 与 BP 预测模型，输出测试集评价指标", "第4章"],
        ["step4_capacity_optimization.m", "建立 NSGA-II 容量优化模型并进行 TOPSIS 排序", "第5章"],
        ["main.m", "串联数据处理、模型训练、优化求解和结果导出流程", "全文计算流程"],
    ]
    for i, h in enumerate(headers):
        table1.rows[0].cells[i].text = h
    for row in rows:
        cells = table1.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    set_table_widths(table1, [4.2, 8.0, 3.8])
    format_table(table1)

    exp1 = append_caption_after_table(
        table1,
        "表A-1 汇总了主要程序文件的功能分工，说明本文计算流程由数据准备、特征分析、负荷预测和容量优化四个环节构成。",
    )
    format_body(exp1)

    p2 = insert_paragraph_after(
        exp1,
        "关键字段按变量来源可分为时间字段、客流字段、环境字段、负荷字段和模型辅助字段。不同字段在模型中的作用不同：客流和环境变量用于刻画外部扰动，历史负荷变量用于反映系统惯性，聚类标签用于描述典型日运行模式。",
    )
    format_body(p2)

    caption2 = insert_paragraph_after(p2, "表A-2 关键数据字段及含义说明")
    caption2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption2.runs:
        set_run_font(run, 10.5)

    table2 = doc.add_table(rows=1, cols=3)
    caption2._p.addnext(table2._tbl)
    headers = ["字段或变量", "含义", "用途"]
    rows = [
        ["timestamp", "15 min 粒度时间戳", "统一客流、气象、站内环境和负荷数据时间尺度"],
        ["entry_flow / exit_flow", "进站客流和出站客流", "表征乘客活动强度及人员负荷变化"],
        ["platform_passengers", "站台滞留人数", "反映站台区域人员密度和新风需求"],
        ["outdoor_temp / outdoor_rh", "室外温度和相对湿度", "描述新风处理和围护结构传热边界"],
        ["platform_temp / platform_rh", "站台温度和相对湿度", "反映站内热湿环境状态"],
        ["co2", "站台二氧化碳浓度", "间接表征客流密度与通风需求"],
        ["total_cooling_load_kw", "站台区域总冷负荷", "作为预测目标和容量优化负荷基础"],
        ["load_lag_1 / load_lag_4 / load_lag_8 / load_lag_96", "前15 min、前1 h、前2 h和前1 d同刻负荷", "刻画负荷短时惯性和日周期记忆"],
        ["cluster_label", "典型日聚类标签", "作为辅助特征区分不同日负荷模式"],
    ]
    for i, h in enumerate(headers):
        table2.rows[0].cells[i].text = h
    for row in rows:
        cells = table2.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    set_table_widths(table2, [5.1, 5.3, 5.6])
    format_table(table2)

    exp2 = append_caption_after_table(
        table2,
        "表A-2 说明了本文关键数据字段的工程含义及建模用途，可为后续复现实验和扩展研究提供变量口径参考。",
    )
    format_body(exp2)


def polish_text(doc: Document) -> None:
    replacements = {
        "本文研究具有一定理论意义和工程意义。理论上，本文将负荷影响因素分析、时序预测模型和多目标优化方法结合起来，形成了较完整的地铁车站通风空调系统容量配置优化方法链。工程上，研究结果可为新建地铁车站通风空调系统设计、既有车站设备改造和运行节能优化提供参考，有助于降低设备容量冗余和全生命周期成本，提高系统运行经济性与调节灵活性": "本文研究具有一定理论意义和工程意义。理论上，本文将负荷影响因素分析、时序预测模型和多目标优化方法结合起来，形成了较完整的地铁车站通风空调系统容量配置优化方法链。工程上，研究结果可为新建地铁车站通风空调系统设计、既有车站设备改造和运行节能优化提供参考，有助于降低设备容量冗余和全生命周期成本，提高系统运行经济性与调节灵活性。",
        "本文采用 Pearson 相关系数法对候选特征与系统总负荷之间的线性相关程度进行分析。Pearson 相关系数计算公式为：": "本文采用 Pearson 相关系数法对候选特征与系统总负荷之间的线性相关程度进行分析，其计算公式为：",
        "模型结构如下：   增加模型框图": "结合图4-2，本文 LSTM 模型结构可概括为以下五个层次：",
    }
    for paragraph in doc.paragraphs:
        txt = text_of(paragraph)
        if txt in replacements:
            set_paragraph_text(paragraph, replacements[txt])
            format_body(paragraph)


def main() -> None:
    if len(sys.argv) != 3:
        print("usage: revise_draft07.py <input.docx> <output.docx>")
        raise SystemExit(2)

    src = Path(sys.argv[1])
    dst = Path(sys.argv[2])
    shutil.copyfile(src, dst)
    doc = Document(str(dst))
    replace_toc(doc)
    polish_text(doc)
    fill_appendix(doc)
    doc.save(str(dst))
    print(dst)


if __name__ == "__main__":
    main()
