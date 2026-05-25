from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "毕业设计论文初稿04_规范与结果修订最终版.docx"
OUT = ROOT / "毕业设计论文初稿05_文献与引用修订稿.docx"


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    if not text:
        return
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)


def insert_paragraph_after(
    paragraph: Paragraph, text: str, style: str | None = None
) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        new_para.style = style
    set_paragraph_text(new_para, text)
    return new_para


def normalize_fonts(paragraph: Paragraph) -> None:
    for run in paragraph.runs:
        if not run.text:
            continue
        run.font.name = "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        if run.font.size is None:
            run.font.size = Pt(12)
        run.font.highlight_color = None


def replace_research_status(doc: Document) -> None:
    start = None
    end = None
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if start is None and "1.3" in text and "研究现状" in text:
            start = i
            continue
        if (
            start is not None
            and i > start
            and paragraph.style.name.startswith("Heading")
            and text.startswith("1.4")
        ):
            end = i
            break
    if start is None or end is None:
        raise RuntimeError(f"Could not locate 1.3 section: start={start}, end={end}")

    new_paragraphs = [
        "地铁车站通风空调系统负荷受室外气象、客流强度、列车运行、站内设备散热、围护结构传热以及新风需求等多因素共同影响。与普通公共建筑相比，地铁车站处于地下空间，太阳辐射直接影响较弱，但客流时变性、列车活塞风、屏蔽门开闭、站内设备长期运行等因素使其负荷呈现明显的时段性和场景差异。已有实测和模拟研究表明，地铁车站公共区实际冷负荷常低于传统设计峰值，设备散热、新风处理和乘客密度是影响空调负荷的重要因素，环控系统容易出现“大马拉小车”和长期低负荷运行问题[1-7]。",
        "在节能控制和运行诊断方面，国内研究主要集中在控制策略优化、控制参数调整、用能诊断和既有车站改造适配等方向。相关综述认为，地铁站通风空调控制系统的节能优化可从控制策略、控制参数和智能控制方法三个层面展开[8]；热舒适与节能策略研究指出，站内环境控制不能只追求低温，还应兼顾乘客热舒适、空气品质和能耗水平[9]。近年的研究进一步关注环控系统能耗诊断、控制方案对比、前馈控制、智能环控系统以及制式改造可行性，为既有车站节能改造和新建车站系统设计提供了工程参考[10-18]。",
        "在负荷预测方面，传统方法包括线性回归、时间序列模型、灰色预测和 BP 神经网络等。这类方法结构相对简单、实现方便，但在处理非线性、多变量和长短期时序依赖时存在一定局限。面向地铁车站空调负荷，已有研究尝试将动态客流模型、客流规律分析和数据挖掘方法引入负荷预测与用能诊断[7,19-20]；国内外相关研究也从设备协同优化、通风空调系统自主控制、能耗影响因素分析、神经网络负荷预测和模糊控制等角度验证了数据驱动方法在地铁车站能耗分析中的适用性[23-26]。LSTM 通过门控结构保存历史状态，适合描述负荷惯性、客流高峰持续和环境参数滞后影响等问题[27]。",
        "在容量配置和综合优化方面，传统工程设计通常依据设计峰值负荷、远期客流和经验安全裕度确定设备容量，方法可靠但容易造成装机容量偏大。多目标优化方法能够同时考虑经济性、冗余率和工程约束，NSGA-II 可输出 Pareto 非劣解集[28]，TOPSIS 可进一步根据成本和冗余等指标进行综合排序[29]，适合用于设备容量组合方案比选。综合来看，现有研究已在地铁车站负荷特性、空调负荷预测、节能控制和运行诊断方面取得进展，但预测结果与设备容量配置之间的衔接仍不够充分，容量冗余率、全生命周期成本和工程离散选型约束仍需进一步结合。因此，本文将负荷特性分析、LSTM 预测和 NSGA-II/TOPSIS 容量优化串联起来，形成面向站台区域通风空调系统的完整方法链。",
    ]

    body_slots = list(range(start + 1, end))
    if len(body_slots) < len(new_paragraphs):
        raise RuntimeError("Not enough paragraph slots in 1.3 section")
    for idx, text in zip(body_slots, new_paragraphs):
        set_paragraph_text(doc.paragraphs[idx], text)
    for idx in body_slots[len(new_paragraphs) :]:
        set_paragraph_text(doc.paragraphs[idx], "")


def clean_markers_and_fix_citations(doc: Document) -> None:
    replacements = {
        "为了识别站台区域总冷负荷的主要驱动因素，本文在完成时间对齐和标准化后，对候选特征与目标负荷进行 Pearson 相关分析[6]。Pearson 相关系数用于衡量两个连续变量之间的线性相关强度，其计算公式为：": "为了识别站台区域总冷负荷的主要驱动因素，本文在完成时间对齐和标准化后，对候选特征与目标负荷进行 Pearson 相关分析。Pearson 相关系数用于衡量两个连续变量之间的线性相关强度，其计算公式为：",
        "K-Means 聚类通过最小化簇内平方误差来划分样本[7]，其目标函数为：": "K-Means 聚类通过最小化簇内平方误差来划分样本，其目标函数为：",
        "NSGA-II 是一种基于非支配排序和拥挤距离机制的多目标遗传算法。其主要流程包括初始化种群、计算目标函数、非支配排序、拥挤距离计算、选择、交叉、变异和种群更新。该算法能够在一次优化过程中获得多个 Pareto 非劣解，从而反映不同优化目标之间的权衡关系。": "NSGA-II 是一种基于非支配排序和拥挤距离机制的多目标遗传算法[28]。其主要流程包括初始化种群、计算目标函数、非支配排序、拥挤距离计算、选择、交叉、变异和种群更新。该算法能够在一次优化过程中获得多个 Pareto 非劣解，从而反映不同优化目标之间的权衡关系。",
        "TOPSIS 方法的基本思想是：最优方案应尽可能接近正理想解，同时尽可能远离负理想解。具体步骤包括指标标准化、确定正负理想解、计算各方案到正负理想解的距离、计算贴近度并排序。对于容量配置方案，评价指标可包括全生命周期成本、容量冗余率、年运行能耗、设备匹配性和调节灵活性等。": "TOPSIS 方法的基本思想是：最优方案应尽可能接近正理想解，同时尽可能远离负理想解[29]。具体步骤包括指标标准化、确定正负理想解、计算各方案到正负理想解的距离、计算贴近度并排序。对于容量配置方案，评价指标可包括全生命周期成本、容量冗余率、年运行能耗、设备匹配性和调节灵活性等。",
        "6  优化结果分析    去掉 加到第五章": "6  优化结果分析",
        "本文以地铁车站站台区域通风空调系统为研究对象，围绕容量配置偏大、系统长期低负荷运行和运行能耗较高等问题，构建了“数据预处理、负荷特性分析、负荷预测建模、容量配置优化、工程评价”的研究流程。基于福州地铁东街口站 2025 年全年 15 分钟粒度运行数据，采用 Pearson 相关分析、负荷分项分解、K-Means 聚类、LSTM 神经网络、NSGA-II 多目标优化和 TOPSIS 综合评价等方法，开展了地铁车站通风空调系统容量配置优化研究。主要结论如下。三段式凝练": "本文以地铁车站站台区域通风空调系统为研究对象，围绕容量配置偏大、系统长期低负荷运行和运行能耗较高等问题，构建了“数据预处理、负荷特性分析、负荷预测建模、容量配置优化、工程评价”的研究流程。基于福州地铁东街口站 2025 年全年 15 分钟粒度运行数据，采用 Pearson 相关分析、负荷分项分解、K-Means 聚类、LSTM 神经网络、NSGA-II 多目标优化和 TOPSIS 综合评价等方法，开展了地铁车站通风空调系统容量配置优化研究。主要结论如下。",
    }

    for paragraph in doc.paragraphs:
        if paragraph.text in replacements:
            set_paragraph_text(paragraph, replacements[paragraph.text])
        elif paragraph.text.strip() == "精简描述":
            set_paragraph_text(paragraph, "")


def insert_references(doc: Document) -> None:
    references = [
        "[1] 刘正宁，王文强，孙春辉．地铁站通风空调系统实测分析及负荷特性研究[J]．铁道科学与工程学报，2024，21(11)：4746-4755．",
        "[2] 齐江浩，赵蕾，王君，等．西安地铁车站环境实测及公共区空调负荷计算分析[J]．铁道科学与工程学报，2016，13(6)：1206-1211．",
        "[3] 马晓明，李晓锋，朱颖心．基于TRNSYS的地铁车站公共区冷负荷预测模型[J]．都市快轨交通，2021，34(2)：130-136．",
        "[4] 隋学敏，王靖宜，郭磊，等．屏蔽门系统地铁车站空调负荷研究现状及展望[J]．铁道标准设计，2019，63(12)：141-149．",
        "[5] 黄莉，苏子怡，李晓锋．地铁车站公共区域空调能耗影响因素的敏感性分析[J]．都市快轨交通，2021，34(6)：125-130．",
        "[6] 邓光蔚，钱程，张丹．上海典型地铁车站环控系统能耗特点分析及运行诊断[J]．暖通空调，2021，51(10)：83-86．",
        "[7] 刘佳慧，龙静，潘志刚，等．基于数据挖掘技术的地铁站环控系统用能诊断[J]．制冷学报，2018，39(3)：1-6．",
        "[8] 曹勇，丁天一，于震．地铁站通风空调控制系统节能优化研究综述[J]．建筑科学，2022，38(4)：213-228．",
        "[9] 曾逸婷，赵蕾．地铁车站环境热舒适与通风空调系统节能策略研究进展[J]．铁道标准设计，2019，63(3)：178-183．",
        "[10] 王丽慧，马嘉楠，尹立元．基于迭代模型的地铁车站空气温度逐年演化特性研究[J]．能源研究与信息，2025，41(3)：144-151．",
        "[11] 张云霞，苏子怡，李晓锋．地铁站环控系统控制方案的节能对比研究[J]．都市快轨交通，2026．",
        "[12] 张乔．关于地铁通风空调系统节能的有效策略探讨[J]．北方建筑，2019，4(6)：43-46．",
        "[13] 吴炜．智能环控系统在城市轨道交通中的应用分析[J]．暖通空调，2020，50(S1)：33-39．",
        "[14] 林菁，付战莹，江洪泽，等．地铁环控大系统基于已知负荷模型的前馈控制策略讨论[J]．都市快轨交通，2021，34(1)：127-131．",
        "[15] 肖宾杰，黄亮亮．夏热冬冷地区地铁车站通风空调系统节能控制系统：以汇金路站为例[J]．城市道桥与防洪，2019(3)：210-213，215．",
        "[16] 王佳琦．轨道交通地下车站环控系统制式改造方案可行性分析[J]．隧道与轨道交通，2021(1)：31-34，61．",
        "[17] 李思玮．基于动态客流量的地铁环控系统节能策略[D]．广州：广州大学，2020．",
        "[18] 马江燕．北方地区冬季地铁车站热环境及其控制策略[D]．西安：西安建筑科技大学，2021．",
        "[19] 苏醒，王磊，田少宸，等．基于动态客流量模型的地铁车站空调负荷预测[J]．同济大学学报(自然科学版)，2022，50(1)．",
        "[20] 李子浩，田向亮，黎忠文，等．基于客流规律的地铁车站客流风险分析[J]．清华大学学报(自然科学版)，2019，59(10)：854-860．",
        "[21] 李明，张骄，崔霆锐，等．北京地铁绿色低碳技术创新研究与应用[J]．机车电传动，2022(3)：29-36．",
        "[22] 岳渤雨．轨道交通地下车站环控系统气流组织模拟与优化[J]．工程技术与管理，2025，9(6)．",
        "[23] 王升，郑懿，常晓敏．基于混合整数非线性规划模型的地铁站通风空调水系统节能措施及潜力研究[J]．暖通空调，2024，54(S1)：302-306．",
        "[24] WANG Yongcai，FENG Haoran，XI Xiangyu．Monitoring and Autonomous Control of Beijing Subway HVAC System for Energy Sustainability[J]．Energy for Sustainable Development，2017，39：1-12．",
        "[25] GUAN Bowen，LIU Xiaohua，ZHANG Tao，et al．Energy Consumption of Subway Stations in China：Data and Influencing Factors[J]．Sustainable Cities and Society，2018，43：451-461．",
        "[26] BI Haiquan，ZHOU Yuanlong，LIU Jin，et al．Load Forecast and Fuzzy Control of the Air-Conditioning Systems at the Subway Stations[J]．Journal of Building Engineering，2022，49：104029．",
        "[27] HOCHREITER S，SCHMIDHUBER J．Long Short-Term Memory[J]．Neural Computation，1997，9(8)：1735-1780．",
        "[28] DEB K，PRATAP A，AGARWAL S，et al．A Fast and Elitist Multiobjective Genetic Algorithm：NSGA-II[J]．IEEE Transactions on Evolutionary Computation，2002，6(2)：182-197．",
        "[29] HWANG C L，YOON K．Multiple Attribute Decision Making：Methods and Applications[M]．Berlin：Springer，1981．",
    ]

    ref_heading = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "参考文献":
            ref_heading = paragraph
            break
    if ref_heading is None:
        raise RuntimeError("Could not find references heading")

    current = ref_heading
    for item in references:
        current = insert_paragraph_after(current, item, style="Normal")
        current.paragraph_format.first_line_indent = None
        current.paragraph_format.left_indent = None
        current.paragraph_format.space_after = Pt(0)


def main() -> None:
    doc = Document(SRC)
    replace_research_status(doc)
    clean_markers_and_fix_citations(doc)
    insert_references(doc)
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            normalize_fonts(paragraph)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
