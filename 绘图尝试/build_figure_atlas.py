from __future__ import annotations

import math
import shutil
from dataclasses import dataclass
from pathlib import Path

import numpy as np
import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "绘图尝试"
FIG_DIR = OUT_DIR / "figures"
DOCX_PATH = OUT_DIR / "地铁环控系统论文绘图尝试_全图集.docx"


FONT_REGULAR = [
    Path(r"C:\Windows\Fonts\NotoSansSC-VF.ttf"),
    Path(r"C:\Windows\Fonts\msyh.ttc"),
    Path(r"C:\Windows\Fonts\simhei.ttf"),
    Path(r"C:\Windows\Fonts\simsun.ttc"),
]
FONT_SERIF = [
    Path(r"C:\Windows\Fonts\NotoSerifSC-VF.ttf"),
    Path(r"C:\Windows\Fonts\simsun.ttc"),
    Path(r"C:\Windows\Fonts\simkai.ttf"),
]

PALETTE = {
    "ink": "#1F2937",
    "muted": "#6B7280",
    "grid": "#D7DEE8",
    "light": "#F6F8FB",
    "blue": "#2563EB",
    "cyan": "#0891B2",
    "green": "#059669",
    "orange": "#D97706",
    "red": "#DC2626",
    "purple": "#7C3AED",
    "teal": "#0F766E",
    "slate": "#475569",
    "gold": "#B7791F",
}
SERIES_COLORS = [
    PALETTE["blue"],
    PALETTE["orange"],
    PALETTE["green"],
    PALETTE["red"],
    PALETTE["purple"],
    PALETTE["cyan"],
    PALETTE["teal"],
    PALETTE["gold"],
]


@dataclass
class FigureItem:
    no: int
    title: str
    path: Path
    section: str
    note: str


figures: list[FigureItem] = []


def ensure_dirs() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)


def pick_font(paths: list[Path], size: int) -> ImageFont.FreeTypeFont:
    for path in paths:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def font(size: int, serif: bool = False) -> ImageFont.FreeTypeFont:
    return pick_font(FONT_SERIF if serif else FONT_REGULAR, size)


def hex_to_rgb(hex_color: str) -> tuple[int, int, int]:
    hex_color = hex_color.lstrip("#")
    return tuple(int(hex_color[i : i + 2], 16) for i in (0, 2, 4))


def blend(c1: str, c2: str, t: float) -> tuple[int, int, int]:
    a = np.array(hex_to_rgb(c1), dtype=float)
    b = np.array(hex_to_rgb(c2), dtype=float)
    return tuple(np.clip(a * (1 - t) + b * t, 0, 255).astype(int))


def text_size(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.ImageFont) -> tuple[int, int]:
    box = draw.textbbox((0, 0), text, font=fnt)
    return box[2] - box[0], box[3] - box[1]


def wrap_text(draw: ImageDraw.ImageDraw, text: str, fnt: ImageFont.ImageFont, max_width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for ch in str(text):
        if ch == "\n":
            if current:
                lines.append(current)
            current = ""
            continue
        trial = current + ch
        if text_size(draw, trial, fnt)[0] <= max_width or not current:
            current = trial
        else:
            lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines


def draw_wrapped(
    draw: ImageDraw.ImageDraw,
    text: str,
    xy: tuple[int, int],
    fnt: ImageFont.ImageFont,
    fill: str | tuple[int, int, int] = PALETTE["ink"],
    max_width: int = 600,
    line_gap: int = 8,
    align: str = "left",
) -> int:
    x, y = xy
    lines = wrap_text(draw, text, fnt, max_width)
    line_height = text_size(draw, "测试", fnt)[1] + line_gap
    for line in lines:
        w, _ = text_size(draw, line, fnt)
        lx = x
        if align == "center":
            lx = x + (max_width - w) // 2
        elif align == "right":
            lx = x + max_width - w
        draw.text((lx, y), line, font=fnt, fill=fill)
        y += line_height
    return y


def new_canvas(title: str, subtitle: str | None = None, size: tuple[int, int] = (1800, 1120)) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", size, "white")
    draw = ImageDraw.Draw(img)
    draw.rectangle((0, 0, size[0], 88), fill="#EEF3FB")
    draw.text((80, 26), title, font=font(36), fill=PALETTE["ink"])
    if subtitle:
        draw.text((80, 74), subtitle, font=font(18), fill=PALETTE["muted"])
    draw.line((60, 104, size[0] - 60, 104), fill=PALETTE["grid"], width=2)
    return img, draw


def save_figure(img: Image.Image, filename: str, title: str, section: str, note: str) -> Path:
    path = FIG_DIR / filename
    img.save(path, quality=95)
    figures.append(FigureItem(len(figures) + 1, title, path, section, note))
    return path


def read_csv_any(path: Path) -> pd.DataFrame:
    for enc in ("utf-8-sig", "utf-8", "gb18030"):
        try:
            return pd.read_csv(path, encoding=enc)
        except UnicodeDecodeError:
            continue
    return pd.read_csv(path)


def load_data() -> dict[str, pd.DataFrame]:
    data = read_csv_any(ROOT / "data" / "fuzhou_metro_dongjiekou_2025.csv")
    data["timestamp"] = pd.to_datetime(data["timestamp"])
    tables: dict[str, pd.DataFrame] = {"raw": data}
    for path in (ROOT / "output" / "tables").glob("*.csv"):
        tables[path.stem] = read_csv_any(path)
    return tables


def clean_series(values: pd.Series) -> np.ndarray:
    arr = pd.to_numeric(values, errors="coerce").to_numpy(dtype=float)
    return arr[np.isfinite(arr)]


def scale_values(values: np.ndarray, vmin: float, vmax: float, pmin: int, pmax: int, invert: bool = False) -> np.ndarray:
    if not np.isfinite(vmin) or not np.isfinite(vmax) or abs(vmax - vmin) < 1e-12:
        return np.full_like(values, (pmin + pmax) / 2, dtype=float)
    t = (values - vmin) / (vmax - vmin)
    t = np.clip(t, 0, 1)
    if invert:
        t = 1 - t
    return pmin + t * (pmax - pmin)


def draw_axes(
    draw: ImageDraw.ImageDraw,
    area: tuple[int, int, int, int],
    y_min: float,
    y_max: float,
    x_labels: list[tuple[float, str]] | None = None,
    y_label: str = "",
) -> None:
    x0, y0, x1, y1 = area
    draw.line((x0, y1, x1, y1), fill=PALETTE["ink"], width=3)
    draw.line((x0, y0, x0, y1), fill=PALETTE["ink"], width=3)
    tick_font = font(18)
    for i in range(6):
        yy = int(y1 - (y1 - y0) * i / 5)
        val = y_min + (y_max - y_min) * i / 5
        draw.line((x0, yy, x1, yy), fill=PALETTE["grid"], width=1)
        label = f"{val:.0f}" if abs(y_max - y_min) > 10 else f"{val:.2f}"
        tw, th = text_size(draw, label, tick_font)
        draw.text((x0 - tw - 12, yy - th // 2), label, font=tick_font, fill=PALETTE["muted"])
    if x_labels:
        for frac, label in x_labels:
            xx = int(x0 + frac * (x1 - x0))
            draw.line((xx, y1, xx, y1 + 8), fill=PALETTE["ink"], width=2)
            tw, _ = text_size(draw, label, tick_font)
            draw.text((xx - tw // 2, y1 + 18), label, font=tick_font, fill=PALETTE["muted"])
    if y_label:
        draw.text((x0 - 80, y0 - 42), y_label, font=font(20), fill=PALETTE["muted"])


def line_chart(
    filename: str,
    title: str,
    section: str,
    note: str,
    series: list[tuple[str, np.ndarray, str]],
    x_labels: list[tuple[float, str]],
    y_label: str,
    subtitle: str | None = None,
) -> None:
    img, draw = new_canvas(title, subtitle)
    area = (150, 170, 1660, 900)
    all_vals = np.concatenate([s[1][np.isfinite(s[1])] for s in series if len(s[1])])
    y_min = float(np.nanmin(all_vals))
    y_max = float(np.nanmax(all_vals))
    pad = max((y_max - y_min) * 0.08, 1)
    y_min -= pad
    y_max += pad
    draw_axes(draw, area, y_min, y_max, x_labels, y_label)
    x0, y0, x1, y1 = area
    for label, vals, color in series:
        vals = np.asarray(vals, dtype=float)
        idx = np.arange(len(vals))
        ok = np.isfinite(vals)
        if ok.sum() < 2:
            continue
        xs = scale_values(idx[ok], idx[ok].min(), idx[ok].max(), x0, x1)
        ys = scale_values(vals[ok], y_min, y_max, y0, y1, invert=True)
        points = list(zip(xs.astype(int), ys.astype(int)))
        if len(points) > 1400:
            step = max(1, len(points) // 1400)
            points = points[::step]
        draw.line(points, fill=color, width=4, joint="curve")
    legend_x = 160
    legend_y = 950
    for label, _, color in series:
        draw.rectangle((legend_x, legend_y, legend_x + 32, legend_y + 16), fill=color)
        draw.text((legend_x + 44, legend_y - 5), label, font=font(20), fill=PALETTE["ink"])
        legend_x += text_size(draw, label, font(20))[0] + 120
    save_figure(img, filename, title, section, note)


def horizontal_bar(
    filename: str,
    title: str,
    section: str,
    note: str,
    labels: list[str],
    values: list[float],
    value_fmt: str = "{:.2f}",
    subtitle: str | None = None,
    color: str = PALETTE["blue"],
) -> None:
    img, draw = new_canvas(title, subtitle)
    x0, y0, x1, y1 = (420, 160, 1610, 930)
    vals = np.asarray(values, dtype=float)
    vmax = max(float(np.nanmax(vals)), 1e-9)
    row_h = min(62, max(30, (y1 - y0) // max(1, len(labels))))
    lab_font = font(22)
    val_font = font(20)
    for i, (lab, val) in enumerate(zip(labels, vals)):
        cy = y0 + i * row_h + row_h // 2
        draw.text((70, cy - 16), lab, font=lab_font, fill=PALETTE["ink"])
        bar_w = int((x1 - x0) * max(val, 0) / vmax)
        draw.rounded_rectangle((x0, cy - 18, x0 + bar_w, cy + 18), radius=6, fill=color)
        draw.text((x0 + bar_w + 14, cy - 13), value_fmt.format(val), font=val_font, fill=PALETTE["muted"])
    save_figure(img, filename, title, section, note)


def vertical_bar(
    filename: str,
    title: str,
    section: str,
    note: str,
    labels: list[str],
    values: list[float],
    y_label: str,
    subtitle: str | None = None,
    color: str = PALETTE["cyan"],
) -> None:
    img, draw = new_canvas(title, subtitle)
    area = (150, 170, 1660, 850)
    vals = np.asarray(values, dtype=float)
    y_min = min(0.0, float(np.nanmin(vals)))
    y_max = float(np.nanmax(vals)) * 1.15 if len(vals) else 1.0
    draw_axes(draw, area, y_min, y_max, None, y_label)
    x0, y0, x1, y1 = area
    gap = 24
    bw = max(18, int((x1 - x0 - gap * (len(vals) + 1)) / max(1, len(vals))))
    for i, val in enumerate(vals):
        bx = x0 + gap + i * (bw + gap)
        by = int(scale_values(np.array([val]), y_min, y_max, y0, y1, invert=True)[0])
        draw.rounded_rectangle((bx, by, bx + bw, y1), radius=8, fill=color)
        lab_lines = wrap_text(draw, labels[i], font(18), bw + 30)
        ly = y1 + 20
        for line in lab_lines[:2]:
            tw, _ = text_size(draw, line, font(18))
            draw.text((bx + bw / 2 - tw / 2, ly), line, font=font(18), fill=PALETTE["muted"])
            ly += 24
    save_figure(img, filename, title, section, note)


def grouped_bar(
    filename: str,
    title: str,
    section: str,
    note: str,
    groups: list[str],
    series: list[tuple[str, list[float], str]],
    y_label: str,
    subtitle: str | None = None,
) -> None:
    img, draw = new_canvas(title, subtitle)
    area = (150, 170, 1660, 830)
    all_vals = np.concatenate([np.asarray(vals, dtype=float) for _, vals, _ in series])
    y_min = min(0.0, float(np.nanmin(all_vals)))
    y_max = float(np.nanmax(all_vals)) * 1.18 if len(all_vals) else 1.0
    draw_axes(draw, area, y_min, y_max, None, y_label)
    x0, y0, x1, y1 = area
    group_w = (x1 - x0) / max(1, len(groups))
    bar_w = min(55, group_w / (len(series) + 1.4))
    for gi, group in enumerate(groups):
        center = x0 + group_w * (gi + 0.5)
        for si, (_, vals, color) in enumerate(series):
            val = vals[gi]
            bx0 = center - (len(series) * bar_w) / 2 + si * bar_w
            bx1 = bx0 + bar_w * 0.82
            by = int(scale_values(np.array([val]), y_min, y_max, y0, y1, invert=True)[0])
            draw.rounded_rectangle((bx0, by, bx1, y1), radius=6, fill=color)
        lines = wrap_text(draw, group, font(18), int(group_w * 0.9))
        ly = y1 + 18
        for line in lines[:2]:
            tw, _ = text_size(draw, line, font(18))
            draw.text((center - tw / 2, ly), line, font=font(18), fill=PALETTE["muted"])
            ly += 24
    lx, ly = 160, 935
    for name, _, color in series:
        draw.rectangle((lx, ly, lx + 32, ly + 18), fill=color)
        draw.text((lx + 44, ly - 4), name, font=font(20), fill=PALETTE["ink"])
        lx += text_size(draw, name, font(20))[0] + 110
    save_figure(img, filename, title, section, note)


def heatmap(
    filename: str,
    title: str,
    section: str,
    note: str,
    matrix: np.ndarray,
    row_labels: list[str],
    col_labels: list[str],
    subtitle: str | None = None,
    low: str = "#F7FBFF",
    high: str = "#1D4ED8",
) -> None:
    img, draw = new_canvas(title, subtitle)
    x0, y0, x1, y1 = (170, 160, 1640, 870)
    rows, cols = matrix.shape
    cell_w = (x1 - x0) / cols
    cell_h = (y1 - y0) / rows
    finite = matrix[np.isfinite(matrix)]
    vmin = float(np.nanmin(finite)) if finite.size else 0
    vmax = float(np.nanmax(finite)) if finite.size else 1
    for r in range(rows):
        for c in range(cols):
            val = matrix[r, c]
            t = 0 if vmax == vmin or not np.isfinite(val) else (val - vmin) / (vmax - vmin)
            color = blend(low, high, float(np.clip(t, 0, 1)))
            draw.rectangle((x0 + c * cell_w, y0 + r * cell_h, x0 + (c + 1) * cell_w, y0 + (r + 1) * cell_h), fill=color)
    for r, lab in enumerate(row_labels):
        draw.text((80, y0 + r * cell_h + cell_h / 2 - 12), lab, font=font(18), fill=PALETTE["ink"])
    step = max(1, len(col_labels) // 12)
    for c in range(0, len(col_labels), step):
        lab = col_labels[c]
        tw, _ = text_size(draw, lab, font(16))
        draw.text((x0 + c * cell_w + cell_w / 2 - tw / 2, y1 + 18), lab, font=font(16), fill=PALETTE["muted"])
    draw.rectangle((x0, y0, x1, y1), outline=PALETTE["ink"], width=2)
    draw.text((x0, y1 + 60), f"颜色越深表示数值越高；范围：{vmin:.1f} 至 {vmax:.1f}", font=font(20), fill=PALETTE["muted"])
    save_figure(img, filename, title, section, note)


def scatter_plot(
    filename: str,
    title: str,
    section: str,
    note: str,
    x: np.ndarray,
    y: np.ndarray,
    x_label: str,
    y_label: str,
    subtitle: str | None = None,
) -> None:
    img, draw = new_canvas(title, subtitle)
    area = (160, 170, 1640, 860)
    ok = np.isfinite(x) & np.isfinite(y)
    x = x[ok]
    y = y[ok]
    if len(x) > 3500:
        idx = np.linspace(0, len(x) - 1, 3500).astype(int)
        x = x[idx]
        y = y[idx]
    xmin, xmax = float(np.min(x)), float(np.max(x))
    ymin, ymax = float(np.min(y)), float(np.max(y))
    xpad = max((xmax - xmin) * 0.06, 1)
    ypad = max((ymax - ymin) * 0.06, 1)
    xmin -= xpad
    xmax += xpad
    ymin -= ypad
    ymax += ypad
    draw_axes(draw, area, ymin, ymax, None, y_label)
    x0, y0, x1, y1 = area
    xs = scale_values(x, xmin, xmax, x0, x1)
    ys = scale_values(y, ymin, ymax, y0, y1, invert=True)
    overlay = Image.new("RGBA", img.size, (255, 255, 255, 0))
    od = ImageDraw.Draw(overlay)
    for xx, yy in zip(xs, ys):
        od.ellipse((xx - 3, yy - 3, xx + 3, yy + 3), fill=(37, 99, 235, 70))
    img.alpha_composite(overlay) if img.mode == "RGBA" else None
    img = Image.alpha_composite(img.convert("RGBA"), overlay).convert("RGB")
    draw = ImageDraw.Draw(img)
    draw.text((x0, y1 + 50), x_label, font=font(22), fill=PALETTE["muted"])
    draw.text((x0, y1 + 85), f"x范围：{xmin + xpad:.1f} 至 {xmax - xpad:.1f}", font=font(18), fill=PALETTE["muted"])
    save_figure(img, filename, title, section, note)


def donut_chart(
    filename: str,
    title: str,
    section: str,
    note: str,
    labels: list[str],
    values: list[float],
    subtitle: str | None = None,
) -> None:
    img, draw = new_canvas(title, subtitle)
    cx, cy, r = 580, 535, 330
    total = sum(values)
    start = -90
    for i, (lab, val) in enumerate(zip(labels, values)):
        extent = 360 * val / total if total else 0
        draw.pieslice((cx - r, cy - r, cx + r, cy + r), start, start + extent, fill=SERIES_COLORS[i % len(SERIES_COLORS)])
        start += extent
    draw.ellipse((cx - 155, cy - 155, cx + 155, cy + 155), fill="white")
    draw.text((cx - 95, cy - 26), "平均占比", font=font(32), fill=PALETTE["ink"])
    lx, ly = 1050, 260
    for i, (lab, val) in enumerate(zip(labels, values)):
        color = SERIES_COLORS[i % len(SERIES_COLORS)]
        draw.rectangle((lx, ly, lx + 34, ly + 24), fill=color)
        draw.text((lx + 52, ly - 4), f"{lab}  {val / total * 100:.1f}%", font=font(26), fill=PALETTE["ink"])
        ly += 70
    save_figure(img, filename, title, section, note)


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color: str = PALETTE["slate"], width: int = 4) -> None:
    draw.line((start, end), fill=color, width=width)
    sx, sy = start
    ex, ey = end
    ang = math.atan2(ey - sy, ex - sx)
    size = 16
    pts = [
        (ex, ey),
        (ex - size * math.cos(ang - math.pi / 6), ey - size * math.sin(ang - math.pi / 6)),
        (ex - size * math.cos(ang + math.pi / 6), ey - size * math.sin(ang + math.pi / 6)),
    ]
    draw.polygon(pts, fill=color)


def box(draw: ImageDraw.ImageDraw, rect: tuple[int, int, int, int], title: str, body: str = "", fill: str = "#F8FAFC", outline: str = PALETTE["blue"]) -> None:
    draw.rounded_rectangle(rect, radius=18, fill=fill, outline=outline, width=3)
    x0, y0, x1, y1 = rect
    draw_wrapped(draw, title, (x0 + 22, y0 + 18), font(25), PALETTE["ink"], x1 - x0 - 44)
    if body:
        draw_wrapped(draw, body, (x0 + 22, y0 + 64), font(18), PALETTE["muted"], x1 - x0 - 44, line_gap=6)


def flow_figure(filename: str, title: str, section: str, note: str, nodes: list[dict], arrows: list[tuple[int, int]], subtitle: str | None = None) -> None:
    img, draw = new_canvas(title, subtitle)
    for n in nodes:
        box(draw, n["rect"], n["title"], n.get("body", ""), n.get("fill", "#F8FAFC"), n.get("outline", PALETTE["blue"]))
    for a, b in arrows:
        r1 = nodes[a]["rect"]
        r2 = nodes[b]["rect"]
        start = (r1[2], (r1[1] + r1[3]) // 2)
        end = (r2[0], (r2[1] + r2[3]) // 2)
        if abs(start[0] - end[0]) < 30:
            start = ((r1[0] + r1[2]) // 2, r1[3])
            end = ((r2[0] + r2[2]) // 2, r2[1])
        draw_arrow(draw, start, end)
    save_figure(img, filename, title, section, note)


def table_figure(
    filename: str,
    title: str,
    section: str,
    note: str,
    headers: list[str],
    rows: list[list[str]],
    subtitle: str | None = None,
) -> None:
    img, draw = new_canvas(title, subtitle)
    x0, y0, x1 = 90, 170, 1710
    col_widths = [int((x1 - x0) * w) for w in np.ones(len(headers)) / len(headers)]
    row_h = 72
    header_h = 58
    draw.rounded_rectangle((x0, y0, x1, y0 + header_h + row_h * len(rows)), radius=14, fill="white", outline=PALETTE["grid"], width=2)
    draw.rectangle((x0, y0, x1, y0 + header_h), fill="#E8F0FE")
    cx = x0
    for i, h in enumerate(headers):
        draw_wrapped(draw, h, (cx + 14, y0 + 14), font(18), PALETTE["ink"], col_widths[i] - 28, line_gap=4)
        cx += col_widths[i]
        draw.line((cx, y0, cx, y0 + header_h + row_h * len(rows)), fill=PALETTE["grid"], width=1)
    for ri, row in enumerate(rows):
        ry = y0 + header_h + ri * row_h
        if ri % 2 == 1:
            draw.rectangle((x0, ry, x1, ry + row_h), fill="#FBFCFE")
        draw.line((x0, ry, x1, ry), fill=PALETTE["grid"], width=1)
        cx = x0
        for ci, txt in enumerate(row):
            draw_wrapped(draw, str(txt), (cx + 14, ry + 12), font(17), PALETTE["ink"], col_widths[ci] - 28, line_gap=3)
            cx += col_widths[ci]
    save_figure(img, filename, title, section, note)


def copy_existing(src_name: str, filename: str, title: str, section: str, note: str) -> None:
    src = ROOT / "output" / "figures" / src_name
    if src.exists():
        dst = FIG_DIR / filename
        shutil.copy2(src, dst)
        figures.append(FigureItem(len(figures) + 1, title, dst, section, note))


def generate_flow_diagrams() -> None:
    section = "一、研究流程与系统架构"
    flow_figure(
        "01_技术路线总图.png",
        "技术路线总图",
        section,
        "展示论文从数据获取、特征分析、负荷预测到容量优化和工程评价的完整研究路径。",
        [
            {"rect": (80, 190, 330, 310), "title": "数据生成/采集", "body": "客流、气象、站内环境、设备负荷"},
            {"rect": (420, 190, 670, 310), "title": "预处理", "body": "缺失填补、时间特征、滞后负荷"},
            {"rect": (760, 190, 1010, 310), "title": "特性分析", "body": "Pearson、负荷分项、典型日聚类"},
            {"rect": (1100, 190, 1350, 310), "title": "负荷预测", "body": "LSTM 与 BP 对比、消融实验"},
            {"rect": (1440, 190, 1690, 310), "title": "容量优化", "body": "NSGA-II、TOPSIS、工程校核"},
            {"rect": (760, 530, 1010, 660), "title": "论文结论", "body": "成本、冗余率、能耗与工程边界"},
        ],
        [(0, 1), (1, 2), (2, 3), (3, 4), (2, 5), (4, 5)],
    )
    flow_figure(
        "02_系统架构与数据流图.png",
        "系统架构与数据流图",
        section,
        "说明 MATLAB 主流程、Python 数据生成器、配置中心、输出图表和测试之间的关系。",
        [
            {"rect": (90, 180, 360, 300), "title": "Python 数据生成", "body": "generate_station_data -> CSV"},
            {"rect": (90, 410, 360, 530), "title": "config.m", "body": "路径、特征、模型参数、工程约束"},
            {"rect": (500, 180, 770, 300), "title": "step1", "body": "清洗与特征工程"},
            {"rect": (840, 180, 1110, 300), "title": "step2", "body": "相关性与聚类"},
            {"rect": (1180, 180, 1450, 300), "title": "step3", "body": "预测与需求情景"},
            {"rect": (1180, 470, 1450, 590), "title": "step4", "body": "容量优化与校核"},
            {"rect": (500, 470, 770, 590), "title": "output", "body": "tables / figures / models"},
            {"rect": (840, 470, 1110, 590), "title": "tests", "body": "pytest + MATLAB tests"},
        ],
        [(0, 2), (2, 3), (3, 4), (4, 5), (1, 2), (1, 5), (5, 6), (6, 7)],
    )
    flow_figure(
        "03_数据生成机制图.png",
        "合成数据生成机制图",
        section,
        "概括客流峰型、季节项、环境项和设备控制项如何共同生成总冷负荷。",
        [
            {"rect": (80, 170, 360, 290), "title": "日类型", "body": "工作日高峰/中等、周末单峰、低流量日"},
            {"rect": (80, 370, 360, 490), "title": "时间峰型", "body": "早高峰、晚高峰、午间、夜间"},
            {"rect": (470, 270, 750, 390), "title": "客流特征", "body": "进站、出站、站台人数、CO2"},
            {"rect": (860, 170, 1140, 290), "title": "气象特征", "body": "室外温湿度、太阳辐射、季节修正"},
            {"rect": (860, 370, 1140, 490), "title": "站内环境", "body": "站台温湿度、热惯性"},
            {"rect": (1260, 270, 1600, 390), "title": "冷负荷输出", "body": "人员、新风、围护结构、设备、总冷负荷"},
        ],
        [(0, 2), (1, 2), (2, 5), (3, 5), (4, 5)],
    )
    flow_figure(
        "04_预测到容量优化映射图.png",
        "预测结果到容量优化映射图",
        section,
        "解释 LSTM 完整预测曲线如何转换为子系统 P99 需求、误差裕量和设备容量约束。",
        [
            {"rect": (80, 190, 370, 310), "title": "LSTM预测曲线", "body": "totalCoolingLoadKw"},
            {"rect": (470, 190, 760, 310), "title": "分位数情景", "body": "P50 typical / P95 peak / P99 extreme"},
            {"rect": (860, 190, 1150, 310), "title": "子系统需求", "body": "冷机、风机、水泵、AHU"},
            {"rect": (1250, 190, 1600, 310), "title": "误差裕量", "body": "按 LSTM RMSE 修正 P99 需求"},
            {"rect": (470, 500, 760, 620), "title": "工程约束", "body": "安全系数、最小PLR、匹配比例"},
            {"rect": (860, 500, 1150, 620), "title": "候选设备组合", "body": "容量规格索引 + 台数"},
            {"rect": (1250, 500, 1600, 620), "title": "推荐方案", "body": "Pareto + TOPSIS"},
        ],
        [(0, 1), (1, 2), (2, 3), (3, 5), (4, 5), (5, 6)],
    )


def generate_data_figures(tables: dict[str, pd.DataFrame]) -> None:
    section = "二、数据质量与负荷特性"
    df = tables["raw"].copy()
    df["date"] = df["timestamp"].dt.date
    df["month"] = df["timestamp"].dt.month
    df["hour_float"] = df["timestamp"].dt.hour + df["timestamp"].dt.minute / 60
    df["week"] = df["timestamp"].dt.isocalendar().week.astype(int)
    load = pd.to_numeric(df["total_cooling_load_kw"], errors="coerce")

    daily = df.groupby(pd.Grouper(key="timestamp", freq="D"))["total_cooling_load_kw"].mean()
    line_chart(
        "05_全年日均冷负荷时序图.png",
        "全年日均冷负荷时序图",
        section,
        "展示全年冷负荷随季节变化的长期趋势，夏季高负荷区间明显高于过渡季和冬季。",
        [("日均总冷负荷", daily.to_numpy(dtype=float), PALETTE["blue"])],
        [(0, "1月"), (0.25, "4月"), (0.5, "7月"), (0.75, "10月"), (1, "12月")],
        "kW",
        "按日平均聚合，减少15分钟波动对全年趋势判断的干扰。",
    )

    week_df = df[(df["timestamp"] >= "2025-07-07") & (df["timestamp"] < "2025-07-14")]
    line_chart(
        "06_夏季典型周负荷时序图.png",
        "夏季典型周负荷时序图",
        section,
        "展示论文可引用的一周负荷时序，能体现工作日峰值、周末单峰和夜间低负荷。",
        [("15分钟总冷负荷", pd.to_numeric(week_df["total_cooling_load_kw"], errors="coerce").to_numpy(), PALETTE["blue"])],
        [(0, "周一"), (1 / 6, "周二"), (2 / 6, "周三"), (3 / 6, "周四"), (4 / 6, "周五"), (5 / 6, "周六"), (1, "周日")],
        "kW",
        "选取 2025-07-07 至 2025-07-13。",
    )

    norm_load = pd.to_numeric(week_df["total_cooling_load_kw"], errors="coerce")
    norm_flow = pd.to_numeric(week_df["entry_flow"], errors="coerce") + pd.to_numeric(week_df["exit_flow"], errors="coerce")
    norm_load = (norm_load - norm_load.min()) / (norm_load.max() - norm_load.min())
    norm_flow = (norm_flow - norm_flow.min()) / (norm_flow.max() - norm_flow.min())
    line_chart(
        "07_典型周负荷与客流归一化对比图.png",
        "典型周负荷与客流归一化对比图",
        section,
        "将冷负荷与总客流归一化后对比，展示客流高峰对负荷峰值的同步影响。",
        [
            ("总冷负荷归一化", norm_load.to_numpy(dtype=float), PALETTE["blue"]),
            ("进出站客流归一化", norm_flow.to_numpy(dtype=float), PALETTE["orange"]),
        ],
        [(0, "周一"), (1 / 3, "周三"), (2 / 3, "周五"), (1, "周日")],
        "归一化值",
    )

    monthly = df.groupby("month")["total_cooling_load_kw"].mean()
    vertical_bar(
        "08_月平均冷负荷柱状图.png",
        "月平均冷负荷柱状图",
        section,
        "比较各月平均冷负荷，适合论文中说明季节性负荷差异。",
        [f"{i}月" for i in monthly.index],
        monthly.to_list(),
        "kW",
        color=PALETTE["cyan"],
    )

    month_hour = df.pivot_table(index="month", columns=df["timestamp"].dt.hour, values="total_cooling_load_kw", aggfunc="mean")
    heatmap(
        "09_月小时平均冷负荷热力图.png",
        "月-小时平均冷负荷热力图",
        section,
        "横向展示日内规律，纵向展示季节规律；颜色较深区域为容量设计需要重点关注的高负荷时段。",
        month_hour.to_numpy(dtype=float),
        [f"{i}月" for i in month_hour.index],
        [str(i) for i in month_hour.columns],
        "行表示月份，列表示一天24小时。",
    )

    profile_weekday = df[~df["is_weekend"].astype(bool)].groupby("hour_float")["total_cooling_load_kw"].mean()
    profile_weekend = df[df["is_weekend"].astype(bool)].groupby("hour_float")["total_cooling_load_kw"].mean()
    line_chart(
        "10_工作日与周末日内负荷对比图.png",
        "工作日与周末日内负荷对比图",
        section,
        "对比工作日和周末的日内负荷形态，支撑典型日模式分析。",
        [
            ("工作日平均", profile_weekday.to_numpy(dtype=float), PALETTE["blue"]),
            ("周末平均", profile_weekend.to_numpy(dtype=float), PALETTE["orange"]),
        ],
        [(0, "0:00"), (0.25, "6:00"), (0.5, "12:00"), (0.75, "18:00"), (1, "24:00")],
        "kW",
    )

    day_type_order = ["weekday_high", "weekday_medium", "weekend_single", "low_flow"]
    day_type_names = ["工作日高峰", "工作日中等", "周末单峰", "低流量日"]
    series = []
    for i, dt in enumerate(day_type_order):
        prof = df[df["day_type"] == dt].groupby("hour_float")["total_cooling_load_kw"].mean()
        series.append((day_type_names[i], prof.to_numpy(dtype=float), SERIES_COLORS[i]))
    line_chart(
        "11_典型日负荷曲线对比图.png",
        "典型日负荷曲线对比图",
        section,
        "按数据生成时的日类型对比四类典型日负荷曲线，是论文中解释日负荷模式的核心图。",
        series,
        [(0, "0:00"), (0.25, "6:00"), (0.5, "12:00"), (0.75, "18:00"), (1, "24:00")],
        "kW",
    )

    sorted_load = np.sort(clean_series(df["total_cooling_load_kw"]))[::-1]
    line_chart(
        "12_负荷持续曲线图.png",
        "负荷持续曲线图",
        section,
        "负荷持续曲线用于观察高负荷持续时间和设备长期部分负荷运行特征。",
        [("总冷负荷持续曲线", sorted_load, PALETTE["green"])],
        [(0, "最高负荷"), (0.5, "50%时长"), (1, "最低负荷")],
        "kW",
    )

    missing = df[["entry_flow", "outdoor_temp", "platform_temp", "co2", "total_cooling_load_kw"]].isna().sum()
    horizontal_bar(
        "13_原始数据缺失值统计图.png",
        "原始数据缺失值统计图",
        section,
        "展示合成数据中用于验证预处理流程的缺失值分布。",
        ["进站客流", "室外温度", "站台温度", "CO2", "总冷负荷"],
        missing.to_list(),
        "{:.0f}",
        color=PALETTE["red"],
    )

    ratio = tables["step2_load_component_ratio"]
    labels = ratio.iloc[:, 0].astype(str).to_list()
    vals = pd.to_numeric(ratio.iloc[:, 2], errors="coerce").fillna(0).to_list()
    donut_chart(
        "14_负荷分项占比环形图.png",
        "负荷分项占比环形图",
        section,
        "展示总冷负荷中人员、新风、围护结构和设备负荷的平均占比。",
        labels,
        vals,
    )

    comp_cols = ["people_load_kw", "fresh_air_load_kw", "envelope_load_kw", "equipment_load_kw"]
    comp_names = ["人员负荷", "新风负荷", "围护结构负荷", "设备负荷"]
    monthly_comp = df.groupby("month")[comp_cols].mean()
    img, draw = new_canvas("月平均分项负荷堆叠图", "展示不同月份总冷负荷组成变化。")
    area = (150, 170, 1660, 850)
    totals = monthly_comp.sum(axis=1).to_numpy(dtype=float)
    draw_axes(draw, area, 0, float(totals.max()) * 1.15, None, "kW")
    x0, y0, x1, y1 = area
    gap = 25
    bw = (x1 - x0 - gap * 13) / 12
    for i, (_, row) in enumerate(monthly_comp.iterrows()):
        bx0 = x0 + gap + i * (bw + gap)
        bottom = y1
        for ci, col in enumerate(comp_cols):
            val = float(row[col])
            top = int(scale_values(np.array([sum(row[comp_cols[: ci + 1]])]), 0, totals.max() * 1.15, y0, y1, invert=True)[0])
            draw.rectangle((bx0, top, bx0 + bw, bottom), fill=SERIES_COLORS[ci])
            bottom = top
        draw.text((bx0 + bw / 2 - 14, y1 + 18), f"{i+1}月", font=font(16), fill=PALETTE["muted"])
    lx, ly = 170, 940
    for ci, name in enumerate(comp_names):
        draw.rectangle((lx, ly, lx + 28, ly + 18), fill=SERIES_COLORS[ci])
        draw.text((lx + 40, ly - 4), name, font=font(19), fill=PALETTE["ink"])
        lx += text_size(draw, name, font(19))[0] + 92
    save_figure(img, "15_月平均分项负荷堆叠图.png", "月平均分项负荷堆叠图", section, "用堆叠方式展示各月人员、新风、围护结构和设备负荷贡献。")

    scatter_plot(
        "16_客流与总冷负荷散点图.png",
        "客流与总冷负荷散点图",
        section,
        "进出站总客流越高，站台人员、CO2 和人员散热上升，总冷负荷整体抬升。",
        (pd.to_numeric(df["entry_flow"], errors="coerce") + pd.to_numeric(df["exit_flow"], errors="coerce")).to_numpy(dtype=float),
        pd.to_numeric(df["total_cooling_load_kw"], errors="coerce").to_numpy(dtype=float),
        "进出站总客流",
        "kW",
    )
    scatter_plot(
        "17_室外温度与总冷负荷散点图.png",
        "室外温度与总冷负荷散点图",
        section,
        "室外温度升高会增强围护结构和新风处理负荷，是夏季高负荷的重要原因。",
        pd.to_numeric(df["outdoor_temp"], errors="coerce").to_numpy(dtype=float),
        pd.to_numeric(df["total_cooling_load_kw"], errors="coerce").to_numpy(dtype=float),
        "室外温度 / ℃",
        "kW",
    )


def generate_analysis_figures(tables: dict[str, pd.DataFrame]) -> None:
    section = "三、特征分析与聚类"
    pearson = tables["step2_pearson_features"].head(12)
    labels = pearson.iloc[:, 0].astype(str).to_list()
    vals = pd.to_numeric(pearson.iloc[:, 2], errors="coerce").fillna(0).to_list()
    horizontal_bar(
        "18_Pearson特征相关性排名图.png",
        "Pearson 特征相关性排名图",
        section,
        "显示与总冷负荷相关性最高的前12个特征，滞后负荷与客流特征占据主要位置。",
        labels,
        vals,
        "{:.3f}",
        color=PALETTE["purple"],
    )

    df = tables["raw"].copy()
    corr_cols = [
        "entry_flow",
        "exit_flow",
        "platform_passengers",
        "outdoor_temp",
        "solar_radiation",
        "platform_temp",
        "co2",
        "people_load_kw",
        "fresh_air_load_kw",
        "envelope_load_kw",
        "equipment_load_kw",
        "total_cooling_load_kw",
    ]
    names = ["进站", "出站", "站台人数", "室外温度", "太阳辐射", "站台温度", "CO2", "人员", "新风", "围护", "设备", "总负荷"]
    corr = df[corr_cols].corr(numeric_only=True).to_numpy(dtype=float)
    heatmap(
        "19_关键变量相关性热力图.png",
        "关键变量相关性热力图",
        section,
        "展示客流、环境、分项负荷与总冷负荷之间的相关结构。",
        corr,
        names,
        names,
        "相关系数范围 -1 到 1，颜色越深表示正相关越强。",
        low="#FFF7ED",
        high="#1D4ED8",
    )

    sil = tables["step2_cluster_silhouette"]
    vertical_bar(
        "20_K值轮廓系数对比图.png",
        "K 值轮廓系数对比图",
        section,
        "比较候选聚类数的平均轮廓系数，当前结果支持选择 K=4。",
        sil.iloc[:, 0].astype(str).to_list(),
        pd.to_numeric(sil.iloc[:, 1], errors="coerce").to_list(),
        "平均轮廓系数",
        color=PALETTE["green"],
    )

    copy_existing(
        "step2_daily_load_clusters.png",
        "21_主流程输出_日负荷聚类曲线图.png",
        "主流程输出：日负荷聚类曲线图",
        section,
        "复用主流程输出的聚类曲线图，体现 K-Means 聚类后各类日负荷形态。",
    )

    # Day-type monthly distribution heatmap.
    df["month"] = df["timestamp"].dt.month
    day_df = df.assign(date_only=df["timestamp"].dt.date).drop_duplicates("date_only")
    day_counts = day_df.groupby(["month", "day_type"]).size().unstack(fill_value=0)
    order = ["weekday_high", "weekday_medium", "weekend_single", "low_flow"]
    day_counts = day_counts.reindex(columns=order, fill_value=0)
    heatmap(
        "22_日类型月度分布热力图.png",
        "日类型月度分布热力图",
        section,
        "说明四类典型日样本在全年各月份中的覆盖情况，辅助解释聚类样本基础。",
        day_counts.T.to_numpy(dtype=float),
        ["工作日高峰", "工作日中等", "周末单峰", "低流量日"],
        [f"{m}月" for m in day_counts.index],
        "颜色越深表示该月份包含该类型日越多。",
        low="#F0FDFA",
        high="#0F766E",
    )


def generate_model_figures(tables: dict[str, pd.DataFrame]) -> None:
    section = "四、预测模型与模型评价"
    flow_figure(
        "23_训练验证测试集划分图.png",
        "训练-验证-测试集时间划分图",
        section,
        "展示项目按时间顺序划分数据集，避免时间序列任务中随机划分带来的未来信息泄漏。",
        [
            {"rect": (100, 300, 760, 430), "title": "训练集 70%", "body": "学习模型参数"},
            {"rect": (760, 300, 1250, 430), "title": "验证集 20%", "body": "早停与超参数校核"},
            {"rect": (1250, 300, 1680, 430), "title": "测试集 10%", "body": "最终泛化评价"},
            {"rect": (100, 550, 1680, 650), "title": "时间轴", "body": "2025-01-01 00:00 -> 2025-12-31 23:45"},
        ],
        [(0, 1), (1, 2)],
    )
    flow_figure(
        "24_LSTM输入序列窗口示意图.png",
        "LSTM 输入序列窗口示意图",
        section,
        "每个 LSTM 样本使用连续16个15分钟时间步作为输入，对应4小时历史窗口。",
        [
            {"rect": (110, 220, 330, 350), "title": "t-15", "body": "多变量特征"},
            {"rect": (390, 220, 610, 350), "title": "...", "body": "滑动窗口"},
            {"rect": (670, 220, 890, 350), "title": "t-1", "body": "多变量特征"},
            {"rect": (1040, 220, 1320, 350), "title": "LSTM", "body": "记忆短期惯性与滞后响应"},
            {"rect": (1450, 220, 1690, 350), "title": "t时刻负荷", "body": "kW"},
        ],
        [(0, 1), (1, 2), (2, 3), (3, 4)],
    )
    flow_figure(
        "25_LSTM网络结构图.png",
        "LSTM 网络结构图",
        section,
        "展示项目中两层 LSTM 结构：Sequence Input -> LSTM(96) -> LSTM(48) -> FC -> Regression。",
        [
            {"rect": (90, 250, 330, 380), "title": "Sequence Input", "body": "特征数 × 16步"},
            {"rect": (430, 250, 670, 380), "title": "LSTM Layer 1", "body": "96 hidden units\nOutputMode=sequence"},
            {"rect": (770, 250, 1010, 380), "title": "LSTM Layer 2", "body": "48 hidden units\nOutputMode=last"},
            {"rect": (1110, 250, 1350, 380), "title": "Fully Connected", "body": "1个回归输出"},
            {"rect": (1450, 250, 1690, 380), "title": "Regression", "body": "总冷负荷 / kW"},
        ],
        [(0, 1), (1, 2), (2, 3), (3, 4)],
    )
    flow_figure(
        "26_LSTM门控机制示意图.png",
        "LSTM 单元门控机制示意图",
        section,
        "用概念图解释遗忘门、输入门和输出门如何保留或更新历史负荷状态。",
        [
            {"rect": (90, 260, 350, 390), "title": "上一状态", "body": "h(t-1), C(t-1)"},
            {"rect": (460, 160, 730, 290), "title": "遗忘门 f(t)", "body": "决定保留多少历史热惯性"},
            {"rect": (460, 360, 730, 490), "title": "输入门 i(t)", "body": "吸收当前客流和环境变化"},
            {"rect": (860, 260, 1130, 390), "title": "单元状态 C(t)", "body": "更新长期状态"},
            {"rect": (1260, 260, 1530, 390), "title": "输出门 o(t)", "body": "形成当前隐含状态 h(t)"},
        ],
        [(0, 1), (0, 2), (1, 3), (2, 3), (3, 4)],
    )
    flow_figure(
        "27_BP网络结构图.png",
        "BP 神经网络结构图",
        section,
        "展示 BP 基准模型的前馈结构，用于与具备时序记忆的 LSTM 模型比较。",
        [
            {"rect": (140, 250, 430, 390), "title": "静态特征输入", "body": "客流、气象、站内环境、时间特征"},
            {"rect": (560, 250, 830, 390), "title": "隐藏层1", "body": "20个神经元"},
            {"rect": (960, 250, 1230, 390), "title": "隐藏层2", "body": "10个神经元"},
            {"rect": (1360, 250, 1630, 390), "title": "输出层", "body": "总冷负荷 / kW"},
        ],
        [(0, 1), (1, 2), (2, 3)],
    )

    metrics = tables["step3_prediction_metrics"]
    grouped_bar(
        "28_LSTM与BP预测指标对比图.png",
        "LSTM 与 BP 预测指标对比图",
        section,
        "比较 LSTM 与 BP 的 RMSE、MAE、MAPE，当前 LSTM 在主指标上明显优于静态 BP 基准。",
        ["RMSE/kW", "MAE/kW", "MAPE/%"],
        [
            ("LSTM", [float(metrics.iloc[0, 1]), float(metrics.iloc[0, 2]), float(metrics.iloc[0, 3])], PALETTE["blue"]),
            ("BP", [float(metrics.iloc[1, 1]), float(metrics.iloc[1, 2]), float(metrics.iloc[1, 3])], PALETTE["orange"]),
        ],
        "误差值",
    )

    copy_existing("step3_lstm_prediction.png", "29_主流程输出_LSTM预测曲线图.png", "主流程输出：LSTM 预测曲线图", section, "复用主流程输出图，用于展示测试集真实值与 LSTM 预测值的贴合程度。")
    copy_existing("step3_bp_prediction.png", "30_主流程输出_BP预测曲线图.png", "主流程输出：BP 预测曲线图", section, "复用主流程输出图，用于与 LSTM 测试集预测效果进行对比。")
    copy_existing("step3_model_rmse_comparison.png", "31_主流程输出_RMSE对比图.png", "主流程输出：模型 RMSE 对比图", section, "复用主流程输出的模型误差对比图。")

    ablation = tables["step3_cluster_ablation_metrics"]
    grouped_bar(
        "32_聚类标签消融实验图.png",
        "聚类标签消融实验图",
        section,
        "比较 LSTM 是否加入典型日聚类标签后的误差变化，用于说明聚类特征的增益。",
        ["RMSE", "MAE", "MAPE"],
        [
            ("不含聚类标签", [float(ablation.iloc[0, 3]), float(ablation.iloc[0, 4]), float(ablation.iloc[0, 5])], PALETTE["slate"]),
            ("含聚类标签", [float(ablation.iloc[1, 3]), float(ablation.iloc[1, 4]), float(ablation.iloc[1, 5])], PALETTE["green"]),
        ],
        "误差值",
    )

    fair = tables["step3_fair_model_comparison"]
    labels = fair.iloc[:, 0].astype(str).to_list()
    rmse = pd.to_numeric(fair.iloc[:, 4], errors="coerce").to_list()
    horizontal_bar(
        "33_公平输入模型比较图.png",
        "公平输入模型比较图",
        section,
        "在含滞后和不含滞后两类特征组下比较 LSTM 与 BP，帮助避免把输入信息量差异误读为模型结构差异。",
        labels,
        rmse,
        "{:.2f}",
        color=PALETTE["teal"],
    )


def generate_optimization_figures(tables: dict[str, pd.DataFrame]) -> None:
    section = "五、容量优化与工程评价"
    quant = tables["step3_subsystem_demand_quantiles"]
    subs = quant.iloc[:, 0].astype(str).to_list()
    grouped_bar(
        "34_子系统需求分位数对比图.png",
        "子系统需求分位数对比图",
        section,
        "展示总冷负荷、冷机、风机、水泵和 AHU 在 P50/P95/P99 下的容量需求差异。",
        subs[:4],
        [
            ("P50", pd.to_numeric(quant.iloc[:4, 2], errors="coerce").to_list(), PALETTE["cyan"]),
            ("P95", pd.to_numeric(quant.iloc[:4, 4], errors="coerce").to_list(), PALETTE["orange"]),
            ("P99", pd.to_numeric(quant.iloc[:4, 5], errors="coerce").to_list(), PALETTE["red"]),
        ],
        "kW",
        "AHU 另以风量单位计，单独绘制。",
    )
    ahu_vals = [float(quant.iloc[4, 2]), float(quant.iloc[4, 4]), float(quant.iloc[4, 5]), float(quant.iloc[4, 6])]
    vertical_bar(
        "35_AHU风量需求分位数图.png",
        "AHU 风量需求分位数图",
        section,
        "展示 AHU 风量需求从典型情景到极端情景的变化。",
        ["P50", "P95", "P99", "最大值"],
        ahu_vals,
        "m3/h",
        color=PALETTE["green"],
    )

    scenario = tables["step3_scenario_demand"]
    grouped_bar(
        "36_容量设计情景需求图.png",
        "容量设计情景需求图",
        section,
        "比较典型、峰值、极端情景下的子系统需求，是容量优化输入的核心依据。",
        scenario.iloc[:, 0].astype(str).to_list(),
        [
            ("总冷负荷", pd.to_numeric(scenario.iloc[:, 2], errors="coerce").to_list(), PALETTE["blue"]),
            ("冷机需求", pd.to_numeric(scenario.iloc[:, 3], errors="coerce").to_list(), PALETTE["orange"]),
            ("风机需求", pd.to_numeric(scenario.iloc[:, 4], errors="coerce").to_list(), PALETTE["green"]),
            ("水泵需求", pd.to_numeric(scenario.iloc[:, 5], errors="coerce").to_list(), PALETTE["red"]),
        ],
        "kW",
    )

    flow_figure(
        "37_NSGAII_TOPSIS优化流程图.png",
        "NSGA-II 与 TOPSIS 优化流程图",
        section,
        "展示容量优化从候选设备编码、约束检查、目标计算到 TOPSIS 推荐方案的决策过程。",
        [
            {"rect": (80, 200, 340, 320), "title": "设备候选库", "body": "冷机/风机/水泵/AHU规格与台数"},
            {"rect": (430, 200, 690, 320), "title": "整数编码", "body": "8维决策向量"},
            {"rect": (780, 200, 1040, 320), "title": "约束检查", "body": "容量、安全系数、PLR、匹配比例"},
            {"rect": (1130, 200, 1390, 320), "title": "双目标评价", "body": "生命周期成本 + 综合冗余率"},
            {"rect": (1480, 200, 1710, 320), "title": "Pareto 解集", "body": "非劣候选方案"},
            {"rect": (780, 520, 1040, 640), "title": "TOPSIS排序", "body": "成本0.55 / 冗余0.45"},
            {"rect": (1130, 520, 1390, 640), "title": "推荐方案", "body": "输出容量与评价表"},
        ],
        [(0, 1), (1, 2), (2, 3), (3, 4), (4, 5), (5, 6)],
    )

    copy_existing("step4_pareto_front.png", "38_主流程输出_Pareto前沿图.png", "主流程输出：Pareto 前沿图", section, "复用主流程 Pareto 前沿图，展示生命周期成本与综合冗余率之间的权衡关系。")

    topsis = tables["step4_topsis_ranking"].head(8)
    horizontal_bar(
        "39_TOPSIS前八候选方案得分图.png",
        "TOPSIS 前八候选方案得分图",
        section,
        "展示 Pareto 候选方案的 TOPSIS 贴近度排序，分数越高表示越接近低成本、低冗余理想解。",
        [f"方案{i+1}" for i in range(len(topsis))],
        pd.to_numeric(topsis.iloc[:, 2], errors="coerce").to_list(),
        "{:.3f}",
        color=PALETTE["purple"],
    )

    eval_df = tables["step4_scheme_evaluation"]
    groups = ["总制冷容量", "生命周期成本/万", "综合冗余率/%", "年能耗/万kWh"]
    base = [
        float(eval_df.iloc[0, 1]),
        float(eval_df.iloc[0, 2]) / 10000,
        float(eval_df.iloc[0, 3]) * 100,
        float(eval_df.iloc[0, 5]) / 10000,
    ]
    opt = [
        float(eval_df.iloc[1, 1]),
        float(eval_df.iloc[1, 2]) / 10000,
        float(eval_df.iloc[1, 3]) * 100,
        float(eval_df.iloc[1, 5]) / 10000,
    ]
    grouped_bar(
        "40_基准方案与优化方案综合对比图.png",
        "基准方案与优化方案综合对比图",
        section,
        "从容量、成本、冗余率和年能耗四个维度比较基准方案与优化方案。",
        groups,
        [("基准方案", base, PALETTE["slate"]), ("优化方案", opt, PALETTE["blue"])],
        "指标值",
    )
    copy_existing("step4_scheme_comparison.png", "41_主流程输出_方案对比图.png", "主流程输出：方案对比图", section, "复用主流程输出的方案对比图，作为论文原始结果图件之一。")

    red = tables["step4_subsystem_redundancy"]
    subs = red.iloc[:, 0].astype(str).to_list()
    grouped_bar(
        "42_子系统冗余率对比图.png",
        "子系统冗余率对比图",
        section,
        "比较优化前后冷机、风机、水泵和 AHU 的容量冗余率，体现容量优化的直接效果。",
        subs,
        [
            ("基准冗余率", (pd.to_numeric(red.iloc[:, 4], errors="coerce") * 100).to_list(), PALETTE["orange"]),
            ("优化冗余率", (pd.to_numeric(red.iloc[:, 5], errors="coerce") * 100).to_list(), PALETTE["green"]),
        ],
        "%",
    )

    cons = tables["step4_engineering_constraint_check"]
    names = cons.iloc[:, 0].astype(str).to_list()
    actual = pd.to_numeric(cons.iloc[:, 1], errors="coerce").fillna(0).to_list()
    normalized = []
    for i, row in cons.iterrows():
        val = float(pd.to_numeric(row.iloc[1], errors="coerce"))
        lo = pd.to_numeric(row.iloc[2], errors="coerce")
        hi = pd.to_numeric(row.iloc[3], errors="coerce")
        if pd.notna(hi) and hi != 0:
            normalized.append(min(val / float(hi), 1.2))
        elif pd.notna(lo) and lo != 0:
            normalized.append(min(val / float(lo), 1.8))
        else:
            normalized.append(1.0 if val >= 0 else 0.0)
    horizontal_bar(
        "43_工程约束校核通过情况图.png",
        "工程约束校核通过情况图",
        section,
        "将工程约束的实际值相对要求归一化展示，辅助说明推荐方案满足容量和匹配约束。",
        names,
        normalized,
        "{:.2f}",
        color=PALETTE["green"],
    )

    sens = tables["step4_sensitivity_analysis"]
    risk_map = {"低": 1, "中": 2, "高": 3, "low": 1, "medium": 2, "high": 3}
    risks = [risk_map.get(str(x), 2) for x in sens.iloc[:, 4].to_list()]
    vertical_bar(
        "44_敏感性风险等级图.png",
        "敏感性风险等级图",
        section,
        "展示设计分位数、安全系数、TOPSIS权重、电价和预测误差等因素的风险等级。",
        sens.iloc[:, 0].astype(str).to_list(),
        risks,
        "风险等级 1低/2中/3高",
        color=PALETTE["red"],
    )

    boundary = tables["step4_engineering_boundary"]
    rows = []
    for _, row in boundary.iterrows():
        rows.append([str(row.iloc[0]), str(row.iloc[1]), str(row.iloc[2])])
    table_figure(
        "45_工程边界与详细设计补充图.png",
        "工程边界与详细设计补充图",
        section,
        "把论文级方法验证与施工图级设计所需补充条件区分开，避免过度外推。",
        ["工程边界", "当前处理方式", "详细设计所需补充"],
        rows[:6],
    )

    linkage = tables["step4_research_linkage"]
    rows = []
    for _, row in linkage.iterrows():
        rows.append([str(row.iloc[0]), str(row.iloc[1]), str(row.iloc[2]), str(row.iloc[3])])
    table_figure(
        "46_研究链路证据矩阵图.png",
        "研究链路证据矩阵图",
        section,
        "汇总特征分析、聚类分析、预测结果和容量优化之间的证据链路。",
        ["研究链路", "上游输出", "下游用途", "证据指标"],
        rows,
    )


def set_run_font(run, east_asia: str = "宋体", latin: str = "Times New Roman") -> None:
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def set_paragraph_font(paragraph, size: int = 10.5, color: str | None = None, bold: bool = False) -> None:
    for run in paragraph.runs:
        set_run_font(run)
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor(*hex_to_rgb(color))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill.replace("#", ""))
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run)
    run.font.size = Pt(9)
    run.bold = bold


def build_docx() -> None:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Times New Roman"
        styles[name]._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        styles[name].font.color.rgb = RGBColor(*hex_to_rgb(PALETTE["ink"]))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("地铁车站环控系统容量优化论文图集")
    set_run_font(r, east_asia="黑体")
    r.font.size = Pt(22)
    r.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("绘图尝试版：研究流程、负荷特性、预测模型、容量优化与工程校核")
    set_run_font(r)
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(*hex_to_rgb(PALETTE["muted"]))

    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.style = "Table Grid"
    items = [
        ("项目目录", str(ROOT)),
        ("图件数量", f"{len(figures)} 张"),
        ("主要数据源", "data/fuzhou_metro_dongjiekou_2025.csv；output/tables/*.csv；output/figures/*.png"),
        ("说明", "本图集优先使用项目真实数据和主流程输出；每张图均附图名和用途说明。"),
    ]
    for i, (k, v) in enumerate(items):
        set_cell_shading(meta.cell(i, 0), "#E8F0FE")
        set_cell_text(meta.cell(i, 0), k, bold=True)
        set_cell_text(meta.cell(i, 1), v)
        meta.cell(i, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    p = doc.add_paragraph()
    r = p.add_run("使用建议：")
    set_run_font(r, east_asia="黑体")
    r.bold = True
    p.add_run(" 论文正文中优先选用每章最能支撑结论的核心图；附录或答辩材料可使用完整图集。")
    set_paragraph_font(p, 10.5)
    doc.add_page_break()

    doc.add_heading("图件目录", level=1)
    by_section: dict[str, list[FigureItem]] = {}
    for item in figures:
        by_section.setdefault(item.section, []).append(item)
    for sec, items_in_sec in by_section.items():
        p = doc.add_paragraph()
        r = p.add_run(sec)
        set_run_font(r, east_asia="黑体")
        r.bold = True
        r.font.size = Pt(12)
        for item in items_in_sec:
            p = doc.add_paragraph(style=None)
            p.paragraph_format.left_indent = Cm(0.6)
            r = p.add_run(f"图 {item.no}  {item.title}")
            set_run_font(r)
            r.font.size = Pt(10)
    doc.add_page_break()

    current_section = None
    for item in figures:
        if item.section != current_section:
            if current_section is not None:
                doc.add_page_break()
            current_section = item.section
            doc.add_heading(current_section, level=1)

        h = doc.add_paragraph()
        h.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = h.add_run(f"图 {item.no}  {item.title}")
        set_run_font(r, east_asia="黑体")
        r.bold = True
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(*hex_to_rgb(PALETTE["ink"]))

        img_p = doc.add_paragraph()
        img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_p.add_run().add_picture(str(item.path), width=Inches(6.55))

        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(f"图 {item.no}  {item.title}")
        set_run_font(r, east_asia="宋体")
        r.font.size = Pt(10)
        r.bold = True

        note = doc.add_paragraph()
        note.paragraph_format.space_after = Pt(8)
        note.paragraph_format.line_spacing = 1.15
        r = note.add_run("图示说明：")
        set_run_font(r, east_asia="黑体")
        r.bold = True
        r.font.size = Pt(10.5)
        r2 = note.add_run(item.note)
        set_run_font(r2)
        r2.font.size = Pt(10.5)

    doc.save(DOCX_PATH)


def main() -> None:
    ensure_dirs()
    tables = load_data()
    generate_flow_diagrams()
    generate_data_figures(tables)
    generate_analysis_figures(tables)
    generate_model_figures(tables)
    generate_optimization_figures(tables)
    build_docx()
    print(f"Generated {len(figures)} figures")
    print(DOCX_PATH)


if __name__ == "__main__":
    main()
